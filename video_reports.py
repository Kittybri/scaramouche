import asyncio
import json
import os
import sys
import tempfile
from pathlib import Path
from urllib.parse import urljoin

from character_vision import ask_character_bot


ROOT_DIR = Path(__file__).resolve().parent
DOWNLOADS_DIR = ROOT_DIR.parent
SCARA_TEMPLATE_DIR = DOWNLOADS_DIR / "scaramouche_blender_templates"
WANDERER_TEMPLATE_DIR = DOWNLOADS_DIR / "wanderer_blender_templates"
BLENDER_EXE = Path(r"C:\Program Files\Blender Foundation\Blender 4.2\blender.exe")
SCARA_VIDEO_SCRIPT = SCARA_TEMPLATE_DIR / "notes_to_scaramouche_video.py"
WANDERER_VIDEO_SCRIPT = WANDERER_TEMPLATE_DIR / "notes_to_wanderer_video.py"
DUO_VIDEO_SCRIPT = SCARA_TEMPLATE_DIR / "notes_to_duo_debate_video.py"
VIDEO_RENDER_TIMEOUT_S = int(os.getenv("VIDEO_RENDER_TIMEOUT_S", "1200") or "1200")
VIDEO_RENDER_WIDTH = int(os.getenv("VIDEO_RENDER_WIDTH", "960") or "960")
VIDEO_RENDER_HEIGHT = int(os.getenv("VIDEO_RENDER_HEIGHT", "540") or "540")
VIDEO_RENDER_FPS = int(os.getenv("VIDEO_RENDER_FPS", "8") or "8")
VIDEO_RENDER_ENGINE = os.getenv("VIDEO_RENDER_ENGINE", "BLENDER_WORKBENCH").strip() or "BLENDER_WORKBENCH"
VIDEO_RENDER_MODE = (os.getenv("VIDEO_RENDER_MODE", "auto") or "auto").strip().lower()
VIDEO_RENDER_BASE_URL = (os.getenv("VIDEO_RENDER_BASE_URL") or "").strip().rstrip("/")
VIDEO_RENDER_PRIMARY_URL = (os.getenv("VIDEO_RENDER_PRIMARY_URL") or "").strip().rstrip("/")
VIDEO_RENDER_FALLBACK_URL = (os.getenv("VIDEO_RENDER_FALLBACK_URL") or "").strip().rstrip("/")
VIDEO_RENDER_BASE_URLS = (
    os.getenv("VIDEO_RENDER_BASE_URLS")
    or os.getenv("VIDEO_RENDER_REMOTE_URLS")
    or ""
).strip()
VIDEO_RENDER_SECRET = (
    os.getenv("VIDEO_RENDER_SECRET")
    or os.getenv("VIDEO_RENDER_SHARED_SECRET")
    or ""
).strip()
VIDEO_RENDER_POLL_S = max(1.0, float(os.getenv("VIDEO_RENDER_POLL_S", "4") or "4"))
VIDEO_RENDER_REMOTE_TIMEOUT_S = int(
    os.getenv("VIDEO_RENDER_REMOTE_TIMEOUT_S", str(VIDEO_RENDER_TIMEOUT_S + 180))
    or str(VIDEO_RENDER_TIMEOUT_S + 180)
)


def _render_env_overrides() -> dict[str, str]:
    overrides: dict[str, str] = {}
    for name in (
        "FISH_AUDIO_API_KEY",
        "SCARAMOUCHE_FISH_VOICE_ID",
        "SCARAMOUCHE_VOICE_ID",
        "WANDERER_FISH_VOICE_ID",
        "WANDERER_VOICE_ID",
    ):
        value = (os.getenv(name) or "").strip()
        if value:
            overrides[name] = value
    return overrides


def _video_root() -> Path:
    configured = (os.getenv("VIDEO_RENDER_ROOT") or "").strip()
    if configured:
        path = Path(configured)
    else:
        path = Path(tempfile.gettempdir()) / "scara_wanderer_video_jobs"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _script_for_bot(bot_name: str) -> Path:
    key = (bot_name or "").strip().lower()
    if key == "scaramouche":
        return SCARA_VIDEO_SCRIPT
    if key == "wanderer":
        return WANDERER_VIDEO_SCRIPT
    raise ValueError(f"Unsupported bot video renderer: {bot_name}")


def _remote_requested() -> bool:
    if VIDEO_RENDER_MODE == "remote":
        return True
    if VIDEO_RENDER_MODE == "local":
        return False
    return bool(_remote_base_urls())


def _local_requested() -> bool:
    if VIDEO_RENDER_MODE == "remote":
        return False
    return True


def _remote_headers() -> dict[str, str]:
    headers = {"Accept": "application/json"}
    if VIDEO_RENDER_SECRET:
        headers["X-Render-Secret"] = VIDEO_RENDER_SECRET
    return headers


def _remote_base_urls() -> list[str]:
    values: list[str] = []
    for item in [VIDEO_RENDER_PRIMARY_URL, VIDEO_RENDER_BASE_URL]:
        if item:
            values.append(item)
    if VIDEO_RENDER_BASE_URLS:
        for raw in VIDEO_RENDER_BASE_URLS.replace("\n", ",").split(","):
            cleaned = raw.strip().rstrip("/")
            if cleaned:
                values.append(cleaned)
    if VIDEO_RENDER_FALLBACK_URL:
        values.append(VIDEO_RENDER_FALLBACK_URL)

    deduped: list[str] = []
    seen: set[str] = set()
    for value in values:
        if value not in seen:
            deduped.append(value)
            seen.add(value)
    return deduped


def _aiohttp():
    import aiohttp

    return aiohttp


def _make_job_dir(prefix: str) -> Path:
    path = _video_root() / prefix
    suffix = 1
    final = path
    while final.exists():
        suffix += 1
        final = _video_root() / f"{prefix}-{suffix}"
    final.mkdir(parents=True, exist_ok=True)
    return final


def video_renderer_available(bot_name: str | None = None, *, duo: bool = False) -> tuple[bool, str]:
    if _remote_requested():
        if not _remote_base_urls():
            return False, "VIDEO_RENDER_BASE_URL or VIDEO_RENDER_BASE_URLS is not set for remote rendering."
        return True, ""

    if not _local_requested():
        return False, "No video renderer mode is enabled."

    if not BLENDER_EXE.exists():
        return False, (
            f"Blender is missing at {BLENDER_EXE}. "
            "Set VIDEO_RENDER_BASE_URL to a remote render worker if Railway should call another machine."
        )

    scripts = [DUO_VIDEO_SCRIPT, SCARA_VIDEO_SCRIPT, WANDERER_VIDEO_SCRIPT] if duo else [_script_for_bot(bot_name or "scaramouche")]
    for script in scripts:
        if not script.exists():
            return False, f"Missing renderer script: {script}"
    return True, ""


def _can_failover_remote(exc: Exception | str) -> bool:
    text = str(exc).lower()
    return any(
        token in text
        for token in [
            "cannot connect",
            "name or service not known",
            "temporary failure in name resolution",
            "status check failed (530)",
            "status check failed (502)",
            "status check failed (503)",
            "status check failed (504)",
            "timed out",
            "timeout",
            "connection reset",
            "connection refused",
            "cloudflare tunnel error",
        ]
    )


async def _run_render_script(
    script_path: Path,
    notes_text: str,
    *,
    title: str = "",
    width: int = VIDEO_RENDER_WIDTH,
    height: int = VIDEO_RENDER_HEIGHT,
    fps: int = VIDEO_RENDER_FPS,
    engine: str = VIDEO_RENDER_ENGINE,
) -> dict:
    if not script_path.exists():
        raise FileNotFoundError(f"Render script not found: {script_path}")

    output_dir = _make_job_dir(script_path.stem)
    notes_path = output_dir / "notes.txt"
    notes_path.write_text(notes_text, encoding="utf-8")

    cmd = [
        sys.executable,
        str(script_path),
        str(notes_path),
        "--width",
        str(width),
        "--height",
        str(height),
        "--fps",
        str(fps),
        "--engine",
        engine,
        "--output-dir",
        str(output_dir),
    ]
    if title.strip():
        cmd.extend(["--title", title.strip()])

    run_env = os.environ.copy()
    for key, value in _render_env_overrides().items():
        run_env[key] = value

    proc = await asyncio.create_subprocess_exec(
        *cmd,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.STDOUT,
        cwd=str(script_path.parent),
        env=run_env,
    )
    try:
        stdout, _ = await asyncio.wait_for(proc.communicate(), timeout=VIDEO_RENDER_TIMEOUT_S)
    except asyncio.TimeoutError:
        proc.kill()
        raise RuntimeError("Video rendering timed out.")

    combined = (stdout or b"").decode("utf-8", errors="ignore")
    summary_path = output_dir / "run_summary.json"
    if proc.returncode != 0:
        tail = "\n".join(combined.splitlines()[-20:]) or "no output"
        raise RuntimeError(f"Video rendering failed.\n{tail}")
    if not summary_path.exists():
        tail = "\n".join(combined.splitlines()[-20:]) or "no output"
        raise RuntimeError(f"Render finished without a summary file.\n{tail}")
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    summary["_output_dir"] = str(output_dir)
    summary["_logs"] = combined[-4000:]
    return summary


async def _submit_remote_job_once(base_url: str, payload: dict) -> dict:
    aiohttp = _aiohttp()
    timeout = aiohttp.ClientTimeout(total=45)
    async with aiohttp.ClientSession(timeout=timeout) as session:
        async with session.post(
            f"{base_url}/jobs",
            json=payload,
            headers=_remote_headers(),
        ) as resp:
            body_text = await resp.text()
            if resp.status not in {200, 202}:
                raise RuntimeError(f"Remote render worker rejected the job ({resp.status}): {body_text[:400]}")
            body = json.loads(body_text or "{}")
            job_id = (body.get("job_id") or "").strip()
            if not job_id:
                raise RuntimeError("Remote render worker did not return a job id.")

        deadline = asyncio.get_running_loop().time() + VIDEO_RENDER_REMOTE_TIMEOUT_S
        while asyncio.get_running_loop().time() < deadline:
            await asyncio.sleep(VIDEO_RENDER_POLL_S)
            async with session.get(
                f"{base_url}/jobs/{job_id}",
                headers=_remote_headers(),
            ) as resp:
                body_text = await resp.text()
                if resp.status != 200:
                    raise RuntimeError(f"Remote render status check failed ({resp.status}): {body_text[:400]}")
                body = json.loads(body_text or "{}")
                status = (body.get("status") or "").strip().lower()
                if status in {"queued", "running"}:
                    continue
                if status in {"failed", "error"}:
                    raise RuntimeError(body.get("error") or "Remote render worker failed the job.")
                if status in {"done", "completed", "ready"}:
                    summary = body.get("summary") or {}
                    download_url = (body.get("download_url") or "").strip()
                    if download_url:
                        summary["final_video_url"] = urljoin(f"{base_url}/", download_url.lstrip("/"))
                    summary["final_video_name"] = body.get("final_video_name") or summary.get("final_video_name") or "presentation.mp4"
                    summary["_remote_job_id"] = job_id
                    summary["_remote_base_url"] = base_url
                    return summary
                raise RuntimeError(f"Remote render worker returned an unknown status: {status or 'missing'}")

    raise RuntimeError("Remote video rendering timed out while waiting for the worker.")


async def _submit_remote_job(payload: dict) -> dict:
    base_urls = _remote_base_urls()
    if not base_urls:
        raise RuntimeError("VIDEO_RENDER_BASE_URL or VIDEO_RENDER_BASE_URLS is not set.")

    errors: list[str] = []
    for index, base_url in enumerate(base_urls, start=1):
        try:
            return await _submit_remote_job_once(base_url, payload)
        except Exception as exc:
            errors.append(f"[{index}] {base_url} -> {exc}")
            if not _can_failover_remote(exc) or index == len(base_urls):
                break

    raise RuntimeError("Remote render failed across all configured workers.\n" + "\n".join(errors[:4]))


async def render_teaching_video(bot_name: str, notes_text: str, *, title: str = "") -> dict:
    if _remote_requested():
        return await _submit_remote_job(
            {
                "job_type": "teaching",
                "bot_name": (bot_name or "").strip().lower(),
                "notes_text": notes_text,
                "title": title,
                "env_overrides": _render_env_overrides(),
            }
        )
    return await _run_render_script(_script_for_bot(bot_name), notes_text, title=title)


async def render_duo_debate_video(notes_text: str, *, title: str = "") -> dict:
    if _remote_requested():
        return await _submit_remote_job(
            {
                "job_type": "duo",
                "bot_name": "scaramouche",
                "notes_text": notes_text,
                "title": title,
                "env_overrides": _render_env_overrides(),
            }
        )
    return await _run_render_script(DUO_VIDEO_SCRIPT, notes_text, title=title)


async def fetch_rendered_video_bytes(summary: dict) -> tuple[bytes, str]:
    final_video = (summary or {}).get("final_video", "")
    if final_video and os.path.exists(final_video):
        path = Path(final_video)
        return path.read_bytes(), path.name

    final_video_url = (summary or {}).get("final_video_url", "")
    if final_video_url:
        aiohttp = _aiohttp()
        timeout = aiohttp.ClientTimeout(total=VIDEO_RENDER_REMOTE_TIMEOUT_S)
        async with aiohttp.ClientSession(timeout=timeout) as session:
            async with session.get(final_video_url, headers=_remote_headers()) as resp:
                if resp.status != 200:
                    text = await resp.text()
                    raise RuntimeError(f"Remote video download failed ({resp.status}): {text[:300]}")
                data = await resp.read()
        filename = (summary or {}).get("final_video_name") or Path(final_video_url.split("?", 1)[0]).name or "presentation.mp4"
        return data, filename

    raise RuntimeError("The render finished, but there is no local file or remote download URL.")


async def download_attachment_bytes(attachment) -> bytes:
    aiohttp = _aiohttp()
    async with aiohttp.ClientSession() as session:
        async with session.get(attachment.url) as resp:
            if resp.status != 200:
                raise RuntimeError(f"Attachment download failed with HTTP {resp.status}")
            return await resp.read()


async def extract_teaching_material(attachment, *, topic: str = "", bot_name: str = "wanderer") -> str:
    material_content = ""
    if attachment:
        content_type = (attachment.content_type or "").lower()
        filename = (attachment.filename or "attachment").lower()
        file_bytes = await download_attachment_bytes(attachment)

        if "pdf" in content_type or filename.endswith(".pdf"):
            try:
                import io as _io
                import pdfplumber

                with pdfplumber.open(_io.BytesIO(file_bytes)) as pdf:
                    material_content = "\n".join(page.extract_text() or "" for page in pdf.pages)[:8000]
            except Exception:
                try:
                    import io as _io
                    import PyPDF2

                    reader = PyPDF2.PdfReader(_io.BytesIO(file_bytes))
                    material_content = "\n".join(page.extract_text() or "" for page in reader.pages)[:8000]
                except Exception as exc:
                    raise RuntimeError(f"Couldn't read the PDF: {exc}") from exc

        elif "image" in content_type or filename.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
            try:
                material_content = ask_character_bot(
                    bot_name,
                    "Extract all educational content visible in this image. Include every concept, formula, definition, diagram label, and key point.",
                    image_url=attachment.url,
                )
            except Exception as exc:
                raise RuntimeError(f"Couldn't read the image: {exc}") from exc

        elif "text" in content_type or filename.endswith((".txt", ".md", ".csv")):
            try:
                material_content = file_bytes.decode("utf-8", errors="ignore")[:6000]
            except Exception as exc:
                raise RuntimeError(f"Couldn't read the text file: {exc}") from exc

        elif filename.endswith((".pptx", ".ppt")):
            try:
                import io as _io
                from pptx import Presentation as _Prs

                prs = _Prs(_io.BytesIO(file_bytes))
                parts = []
                for index, slide in enumerate(prs.slides, start=1):
                    slide_texts = []
                    for shape in slide.shapes:
                        if hasattr(shape, "text") and shape.text.strip():
                            slide_texts.append(shape.text.strip())
                    if slide_texts:
                        parts.append(f"[Slide {index}]\n" + "\n".join(slide_texts))
                material_content = "\n\n".join(parts)[:6000]
            except Exception as exc:
                raise RuntimeError(f"Couldn't read the PowerPoint: {exc}") from exc

        elif filename.endswith((".docx", ".doc")):
            try:
                import io as _io
                import docx as _docx

                doc = _docx.Document(_io.BytesIO(file_bytes))
                parts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            parts.append(row_text)
                material_content = "\n".join(parts)[:6000]
            except Exception as exc:
                raise RuntimeError(f"Couldn't read the Word document: {exc}") from exc

        else:
            raise RuntimeError("I can read PDFs, images, PowerPoint files, Word documents, and text files for video lessons.")

    full_material = f"Topic: {topic}\n\n{material_content}".strip() if topic else material_content.strip()
    return full_material.strip()


def build_weather_video_notes(place: str, weather_data: dict, news_results: list[dict], *, duo: bool = False) -> str:
    title = f"{place} Weather Report"
    intro = "A duo weather news report." if duo else "A weather news report."
    lines = [
        f"# {title}",
        intro,
        "",
        "## Current conditions",
        f"- Place: {place}",
        f"- Forecast: {weather_data.get('forecast') or 'Unavailable'}",
    ]
    temp = weather_data.get("temperature")
    unit = weather_data.get("temperature_unit") or "F"
    if temp is not None:
        lines.append(f"- Temperature: {temp} {unit}")
    wind_speed = weather_data.get("wind_speed") or ""
    wind_direction = weather_data.get("wind_direction") or ""
    if wind_speed or wind_direction:
        lines.append(f"- Wind: {wind_direction} {wind_speed}".strip())
    precip = weather_data.get("precipitation")
    if precip is not None:
        lines.append(f"- Chance of precipitation: {precip}%")

    if news_results:
        lines.extend(["", "## Recent weather headlines"])
        for index, item in enumerate(news_results[:3], start=1):
            title_text = (item.get("title") or "Weather headline").strip()
            snippet = (item.get("snippet") or "").strip()
            url = (item.get("url") or "").strip()
            lines.append(f"- Headline {index}: {title_text}")
            if snippet:
                lines.append(f"  Detail: {snippet}")
            if url:
                lines.append(f"  Source: {url}")
    return "\n".join(lines).strip()
