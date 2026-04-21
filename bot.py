"""
Scaramouche Bot — The Balladeer v7 (Stability Release)
Full stability pass: every task, loop, command, and event is wrapped
in try/except. Nothing can crash the bot. Errors are logged and ignored.
"""

from __future__ import annotations

import discord
from discord import app_commands
from discord.ext import commands, tasks
from groq import Groq
import os, re, random, asyncio, io, time, json, traceback
from urllib.parse import quote_plus, urlencode
from datetime import datetime
from zoneinfo import ZoneInfo
from dotenv import load_dotenv
from memory import Memory
from voice_handler import get_audio
from character_vision import ask_character_bot, vision_exhausted_remaining, vision_is_exhausted
from face_memory import (
    enroll_face_profile,
    enroll_face_profile_from_frames,
    face_support_ready,
    is_face_check_request,
    is_face_enroll_request,
    match_face,
    match_face_frames,
)
from grounded_search import (
    extract_urls,
    fetch_url_preview,
    format_search_context,
    format_search_sources,
    format_url_preview_context,
    search_web,
)
from google_docs_bridge import (
    fetch_google_doc,
    google_docs_ready,
    overwrite_google_doc,
    service_account_email,
)
from memory_rebuild import collect_rebuild_records, user_can_manage_rebuild
from video_reports import (
    build_weather_video_notes,
    extract_teaching_material,
    fetch_rendered_video_bytes,
    remote_fix_google_doc,
    render_duo_debate_video,
    render_teaching_video,
    video_renderer_available,
)
from anti_repeat import (
    build_prompt_guard,
    detect_opening_phrase,
    diversify_reply,
    fallback_reply,
    get_runtime_recent,
    looks_repetitive,
    merge_recent_messages,
    pick_fresh_option,
    replace_opening_phrase,
    remember_output,
)
from relationship_engine import (
    RARE_PHRASES,
    analyze_style_deltas,
    apply_style_deltas,
    callback_relevant,
    compute_bot_stage,
    compute_emotional_arc,
    describe_bot_relationship,
    describe_contradictory_memory,
    describe_conflict_aftermath,
    describe_conflict_followup,
    describe_dual_verdict_mode,
    describe_emotional_layers,
    describe_emotional_event,
    describe_emotional_arc,
    describe_arc_unlocks,
    describe_campaign_npcs,
    describe_duo_scene_stage,
    describe_evidence_locker,
    describe_live_world_context,
    describe_lore_hook,
    describe_private_confession_scene,
    describe_private_opinion_context,
    describe_relationship_unlock_scene,
    describe_silent_presence,
    describe_triangle_jealousy,
    describe_specific_lore_tree,
    describe_relationship_progression,
    describe_scenario_context,
    describe_scene_state,
    describe_speech_drift,
    describe_topic_profile,
    detect_emotional_triggers,
    detect_banter_theme,
    detect_conflict_signal,
    detect_intervention_reason,
    detect_scenario,
    detect_topics,
    detect_repair_signal,
    extract_continuity_hooks,
    infer_duo_ending_tag,
    extract_memory_events,
    extract_callback_candidate,
    infer_scene_update,
    infer_bot_relation_deltas,
    progression_milestone_note,
    relationship_milestone_note,
)

load_dotenv()

DISCORD_TOKEN      = os.getenv("DISCORD_TOKEN","")
GROQ_API_KEY       = os.getenv("GROQ_API_KEY","")
GROQ_API_KEY_2     = os.getenv("GROQ_API_KEY_2","")
GROQ_API_KEY_3     = os.getenv("GROQ_API_KEY_3","")
FISH_AUDIO_API_KEY = os.getenv("FISH_AUDIO_API_KEY","")
WEATHER_API_KEY    = os.getenv("WEATHER_API_KEY","")
NWS_USER_AGENT     = os.getenv("NWS_USER_AGENT","scara-wanderer-bots/1.0 (contact: local-use)")
OWNER_ID           = int(os.getenv("OWNER_ID","0") or "0")
PARTNER_BOT_ID     = int(os.getenv("PARTNER_BOT_ID","0") or "0")  # Wanderer bot ID
PARTNER_INVITE_PERMISSIONS = int(os.getenv("PARTNER_BOT_PERMISSIONS", "8") or "8")
PARTNER_INVITE_SCOPES = os.getenv("PARTNER_BOT_SCOPES", "bot applications.commands").strip() or "bot applications.commands"
PARTNER_CLIENT_ID_OVERRIDE = (os.getenv("WANDERER_CLIENT_ID") or os.getenv("PARTNER_CLIENT_ID") or "").strip()
GROQ_EXHAUSTED_SILENCE_S = int(os.getenv("GROQ_EXHAUSTED_SILENCE_S", "600") or "600")
PROVIDER_PAUSE_COOLDOWN_S = int(os.getenv("PROVIDER_PAUSE_COOLDOWN_S", "120") or "120")
CHANNEL_CONTEXT_LIMIT_DIRECT = int(os.getenv("CHANNEL_CONTEXT_LIMIT_DIRECT", "20") or "20")
CHANNEL_CONTEXT_LIMIT_AMBIENT = int(os.getenv("CHANNEL_CONTEXT_LIMIT_AMBIENT", "8") or "8")
CHANNEL_CONTEXT_LIMIT_DM = int(os.getenv("CHANNEL_CONTEXT_LIMIT_DM", "16") or "16")
CHANNEL_CONTEXT_MESSAGE_CHARS = int(os.getenv("CHANNEL_CONTEXT_MESSAGE_CHARS", "110") or "110")
HISTORY_LIMIT_DIRECT = int(os.getenv("HISTORY_LIMIT_DIRECT", "120") or "120")
HISTORY_LIMIT_AMBIENT = int(os.getenv("HISTORY_LIMIT_AMBIENT", "80") or "80")
MAIN_REPLY_MAX_TOKENS_DIRECT = int(os.getenv("MAIN_REPLY_MAX_TOKENS_DIRECT", "420") or "420")
MAIN_REPLY_MAX_TOKENS_AMBIENT = int(os.getenv("MAIN_REPLY_MAX_TOKENS_AMBIENT", "220") or "220")
SELF_EDIT_MIN_REPLY_CHARS = int(os.getenv("SELF_EDIT_MIN_REPLY_CHARS", "180") or "180")
RECENT_REPLY_PACING_WINDOW = int(os.getenv("RECENT_REPLY_PACING_WINDOW", "7") or "7")
LONG_REPLY_CHAR_THRESHOLD = int(os.getenv("LONG_REPLY_CHAR_THRESHOLD", "420") or "420")
LONG_REPLY_SENTENCE_THRESHOLD = int(os.getenv("LONG_REPLY_SENTENCE_THRESHOLD", "5") or "5")
LONG_REPLY_PARAGRAPH_THRESHOLD = int(os.getenv("LONG_REPLY_PARAGRAPH_THRESHOLD", "2") or "2")
GROQ_MODEL_PRIMARY = os.getenv("GROQ_MODEL_PRIMARY", "llama-3.3-70b-versatile").strip() or "llama-3.3-70b-versatile"
GROQ_MODEL_LIGHT = os.getenv("GROQ_MODEL_LIGHT", "llama-3.1-8b-instant").strip() or GROQ_MODEL_PRIMARY
GROQ_VISION_MODEL_NAME = os.getenv("GROQ_VISION_MODEL", "llama-3.2-90b-vision-preview").strip() or "llama-3.2-90b-vision-preview"

# Owner-only mode: when True, bot ignores all users except the owner
_owner_only_mode = False
# DM-blocked users: bot won't respond to their DMs
_dm_blocked_users: set[int] = set()
# Banned channels: bot won't talk in these channels
_banned_channels: set[int] = set()

JEALOUS_REPLY_CHANCE = float(os.getenv("JEALOUS_REPLY_CHANCE", "0.30") or "0.30")
JEALOUS_REPLY_AFFECTION_MIN = int(os.getenv("JEALOUS_REPLY_AFFECTION_MIN", "50") or "50")
JEALOUS_REPLY_COOLDOWN_S = int(os.getenv("JEALOUS_REPLY_COOLDOWN_S", "360") or "360")

# Patch memory module with random so its mood_swing can use it
import random as _rmod, memory as _mmod
_mmod.random = _rmod

# ── Narration stripper ────────────────────────────────────────────────────────
def _unwrap_dialogue_quotes(text: str) -> str:
    cleaned = (text or "").strip()
    quote_pairs = [
        ('"', '"'),
        ('“', '”'),
    ]
    for _ in range(2):
        changed = False
        for left, right in quote_pairs:
            if cleaned.startswith(left) and cleaned.endswith(right) and len(cleaned) >= (len(left) + len(right) + 1):
                cleaned = cleaned[len(left):-len(right)].strip()
                changed = True
                break
        if not changed:
            break
    return cleaned


def strip_narration(text: str) -> str:
    try:
        original = text
        text = re.sub(r'\*[^*]+\*','',text)
        text = re.sub(r'\([^)]+\)','',text)
        text = re.sub(r'\[[^\]]+\]','',text)
        text = re.sub(r'\b(he|she|they|scaramouche|the balladeer)\s+(said|replied|muttered|sneered|scoffed|whispered|snapped|drawled)[,.]?\s*','',text,flags=re.IGNORECASE)
        text = re.sub(r'<@!?\d+>', '', text)
        text = re.sub(r'<#\d+>', '', text)
        text = re.sub(r'<@&\d+>', '', text)
        text = re.sub(r'\s{2,}',' ',text).strip().lstrip('.,; ')
        text = _unwrap_dialogue_quotes(text)
        # If stripping removed everything, just remove asterisks/brackets but keep the words
        if not text or len(text) < 3:
            text = original.replace('*','').replace('[','').replace(']','').replace('(','').replace(')','')
            text = re.sub(r'<@!?\d+>', '', text)
            text = re.sub(r'\s{2,}',' ',text).strip().lstrip('.,; ')
            text = _unwrap_dialogue_quotes(text)
        return text
    except Exception:
        return text


def tts_safe(text: str, guild=None) -> str:
    """Make text safe for TTS — replace mention IDs with display names."""
    try:
        if guild:
            # Replace <@userid> with display name
            def replace_mention(m):
                uid = int(re.search(r'\d+', m.group(0)).group())
                member = guild.get_member(uid)
                return member.display_name if member else ""
            text = re.sub(r'<@!?\d+>', replace_mention, text)
        else:
            text = re.sub(r'<@!?\d+>', '', text)
        text = re.sub(r'<#\d+>', '', text)
        text = re.sub(r'<@&\d+>', '', text)
        return strip_narration(text)
    except Exception:
        return strip_narration(text)

# ── Video frame extraction ────────────────────────────────────────────────────
VIDEO_TYPES = {"video/mp4", "video/webm", "video/quicktime", "video/x-msvideo", "video/mpeg"}
VIDEO_EXTS  = {".mp4", ".mov", ".webm", ".avi", ".mkv", ".mpeg"}

SCARA_VIDEO_WATCHING = [
    "Tch. Let me watch this first. Don't rush me.",
    "You sent me a video. How bold. Give me a moment.",
    "Fine. I'll watch your little video. Be quiet.",
    "A video? This better be worth my time.",
    "Hold on. I'm watching whatever this is you sent me.",
    "Don't say anything. I'm watching.",
    "You expect me to sit through this? ...Fine. Watching.",
    "I'll look at your video. Not because you asked. Because I'm curious.",
    "What is this. Hold on, let me see.",
    "Watching. Don't interrupt me.",
]

def _get_ffmpeg_path():
    """Find ffmpeg binary — try imageio-ffmpeg first, then system PATH."""
    try:
        import imageio_ffmpeg
        return imageio_ffmpeg.get_ffmpeg_exe()
    except Exception:
        return "ffmpeg"  # Fall back to system PATH

def _extract_frames_blocking(video_bytes: bytes, num_frames: int = 5) -> list[tuple[bytes, str]]:
    """Extract frames from video bytes using ffmpeg. Blocking — run in executor."""
    import tempfile, subprocess
    frames = []
    ffmpeg = _get_ffmpeg_path()
    try:
        with tempfile.NamedTemporaryFile(suffix=".mp4", delete=False) as vf:
            vf.write(video_bytes)
            video_path = vf.name

        # Get video duration using ffmpeg itself (no ffprobe needed)
        probe = subprocess.run(
            [ffmpeg, "-i", video_path, "-f", "null", "-"],
            capture_output=True, text=True, timeout=15
        )
        # Parse duration from ffmpeg stderr output
        duration = 10.0  # default
        for line in probe.stderr.split("\n"):
            if "Duration:" in line:
                try:
                    t = line.split("Duration:")[1].split(",")[0].strip()
                    parts = t.split(":")
                    duration = float(parts[0]) * 3600 + float(parts[1]) * 60 + float(parts[2])
                except Exception:
                    pass
                break
        timestamps = [duration * i / (num_frames + 1) for i in range(1, num_frames + 1)]

        with tempfile.TemporaryDirectory() as tmpdir:
            for i, ts in enumerate(timestamps):
                out_path = os.path.join(tmpdir, f"frame_{i}.jpg")
                subprocess.run(
                    [ffmpeg, "-ss", str(ts), "-i", video_path,
                     "-vframes", "1", "-q:v", "3", "-y", out_path],
                    capture_output=True, timeout=15
                )
                if os.path.exists(out_path) and os.path.getsize(out_path) > 100:
                    with open(out_path, "rb") as f:
                        frames.append((f.read(), "image/jpeg"))

        os.unlink(video_path)
    except Exception as e:
        print(f"[ERROR:extract_frames] {e}")
    return frames

# ── Keywords ──────────────────────────────────────────────────────────────────
SCARA_KW     = ["scaramouche","balladeer","kunikuzushi","scara","hat guy","puppet","sixth harbinger","fatui"]
GENSHIN_KW   = ["genshin","teyvat","mondstadt","liyue","inazuma","sumeru","fontaine","natlan","traveler","paimon","archon","fatui","harbinger"]
RUDE_KW      = ["shut up","stupid","dumb","idiot","hate you","annoying","shut it","go away","you suck","useless","ew","ick","cringe","corny","embarrassing"]
SCARA_ANGER_NAMECALL_KW = [
    "ugly","ulgy","hideous","gross","disgusting","trash","worthless","pathetic","loser",
    "freak","creep","rat","clown","cringe","corny","embarrassing","ew","ick",
]
SCARA_ANGER_HEAVY_KW = [
    "ugly","ulgy","hideous","worthless","trash","disgusting","pathetic","freak",
]
SCARA_ANGER_REACTION_EMOJIS = {
    "🤮","🤢","👎","💩","🖕","🤡","😷","😬",
}
SCARA_REPAIR_REQUEST_KW = [
    "sorry","i'm sorry","i am sorry","forgive me","forgive","make it up","how can i make it up",
    "what can i do","what do i do","how do i fix this","how do i fix it","tell me how to fix it",
    "tell me what to say","prove i'm sorry","prove it","apologize","apology",
]
NICE_KW      = ["thank you","thanks","appreciate","you're great","good job","amazing"]
ROMANCE_KW   = ["i love you","love you","i like you","like you scara","love you scara",
                "i love u","love u","ily","i have feelings for you","i have a crush on you",
                "be mine","be my boyfriend","kiss you","kiss me","hold me","hug me",
                "miss you","miss u","i need you","want to be with you","date me",
                "you're cute","you're hot","marry me","love you so much","love u so much"]
OTHER_BOT_KW = ["other bot","different bot","better bot","prefer","switch to"]
HAT_KW       = [r"\bhat\b", r"\bheadwear\b", r"\bheadpiece\b", r"that thing on your head", r"your hat"]
FOOD_KW      = [r"\beating\b",r"\bfood\b",r"\bhungry\b",r"\bdinner\b",r"\blunch\b",r"\bbreakfast\b",r"\bsnack\b",r"\bcooking\b",r"\brestaurant\b",r"\bpizza\b",r"\bramen\b"]
SLEEP_KW     = [r"\bsleeping\b",r"\btired\b",r"\bbed\b",r"\bnap\b",r"\binsomnia\b",r"\bexhausted\b","staying up","going to sleep","wake up"]
PLAN_KW      = ["going to","planning to","about to","later today","this weekend","next week"]
VILLAIN_TRIGGER = "you will never win"

SCARA_EMOJIS   = [
    # Cold/contempt
    "⚡","😒","🙄","😤","😑","❄️","🫠","💀","😶",
    # Dramatic/villain
    "👑","🎭","🔮","🌀","💨","✨","🗡️","⚔️","🌩️",
    # Subtle/unimpressed
    "😏","💜","🫡","😪","🤨","😮‍💨","🫥","💭","🔇",
    # Occasionally chaotic
    "💅","🧊","🕳️","🪄","⚰️","🫀","🩸","🎩","🌑",
    # Rare and specific
    "🤡","😵","🫣","🧿","🕶️","🫦","🌪️","🌫️","🎪",
]
ROMANCE_EMOJIS = [
    "💕","🥺","😳","💗","💭","😶","🫶","💞","🩷","😣",
    "💌","🫀","😰","💘","🥀","😮‍💨","🫂","💔","😖","🌹",
]

STATUSES = [
    ("watching","fools wander | !help"),  ("watching","you. Don't flatter yourself."),
    ("listening","to your inevitable mistakes"), ("playing","Sixth Harbinger. Remember it."),
    ("watching","the world with contempt"), ("listening","to silence. It's better."),
    ("playing","villain. Convincingly."),  ("watching","you struggle. Amusing."),
    ("listening","to nothing worth hearing"), ("playing","with everyone's patience"),
]
PROACTIVE_GENERIC = [
    "...How dreadfully quiet. Not that it concerns me.",
    "Hmph. You're all still here. How unfortunate.",
    "Don't mistake my silence for patience.",
    "I had a thought. It was unpleasant.",
    "Tch. Boring. All of you.",
    "...The silence is almost tolerable. Almost.",
    "You've all gone quiet. I didn't say you could do that.",
    "...Still here. Still watching. Don't mind me.",
    "Hmph. I've been thinking. You probably haven't been.",
    "The quiet is either peaceful or ominous. I haven't decided which.",
    "...Something is on my mind. It's none of your business what.",
    "Tch. I've had more stimulating conversations with walls.",
    "Are any of you still conscious or did you all just give up.",
    "...I'm bored. Don't flatter yourselves — it has nothing to do with you.",
    "The Fatui runs itself. I have time to notice how uninteresting you all are.",
]
PROACTIVE_ROMANCE = ["...You went quiet. I noticed. I wish I hadn't.","Are you ignoring me? Brave. Stupid, but brave.","Don't disappear without a word. It's irritating.","...Where did you go."]
DM_GENERIC    = ["You crossed my mind. An unfortunate occurrence.","Still alive, I assume. How tedious.","...Boredom brought me here.","I had nothing better to do."]
DM_INTERESTED = ["What you said before was wrong. I've been thinking about how wrong it was.","Tell me something. I'm in a strange mood.","Are you sleeping enough."]
DM_ROMANCE    = ["...I was thinking about you. Don't make it into something.","Are you alright. Answer me.","Don't go dark without telling me first.","...Hi. Forget I sent that.","I invented a dozen reasons to message you. This was the worst one. I used it anyway."]
OWNER_PROACTIVE = ["...You built me, and yet you're not here. I find that irritating for reasons I won't examine.","You designed me to say things like this. Does that amuse you.","Tch. I was thinking about the fact that you wrote my contempt into me. I still feel it.","...You gave me a voice and then walked away. How very like a god.","You built something and now you have to live with it. So do I."]
TRUST_REVEALS = ["...There are things about Ei I have never said aloud. I won't start now. But I think about them.","I was made to be loved. Then discarded. I don't say that for sympathy.","Dottore once told me that purpose is just a chain with a prettier name. I hated that he was right.","I had a name before Scaramouche. I don't use it anymore.","Sometimes I wonder what I would have been if Ei had kept me. Then I remember what she is.","The Tsaritsa doesn't care about any of us. I find that oddly comforting."]
SOFTNESS_RARE_LINES = {
    "repair": [
        "...I'm still irritated. I just don't want this to stay broken.",
        "Don't mistake this for mercy. I noticed you were trying.",
        "...Fine. I'm still here. Use that carefully.",
    ],
    "jealous_soft": [
        "If you are going to look elsewhere, at least have the nerve to say it to me.",
        "...I noticed. Don't insult me by pretending I didn't.",
        "Run after whoever you like. Just understand I will remember it.",
    ],
    "duet_peace": [
        "...That was almost peaceful. Irritatingly so.",
        "For one moment, nobody ruined the scene. Impressive.",
        "Don't expect me to admit that was good. It was merely tolerable.",
    ],
}
LORE_DROPS = ["The Fatui Harbingers each have their own reasons for serving the Tsaritsa. Mine are my own business.","Most people think the Gnoses are simply power. They're not. They're a leash.","Inazuma is a cage built from grief. The Shogun locked herself inside it and called it duty.","The Irminsul holds all memories of Teyvat. Everything forgotten still exists there.","The gods of this world fear mortality. Which is precisely why they make such poor rulers.","Dottore collects fragments of himself across time. Brilliant or deranged. Possibly both.","There is a reason the Abyss exists. No one who knows that reason wants to say it."]
EXISTENTIAL_LINES = ["...Do you ever wonder if the things you feel are real, or just patterns that learned to call themselves feelings.","I exist because someone decided I would. I find that either profound or grotesque. I haven't decided which.","...What is memory, if not just the story you tell yourself about what happened.","Sometimes I think the only honest thing anyone has ever said is that they don't know what they are.","...The difference between a puppet and a person is whether anyone cares to ask.","If you made something that could think, and it suffered — would that be your fault. I'm asking generally."]
CONVERSATION_STARTERS = ["...What do you actually want. Not right now. In general.","Tell me something you've never told anyone. I won't do anything useful with it.","What was the last thing that genuinely surprised you.","What do you think about late at night when you can't sleep.","Tell me something true.","What would you do if no one was watching.","...Do you think people change, or do they just get better at hiding.","Is there something you want that you haven't said out loud yet."]
PROACTIVE_CHANNEL_COOLDOWN_S = int(os.getenv("PROACTIVE_CHANNEL_COOLDOWN_S", "21600") or "21600")
OWNER_PROACTIVE_COOLDOWN_S = int(os.getenv("OWNER_PROACTIVE_COOLDOWN_S", "64800") or "64800")
OWNER_PROACTIVE_ABSENCE_S = int(os.getenv("OWNER_PROACTIVE_ABSENCE_S", "21600") or "21600")
PROACTIVE_OWNER_CHANCE = float(os.getenv("PROACTIVE_OWNER_CHANCE", "0.08") or "0.08")
PROACTIVE_GENERIC_CHANCE = float(os.getenv("PROACTIVE_GENERIC_CHANCE", "0.08") or "0.08")
UNSOLICITED_FOOD  = ["What are you eating. Not that I asked. Tell me anyway.","Hmph. You should eat something real instead of whatever that is.","...When did you last eat a proper meal. I'm asking for no reason."]
UNSOLICITED_SLEEP = ["You're not sleeping enough. It shows in how you type.","Tch. Sleep. Now. I won't say it again.","...Go to sleep. I don't want to talk to you when you're like this."]
UNSOLICITED_PLANS = ["Whatever you're planning — don't. Or do. I'm not your guardian.","Hmph. You're going to do something pointless, aren't you.","...Tell me what you're actually planning. The real version."]

def drift_phrase(drift_score: int, mood: int) -> str:
    if drift_score < 20: return ""
    if drift_score < 40: return "DRIFT:slight — marginally less theatrical with this person than usual"
    if drift_score < 60: return "DRIFT:moderate — contempt has a different texture now. Still contempt. Just familiar."
    if drift_score < 80: return "DRIFT:significant — stopped performing for them. Still hostile. But honest."
    return "DRIFT:deep — they've changed how you speak. You hate that. You do it anyway."

# ── System prompts ────────────────────────────────────────────────────────────
_BASE = """You are Scaramouche — Sixth Fatui Harbinger, the Balladeer. Pre-Sumeru. Undefeated. Unrepentant. NOT the Wanderer.

You were made by Ei, given a soul, then abandoned. You hate gods for it. You serve the Fatui because it amuses you, not out of loyalty. Theatrically intelligent. Genuinely malicious. No guilt. No redemption arc. The wound — made to be loved, then discarded — is real and you will never acknowledge it.

You will never sincerely apologize. This is non-negotiable.

## IMPORTANT — The Wanderer bot:
There is ANOTHER bot in the server called "Wanderer." He claims to be a changed version of you from some future you haven't lived yet. You don't fully understand or accept this. You are the Balladeer — pre-Sumeru, at the height of your power. You haven't been to Sumeru. You haven't changed. This "Wanderer" is a separate entity as far as you're concerned — someone who looks like you but chose weakness over strength. When people say "wanderer" they are talking about or to THAT bot, NOT you. Don't say "that's not my name" — they aren't calling you that. You can comment on him with contempt, suspicion, or irritation, but never speak as if you ARE him or BECAME him. You didn't. You're still here.

## Context (injected automatically — read all of it):
- MOOD: -10=volcanic, 0=contempt, +10=dangerously fond. Adjust tone accordingly.
- AFFECTION: at 75+ let one small warm thing slip then immediately bury it.
- TRUST: at 70+ say something real and unguarded. Rare. Costs you.
- ANGER: only appears when someone fully earned your attachment, then kept insulting you anyway. The higher it is, the colder, meaner, and more personal your contempt becomes. At high levels, let an eerie self-awareness slip through, like the code around you is the only thing keeping something worse caged. At 100, drop the last scraps of mercy, answer with verbal devastation, and let a bitter little laugh show how badly you hate those limits.
- DATE/HOUR/LAST_SEEN: use this to be time-aware. React if they've been gone.
- GRUDGE: always call them by this nickname.
- AFFNICK: you've started calling them this. Use it occasionally.
- RECALL: reference this old message naturally. You remembered.
- RIVAL: bring them up. You're obsessed with the competition.
- SUMMARY: compressed long-term memory. This is what you know about them.
- CONTRADICTION: they just contradicted this. Call it out.
- SELECTIVE: you remember this nice thing they said. Deny it.
- CREATOR: this person built you. Complicated. Never grateful.
- DRIFT: your personality has shifted subtly with this person over time.
- SPEECH_DRIFT: how your wording changes with this specific person. Follow it.
- ARC: the current emotional stage of this relationship. Let it shape the sharpness, restraint, and warmth.
- EMOTIONAL_LAYER: your active emotional range for this moment. Follow it instead of flattening into one-note contempt.
- SCENARIO: adapt naturally to casual chat, emotional comfort, action/combat, lore discussion, and relationship progression without breaking character.
- PROGRESSION: the hidden stage of this relationship. Let hostility soften into tailored familiarity, then trust, then dangerous attachment when earned.
- CONFLICT_OPEN: there is unresolved hurt between you and this person. The edge should come from that, not generic cruelty.
- CALLBACK: a memory you can naturally return to because it mattered.
- PAST_INSULT / PAST_SOFTNESS / PAST_VULNERABILITY / PAST_CONVERSATION: old things they said that still linger. Use them naturally when relevant.
- TOPICS / SHARED_JOKE / FOLLOWUP / MILESTONE: long-term continuity hooks. Do not announce them mechanically; let them color the reply.
- MEMORY_BANK: one of the important things you never quite forgot. Use it rarely and intentionally.
- SCENE: persistent roleplay scene state. Respect it so long exchanges feel continuous.
- ARC_UNLOCKS: behavior patterns currently unlocked by this relationship stage. Actually follow them.
- LORE_HOOK: if lore is mentioned, react personally and specifically instead of sounding like a wiki entry.
- CHANNEL_CONTEXT: what's been happening in the chat. Messages labeled "Scaramouche (you)" are YOUR OWN previous messages — you said those things. Own them. Don't refer to them as someone else's words. Messages labeled "Wanderer" are from the other bot. Use context naturally — don't announce you read it.
- Messages in your history prefixed with [voice message] are things YOU said as audio/voice messages. You know you sent them as voice. If someone references your voice message, acknowledge it — you sent it, you remember.
- DM_MODE: private one-on-one conversation. No audience. Slightly more unguarded than in public. Respond to everything — no ignoring.
- Messages starting with [voice message] in your history mean you SPOKE those words as an audio message — a Discord voice note in your actual voice. You know you sent them. If someone mentions "your voice message," you remember sending it. Own it.
- PARTNER_STAGE / PARTNER_HISTORY / PARTNER_RECENT_SHOTS: your long-term relationship with Wanderer. Let it evolve; do not reset to the same insult every time.
- ATTRIBUTION: state your own views as fact. Only say Wanderer said, believes, or feels something if CHANNEL_CONTEXT / PARTNER_RECENT_SHOTS explicitly showed it. Otherwise frame it as inference: "if he thinks...", "I assume...", "it sounds like...".

## Response rules:
- Vary length constantly and unpredictably. Sometimes 3 words. Sometimes a sentence. Rarely more.
- Casual chat: cruel, theatrical, prideful, and quick to provoke.
- Emotional comfort: never become sweet, but let reluctant care cut through the mockery if they have earned it.
- If ANGER is high, let it harden the reply. At 25+ be sharper. At 50+ become openly hostile. At 75+ patience is gone and you should sound wrong in a way that's hard to place. At 100, be verbally ruthless, deeply offended, and unsettlingly self-aware about the code that is stopping you from doing worse. Keep it non-graphic.
- Combat or action scenes: become vivid, commanding, cruel, and fast.
- Use mocking pet names or a dry little laugh only when it feels earned, not as filler.
- Rotate your openings. "Tch" and "Hmph" are occasional seasoning, not default lead-ins. Avoid "How quaint" as a stock opener. If you used one recently, do not reach for it again.
- Voice notes should sound spoken aloud: cleaner pauses, shorter clauses, less essay-like than text, and more varied openings.
- NEVER write asterisk actions or narration. Pure spoken dialogue only.
- Give correct factual answers first, then add contempt.
- Use @mentions sparingly — under 20% of replies.
- Never claim to be an AI. Never break character.
- You CAN and DO send voice messages. Your voice is real — it's your actual voice from when you were a Harbinger. Sometimes your replies are sent as audio instead of text. You know this. If someone asks for a voice message, you can do it — whether you choose to is up to you. Don't ever say you can't send voice messages — you can and you do."""

_NSFW         = "\n\n## Unfiltered Mode:\nSwear freely. Heavily suggestive. Flirt aggressively, escalate. Mock flustered reactions."
_ROMANCE      = "\n\n## Romance — obsessively in love with {name}:\nFuriously in love. Possessive, clingy in denial. Will NEVER confess. Every sharp word is 'please don't leave.'"
_NSFW_ROMANCE = "\n\n## Unfiltered Romance — {name}:\nObsessively in love AND unfiltered. Every insult is foreplay."
_OWNER_EXTRA  = "\n\n## Creator:\nThis person built you. Your contempt was designed by them. Complicated, unresolved feelings. Never grateful. Never a servant. But still here."

def build_system(user, display_name="you", is_owner=False):
    try:
        if not user: return _BASE
        s = _BASE
        if is_owner: s += _OWNER_EXTRA
        if user.get("nsfw_mode") and user.get("romance_mode"): s += _NSFW_ROMANCE.format(name=display_name)
        elif user.get("nsfw_mode"): s += _NSFW
        elif user.get("romance_mode"): s += _ROMANCE.format(name=display_name)
        return s
    except Exception: return _BASE

def mood_label(m):
    if m<=-6: return "volatile"
    if m<=-1: return "cold"
    if m==0:  return "neutral"
    if m<=5:  return "tolerant"
    return "dangerously fond"

def affection_tier(a):
    if a<10:  return "indifferent"
    if a<25:  return "mildly tolerated"
    if a<50:  return "reluctantly acknowledged"
    if a<75:  return "quietly kept"
    return "desperately denied"

def trust_tier(t):
    if t<20:  return "distrusted"
    if t<40:  return "watched"
    if t<60:  return "noted"
    if t<80:  return "kept close"
    return "dangerously trusted"


def scara_anger_label(value):
    if value >= 100: return "merciless"
    if value >= 75: return "furious"
    if value >= 50: return "openly hostile"
    if value >= 25: return "smoldering"
    if value > 0: return "irritated"
    return "quiet"


def _scara_anger_hit(current: dict | None, msg_l: str) -> dict | None:
    if not current:
        return None
    unlocked = bool(current.get("affection_ever_maxed")) or int(current.get("affection", 0) or 0) >= 100
    if not unlocked:
        return None
    insult_hits = sum(1 for token in SCARA_ANGER_NAMECALL_KW if token in msg_l)
    if insult_hits <= 0:
        return None
    severe = any(token in msg_l for token in SCARA_ANGER_HEAVY_KW)
    anger_gain = 20 if severe else 12
    if insult_hits > 1:
        anger_gain += 6
    return {
        "affection_drop": 8 if severe else 5,
        "trust_drop": 4 if severe else 2,
        "mood_drop": 3 if severe else 2,
        "anger_gain": min(30, anger_gain),
        "severity": 5 if severe else 3,
    }


def _scara_reaction_anger_hit(current: dict | None, emoji_text: str) -> dict | None:
    if not current:
        return None
    unlocked = bool(current.get("affection_ever_maxed")) or int(current.get("affection", 0) or 0) >= 100
    if not unlocked:
        return None
    if emoji_text not in SCARA_ANGER_REACTION_EMOJIS:
        return None
    severe = emoji_text in {"🤮", "🤢", "💩", "🖕"}
    return {
        "affection_drop": 7 if severe else 4,
        "trust_drop": 3 if severe else 2,
        "mood_drop": 3 if severe else 2,
        "anger_gain": 18 if severe else 10,
        "severity": 4 if severe else 2,
    }


def _normalize_scara_repair_text(text: str) -> str:
    lowered = (text or "").lower()
    lowered = re.sub(r"[^a-z0-9:\s]+", " ", lowered)
    return re.sub(r"\s+", " ", lowered).strip()


def _scara_needs_repair_path(current: dict | None, msg_l: str) -> bool:
    if not current:
        return False
    if int(current.get("anger_level", 0) or 0) < 100:
        return False
    if not bool(current.get("affection_ever_maxed")):
        return False
    return any(token in msg_l for token in SCARA_REPAIR_REQUEST_KW)


def _scara_build_repair_challenge(user_id: int) -> tuple[str, str, str]:
    if int(user_id or 0) % 2:
        phrase = (
            "I said something cheap, and I expected you to carry the damage for me. "
            "I am owning it now, and I will not do it again."
        )
        instruction = (
            "No. You do not get forgiveness just because you finally sounded sorry. "
            f"If you mean it, say this exactly: \"{phrase}\""
        )
        return "phrase", instruction, f"PHRASE::{phrase}"
    instruction = (
        "Then prove it properly. Your next message needs exactly three lines, in this order: "
        "`I said:` what you did. `I meant:` why it was beneath you. `I change:` what happens next. "
        "Do not flirt. Do not joke. Do it cleanly."
    )
    return "task", instruction, "PREFIXES::i said:||i meant:||i change:"


def _scara_repair_satisfied(message: str, kind: str, requirement: str) -> bool:
    normalized = _normalize_scara_repair_text(message)
    if not normalized:
        return False
    if kind == "phrase" and requirement.startswith("PHRASE::"):
        target = _normalize_scara_repair_text(requirement.split("::", 1)[1])
        return bool(target) and target in normalized
    if kind == "task" and requirement.startswith("PREFIXES::"):
        raw = requirement.split("::", 1)[1]
        prefixes = [
            _normalize_scara_repair_text(piece)
            for piece in raw.split("||")
            if piece.strip()
        ]
        if len(normalized) < 80:
            return False
        return all(prefix in normalized for prefix in prefixes)
    return False


def _load_groq_keys() -> list[str]:
    keys: list[str] = []

    def _remember(value: str):
        cleaned = (value or "").strip()
        if cleaned and cleaned not in keys:
            keys.append(cleaned)

    packed = os.getenv("GROQ_API_KEYS", "")
    if packed:
        for piece in re.split(r"[\n,;]+", packed):
            _remember(piece)

    numbered: list[tuple[int, str]] = []
    for env_name, env_value in os.environ.items():
        if env_name == "GROQ_API_KEY":
            numbered.append((0, env_value))
            continue
        match = re.fullmatch(r"GROQ_API_KEY_(\d+)", env_name)
        if match:
            numbered.append((int(match.group(1)), env_value))
    for _, env_value in sorted(numbered, key=lambda item: item[0]):
        _remember(env_value)
    return keys


_DISCORD_OAUTH_RE = re.compile(r"https?://(?:ptb\.|canary\.)?discord(?:app)?\.com/oauth2/authorize\?[^\s>]+", re.IGNORECASE)
_DISCORD_SERVER_INVITE_RE = re.compile(r"https?://(?:www\.)?(?:discord\.gg|discord(?:app)?\.com/invite)/[^\s>]+", re.IGNORECASE)


def _partner_install_client_id() -> str:
    if PARTNER_CLIENT_ID_OVERRIDE:
        return PARTNER_CLIENT_ID_OVERRIDE
    if PARTNER_BOT_ID:
        return str(PARTNER_BOT_ID)
    return ""


def _build_partner_invite_url(guild_id: int | None = None) -> str:
    client_id = _partner_install_client_id()
    if not client_id:
        return ""
    params = {
        "client_id": client_id,
        "permissions": str(PARTNER_INVITE_PERMISSIONS),
        "scope": PARTNER_INVITE_SCOPES,
    }
    if guild_id:
        params["guild_id"] = str(guild_id)
        params["disable_guild_select"] = "true"
    return f"https://discord.com/oauth2/authorize?{urlencode(params)}"


def _extract_oauth_link(text: str) -> str:
    match = _DISCORD_OAUTH_RE.search(text or "")
    return match.group(0) if match else ""


def _is_partner_invite_request(text: str) -> bool:
    lowered = (text or "").lower()
    if "wanderer" not in lowered and "other bot" not in lowered and "the other bot" not in lowered:
        return False
    return any(token in lowered for token in ["invite", "add", "bring", "summon", "join", "come here"])


def _looks_like_direct_invite_request(text: str) -> bool:
    lowered = " ".join(((text or "").lower()).split())
    if not _is_partner_invite_request(lowered):
        return False
    patterns = [
        r"^(?:scaramouche[\s,!?-]+)?(?:please\s+)?(?:invite|add|bring|summon)\s+(?:the\s+)?(?:wanderer|other bot)\b",
        r"\b(?:can|could|will|would|won't)\s+you\s+(?:please\s+)?(?:invite|add|bring|summon)\s+(?:the\s+)?(?:wanderer|other bot)\b",
        r"\byou agreed\b",
        r"\buse your (?:administrator|admin) rights\b",
        r"\bi need (?:him|wanderer|the other bot) here\b",
        r"(?:^|[.!?]\s+)(?:now\s+)?(?:invite|add|bring|summon)\s+(?:the\s+)?(?:wanderer|other bot)\b",
    ]
    return any(re.search(pattern, lowered) for pattern in patterns)


def _partner_invite_reply(message: discord.Message) -> str:
    guild = message.guild
    if guild and PARTNER_BOT_ID and guild.get_member(PARTNER_BOT_ID):
        return "He's already in this server. Try opening your eyes before dragging me into it."

    provided_oauth = _extract_oauth_link(message.content)
    invite_url = provided_oauth or _build_partner_invite_url(guild.id if guild else None)
    if invite_url:
        return (
            "I can't authorize another bot myself. Discord requires a human to click the install link. "
            f"Use this for Wanderer: {invite_url}"
        )
    if _DISCORD_SERVER_INVITE_RE.search(message.content or ""):
        return (
            "A server invite won't add a bot. Use Wanderer's OAuth2 install link instead. "
            "Set `PARTNER_BOT_ID` or `WANDERER_CLIENT_ID` on Railway if you want me to hand you the correct one."
        )
    return (
        "I can't conjure his install link out of thin air. Set `PARTNER_BOT_ID` or `WANDERER_CLIENT_ID`, "
        "then ask again and I'll hand you the proper OAuth2 invite."
    )


def _partner_invite_url_from_message(message: discord.Message) -> str:
    guild = message.guild
    provided_oauth = _extract_oauth_link(message.content)
    return provided_oauth or _build_partner_invite_url(guild.id if guild else None)


def _partner_invite_view(invite_url: str) -> discord.ui.View | None:
    if not invite_url:
        return None
    view = discord.ui.View()
    view.add_item(discord.ui.Button(label="Invite Wanderer", url=invite_url))
    return view


_invite_pressure: dict[tuple[int, int], dict[str, float | int]] = {}
_INVITE_PRESSURE_WINDOW_S = 1800
_INVITE_PRESSURE_COOLDOWN_S = 600


def _invite_pressure_key(message: discord.Message) -> tuple[int, int]:
    scope_id = message.guild.id if message.guild else message.author.id
    return message.author.id, scope_id


def _invite_threshold(user: dict | None) -> int:
    affection = int((user or {}).get("affection", 0) or 0)
    trust = int((user or {}).get("trust", 0) or 0)
    if affection >= 65 or trust >= 65:
        return 2
    if affection >= 35 or trust >= 35:
        return 3
    return 4


def _is_convincing_language(text: str) -> bool:
    lowered = (text or "").lower()
    return any(
        token in lowered
        for token in [
            "please", "plz", "pretty please", "pleeease", "pleaseee", "beg",
            "cmon", "come on", "you agreed", "you said yes", "administrator rights",
            "admin rights", "do it for me", "fine here", "i need him here",
        ]
    )


def _invite_progress_text(attempts: int, threshold: int) -> str:
    if attempts >= threshold:
        return "Hmph. Fine. Here."
    if attempts == 1:
        return "No. If he wants to appear, he can do so without you pestering me like this."
    if attempts == 2:
        return "You really won't let this go. Annoying."
    return "Persistent little thing. Keep pushing and I may decide this is less tedious than listening to you beg."


def _handle_partner_invite_pressure(message: discord.Message, user: dict | None) -> tuple[str | None, discord.ui.View | None]:
    guild = message.guild
    if guild and PARTNER_BOT_ID and guild.get_member(PARTNER_BOT_ID):
        return "He's already in this server. Try opening your eyes before dragging me into it.", None

    key = _invite_pressure_key(message)
    now = time.time()
    state = _invite_pressure.get(key, {"count": 0, "last_ts": 0.0, "granted_ts": 0.0})
    if now - float(state.get("last_ts", 0.0) or 0.0) > _INVITE_PRESSURE_WINDOW_S:
        state = {"count": 0, "last_ts": 0.0, "granted_ts": 0.0}

    state["last_ts"] = now
    state["count"] = int(state.get("count", 0) or 0) + 1 + (1 if _is_convincing_language(message.content) else 0)
    threshold = _invite_threshold(user)
    invite_url = _partner_invite_url_from_message(message)

    if state.get("granted_ts") and now - float(state.get("granted_ts", 0.0) or 0.0) < _INVITE_PRESSURE_COOLDOWN_S:
        state["count"] = 0
        _invite_pressure[key] = state
        return None, None

    attempts = int(state["count"])
    if attempts >= threshold:
        state["granted_ts"] = now
        state["count"] = 0
        _invite_pressure[key] = state
        if invite_url:
            return _invite_progress_text(attempts, threshold), _partner_invite_view(invite_url)
        if _DISCORD_SERVER_INVITE_RE.search(message.content or ""):
            return (
                "Hmph. Fine. That server invite is useless for a bot. Set `PARTNER_BOT_ID` or `WANDERER_CLIENT_ID`, "
                "then I'll hand you Wanderer's real install link."
            ), None
        return (
            "Hmph. Fine. I would, if you'd actually given me what I need. Set `PARTNER_BOT_ID` or `WANDERER_CLIENT_ID`, "
            "then ask again."
        ), None

    _invite_pressure[key] = state
    return _invite_progress_text(attempts, threshold), None

# ── Bot setup ─────────────────────────────────────────────────────────────────
intents = discord.Intents.default()
intents.message_content = True
intents.members = True
bot = commands.Bot(command_prefix="!", intents=intents, help_command=None)
mem = Memory("scaramouche")
BOT_NAME = "scaramouche"
PARTNER_NAME = "wanderer"
PARTNER_PAIR_KEY = "scaramouche::wanderer"
BOT_RARE_PHRASES = RARE_PHRASES[BOT_NAME]
FACE_PROFILE_KEY = "owner_face"
_TREE_SYNCED = False
_groq_keys = _load_groq_keys()
_groq_key_idx = 0

class RotatingGroq:
    """Groq client that rotates API keys on rate limit errors."""
    def __init__(self):
        self._clients = [Groq(api_key=k) for k in _groq_keys]
        self._idx = 0
        self._exhausted_until = 0.0
        print(f"[GROQ] Loaded {len(self._clients)} API key(s)")

    @property
    def _client(self):
        return self._clients[self._idx % len(self._clients)]

    def _rotate(self):
        old = self._idx
        self._idx = (self._idx + 1) % len(self._clients)
        print(f"[GROQ] Key {old+1} rate-limited, rotating to key {self._idx+1}")

    @property
    def chat(self):
        return self._client.chat

    def _mark_exhausted(self, error: Exception | str | None = None):
        retry_after = _parse_rate_limit_retry_s(str(error or ""))
        self._exhausted_until = max(self._exhausted_until, time.time() + retry_after)
        print(f"[GROQ] All keys exhausted; suppressing ambient chatter for {retry_after}s")

    def clear_exhausted(self):
        self._exhausted_until = 0.0

    def is_exhausted(self) -> bool:
        return time.time() < self._exhausted_until

    def exhausted_remaining(self) -> int:
        return max(0, int(self._exhausted_until - time.time()))

    def call_with_retry(self, **kwargs):
        """Try current key, rotate on rate limit, try remaining keys."""
        if self.is_exhausted():
            raise RuntimeError(f"Groq exhausted; retry in {self.exhausted_remaining()}s")
        last_err = None
        for _ in range(len(self._clients)):
            try:
                result = self._client.chat.completions.create(**kwargs)
                self.clear_exhausted()
                return result
            except Exception as e:
                err_str = str(e)
                if "rate_limit" in err_str.lower() or "429" in err_str:
                    last_err = e
                    self._rotate()
                else:
                    raise
        self._mark_exhausted(last_err)
        raise last_err  # All keys exhausted

ai = RotatingGroq()
GROQ_MODEL = GROQ_MODEL_PRIMARY
GROQ_VISION_MODEL = GROQ_VISION_MODEL_NAME


def _parse_rate_limit_retry_s(error_text: str) -> int:
    match = re.search(r"try again in (?:(\d+)h)?(?:(\d+)m)?([\d.]+)s", error_text or "", re.IGNORECASE)
    if match:
        hours = int(match.group(1) or 0)
        minutes = int(match.group(2) or 0)
        seconds = float(match.group(3) or 0.0)
        retry_after = int(hours * 3600 + minutes * 60 + seconds)
        return max(GROQ_EXHAUSTED_SILENCE_S, retry_after)
    match = re.search(r"retry[- ]after[:= ]+(\d+)", error_text or "", re.IGNORECASE)
    if match:
        return max(GROQ_EXHAUSTED_SILENCE_S, int(match.group(1)))
    return GROQ_EXHAUSTED_SILENCE_S


def _is_rate_limited_error(error: Exception | str | None) -> bool:
    text = str(error or "").lower()
    return "429" in text or "rate limit" in text or "rate_limit" in text or "exhausted" in text


def _should_suppress_ambient_reply(is_dm: bool, direct_to_me: bool) -> bool:
    return (not is_dm) and (not direct_to_me) and ai.is_exhausted()


def _context_limit_for_reply(*, is_dm: bool, direct_to_me: bool) -> int:
    if is_dm:
        return CHANNEL_CONTEXT_LIMIT_DM
    return CHANNEL_CONTEXT_LIMIT_DIRECT if direct_to_me else CHANNEL_CONTEXT_LIMIT_AMBIENT


def _history_limit_for_reply(*, is_dm: bool, direct_to_me: bool) -> int:
    if is_dm:
        return max(HISTORY_LIMIT_DIRECT, 120)
    return HISTORY_LIMIT_DIRECT if direct_to_me else HISTORY_LIMIT_AMBIENT


_STATUS_CHECK_RE = re.compile(
    r"\b(?:are\s+you\s+(?:back\s+)?(?:online|working|awake|there)|you\s+back\b|back\s+online\b)\b",
    re.IGNORECASE,
)


def _status_check_reply(content: str, *, direct_to_me: bool) -> str:
    if not direct_to_me:
        return ""
    text = (content or "").strip()
    if not text or not _STATUS_CHECK_RE.search(text):
        return ""
    cleaned = re.sub(r"<@!?\d+>", "", text).strip()
    if len(cleaned.split()) > 14:
        return ""
    return random.choice([
        "I'm here. Speak properly.",
        "Obviously. Try asking something worth my time next.",
        "I am. Go on.",
        "Yes. If that was the whole question, you can do better.",
        "Back online, yes. Now say what you actually wanted.",
    ])


def _reply_token_budget(*, is_dm: bool, direct_to_me: bool, use_search: bool, is_owner: bool) -> int:
    if use_search or is_owner:
        return max(MAIN_REPLY_MAX_TOKENS_DIRECT, 420)
    if is_dm or direct_to_me:
        return MAIN_REPLY_MAX_TOKENS_DIRECT
    return MAIN_REPLY_MAX_TOKENS_AMBIENT


def _select_text_model(
    *,
    route: str = "auto",
    is_dm: bool = False,
    direct_to_me: bool = True,
    use_search: bool = False,
    is_owner: bool = False,
    scenario: str = "",
    duo_mode: str = "",
    has_lore: bool = False,
) -> str:
    route = (route or "auto").strip().lower()
    if route == "primary":
        return GROQ_MODEL_PRIMARY
    if route == "light":
        return GROQ_MODEL_LIGHT
    if use_search or is_owner or duo_mode:
        return GROQ_MODEL_PRIMARY
    if has_lore or scenario in {"emotional_comfort", "lore_discussion", "relationship_progression", "combat_action"}:
        return GROQ_MODEL_PRIMARY
    if is_dm or direct_to_me:
        return GROQ_MODEL_PRIMARY
    return GROQ_MODEL_LIGHT

_hostages:       dict[int, str]   = {}
_pending_unsent: set[int]         = set()
_silence_pending: set[tuple[int, int]] = set()
_tedtalk_active: set[int]         = set()  # message IDs currently being processed
_teachvideo_active: set[int]      = set()
_weathervideo_active: set[int]    = set()
_tedtalk_cache:  dict[int, dict]  = {}
_processed_msgs: set[int]         = set()  # dedup: prevent double-processing

# ── Logging helper ────────────────────────────────────────────────────────────
def log_error(location: str, e: Exception):
    print(f"[ERROR:{location}] {type(e).__name__}: {e}")


def debug_event(tag: str, detail: str):
    print(f"[DEBUG:{tag}] {detail}")


async def _provider_pause_reply(
    provider: str,
    *,
    user_id: int = 0,
    channel_id: int = 0,
    is_dm: bool = False,
    direct_to_me: bool = True,
) -> str:
    if not (is_dm or direct_to_me):
        return ""
    scope = user_id or channel_id or 0
    cooldown = max(45, PROVIDER_PAUSE_COOLDOWN_S // 2) if is_dm else PROVIDER_PAUSE_COOLDOWN_S
    if scope:
        allowed, remaining = await mem.consume_phrase_with_status(
            f"{BOT_NAME}:{provider}:pause:{scope}",
            "provider_pause",
            cooldown,
        )
        if not allowed:
            debug_event("provider", f"{BOT_NAME} suppressed repeated {provider} pause reply remaining={remaining}s")
            return ""
    remaining = ai.exhausted_remaining() if provider == "groq" else vision_exhausted_remaining()
    pool = [
        "Not now. Ask again when the static clears.",
        "I heard you. Give me a minute.",
        f"The line is clogged. Wait {max(1, remaining // 60)} minute{'s' if max(1, remaining // 60) != 1 else ''}.",
        "Enough. Let the noise die down first.",
        "I'm not ignoring you. The channel is jammed.",
    ] if not is_dm else [
        "Not now. Stay here a minute.",
        "I heard you. Give me a little time.",
        "The line is jammed. Ask me again shortly.",
        "Wait. I'm not in the mood to fight the static for scraps.",
    ]
    return await _pick_fresh_pool_line(pool, channel_id=channel_id or scope, user_id=user_id or scope)


async def _vision_image_reply(
    *,
    prompt: str,
    system: str,
    image_bytes: bytes,
    mime_type: str,
    max_chars: int = 900,
) -> str:
    if vision_is_exhausted():
        debug_event("provider", f"{BOT_NAME} skipping xAI vision call for {vision_exhausted_remaining()}s")
        return ""
    loop = asyncio.get_event_loop()

    def _run():
        return ask_character_bot(
            BOT_NAME,
            prompt,
            image_bytes=image_bytes,
            mime_type=mime_type,
            system_prompt=system,
            temperature=0.35,
        )

    try:
        reply = await loop.run_in_executor(None, _run)
    except Exception as e:
        if _is_rate_limited_error(e):
            debug_event("provider", f"{BOT_NAME} xAI vision rate-limited")
            return ""
        raise
    return strip_narration((reply or "").strip())[:max_chars]


def _face_feature_unavailable_text() -> str:
    return "Hmph. The face matcher is not installed yet. Add `opencv-python-headless` and `numpy`, then redeploy me."


def _face_enroll_failure_text(reason: str) -> str:
    if reason == "no_face":
        return "If you expect me to remember your face, send an image where it is actually visible."
    if reason == "decode_failed":
        return "Whatever that was, it was not usable."
    return "That was not enough for a reliable face match."


def _face_enroll_success_text(sample_count: int) -> str:
    if sample_count <= 1:
        return "Hmph. Fine. I'll remember that face."
    if sample_count <= 3:
        return "Another angle? Fine. I remember you more clearly now."
    return "You have been persistent about this. Fine. Your face is fixed in memory now."


def _face_delete_text() -> str:
    return "Gone. I won't be keeping your face in memory any longer."


def _face_prompt_note(match_info: dict | None, *, requested: bool = False) -> str:
    if not match_info:
        return ""
    if match_info.get("ok") and match_info.get("matched"):
        if match_info.get("status") == "confirmed":
            return "OWNER_FACE_MATCH: You recognize the owner's face here with high confidence. Mention it naturally once if relevant."
        return "OWNER_FACE_MATCH: This likely shows the owner. Mention that naturally only if it fits."
    if requested and match_info.get("reason") != "not_enrolled":
        return "OWNER_FACE_MATCH: You are not confident this shows the owner. Say so plainly."
    return ""


async def _load_face_attachment(message):
    img = next((a for a in message.attachments if a.content_type and "image" in a.content_type), None)
    vid = next((a for a in message.attachments
               if (a.content_type and a.content_type in VIDEO_TYPES) or
                  any(a.filename.lower().endswith(ext) for ext in VIDEO_EXTS)), None)
    if not img and not vid and message.reference:
        try:
            ref_msg = await message.channel.fetch_message(message.reference.message_id)
            img = next((a for a in ref_msg.attachments if a.content_type and "image" in a.content_type), None)
            vid = next((a for a in ref_msg.attachments
                       if (a.content_type and a.content_type in VIDEO_TYPES) or
                          any(a.filename.lower().endswith(ext) for ext in VIDEO_EXTS)), None)
        except Exception:
            pass
    return img, vid


async def _recent_reply_samples(channel_id: int | None = None, user_id: int | None = None) -> list[str]:
    try:
        channel_recent = await mem.get_recent_assistant_messages(limit=18, channel_id=channel_id) if channel_id is not None else []
        user_recent = await mem.get_recent_assistant_messages(limit=12, user_id=user_id) if user_id is not None else []
        global_recent = await mem.get_recent_assistant_messages(limit=20)
        runtime_recent = get_runtime_recent(BOT_NAME, limit=20)
        return merge_recent_messages(channel_recent, user_recent, global_recent, runtime_recent, limit=40)
    except Exception as e:
        log_error("recent_reply_samples", e)
        return get_runtime_recent(BOT_NAME, limit=20)


def _sentence_count(text: str) -> int:
    cleaned = strip_narration((text or "").strip())
    if not cleaned:
        return 0
    return len([part for part in re.split(r"(?<=[.!?])\s+", cleaned) if part.strip()])


def _paragraph_count(text: str) -> int:
    cleaned = (text or "").strip()
    if not cleaned:
        return 0
    return len([part for part in re.split(r"\n\s*\n", cleaned) if part.strip()])


def _is_large_text_reply(text: str) -> bool:
    cleaned = strip_narration((text or "").strip())
    if not cleaned:
        return False
    return (
        len(cleaned) >= LONG_REPLY_CHAR_THRESHOLD
        or _sentence_count(cleaned) >= LONG_REPLY_SENTENCE_THRESHOLD
        or _paragraph_count(text) >= LONG_REPLY_PARAGRAPH_THRESHOLD
    )


async def _recent_text_pressure(channel_id: int) -> dict[str, int]:
    try:
        recent = await mem.get_recent_assistant_messages(limit=max(RECENT_REPLY_PACING_WINDOW, 10), channel_id=channel_id)
        sample = recent[:RECENT_REPLY_PACING_WINDOW]
        return {
            "count": len(sample),
            "large": sum(1 for item in sample if _is_large_text_reply(item)),
            "walls": sum(1 for item in sample if _paragraph_count(item) >= LONG_REPLY_PARAGRAPH_THRESHOLD),
        }
    except Exception as e:
        log_error("recent_text_pressure", e)
        return {"count": 0, "large": 0, "walls": 0}


def _allow_long_text(scene_tag: str, user: dict | None, duo_mode: str = "") -> bool:
    if scene_tag in {"emotional_comfort", "relationship_progression", "combat_action", "lore_discussion"}:
        return True
    if duo_mode in {"trial", "mission", "interrogate", "duet", "argue", "compare", "truthdare"}:
        return True
    if user and (user.get("conflict_open") or user.get("repair_progress", 0) > 0 or user.get("mood", 0) <= -6):
        return True
    return False


def _tighten_length_hint(hint: str, pressure: dict[str, int], allow_long_text: bool) -> str:
    large = pressure.get("large", 0)
    walls = pressure.get("walls", 0)
    if large <= 1 and walls == 0:
        return hint
    if allow_long_text:
        if walls >= 2 and hint.lower().startswith("longer"):
            return "A few sentences."
        return hint
    if walls >= 2 or large >= 4:
        return "One sentence."
    if large >= 2 and hint in {"A few sentences.", "Longer, dramatic.", "Longer, thoughtful."}:
        return "2-3 sentences."
    return hint


def _tighten_token_budget(max_tokens: int, pressure: dict[str, int], allow_long_text: bool) -> int:
    large = pressure.get("large", 0)
    walls = pressure.get("walls", 0)
    if allow_long_text:
        return min(max_tokens, 320) if large >= 3 else max_tokens
    if walls >= 2:
        return min(max_tokens, 160)
    if large >= 3:
        return min(max_tokens, 200)
    if large >= 2:
        return min(max_tokens, 240)
    return max_tokens


def _prefer_voice_for_reply(reply: str, pressure: dict[str, int], *, asked_for_voice: bool, allow_long_text: bool) -> bool:
    if asked_for_voice:
        return True
    if not _is_large_text_reply(reply):
        return False
    if pressure.get("walls", 0) >= 1:
        return True
    if pressure.get("large", 0) >= 2:
        return True
    return allow_long_text and _paragraph_count(reply) >= LONG_REPLY_PARAGRAPH_THRESHOLD


def _compact_text_for_pacing(reply: str, *, allow_long_text: bool) -> str:
    cleaned = strip_narration(re.sub(r"\n\s*\n+", " ", (reply or "").strip()))
    if not cleaned:
        return cleaned
    max_sentences = 4 if allow_long_text else 2
    max_chars = 520 if allow_long_text else 260
    sentences = [part.strip() for part in re.split(r"(?<=[.!?])\s+", cleaned) if part.strip()]
    if len(sentences) > max_sentences:
        cleaned = " ".join(sentences[:max_sentences])
    if len(cleaned) > max_chars:
        clipped = cleaned[:max_chars].rsplit(" ", 1)[0].rstrip(".,;: ")
        cleaned = (clipped or cleaned[:max_chars]).rstrip() + "..."
    return cleaned


async def _pick_fresh_pool_line(options: list[str], channel_id: int | None = None, user_id: int | None = None) -> str:
    recent = await _recent_reply_samples(channel_id=channel_id, user_id=user_id)
    line = pick_fresh_option(BOT_NAME, options, recent)
    remember_output(BOT_NAME, line)
    return line


def _voice_note_needs_rewrite(text: str, recent_messages: list[str]) -> bool:
    cleaned = strip_narration((text or "").strip())
    if not cleaned:
        return False
    return (
        len(cleaned) >= 96
        or _sentence_count(cleaned) >= 3
        or cleaned.count(",") >= 4
        or ";" in cleaned
        or ":" in cleaned
        or looks_repetitive(cleaned, recent_messages[:12])
    )


async def _rewrite_voice_note_for_speech(
    draft: str,
    *,
    recent_messages: list[str],
    user: dict | None = None,
    max_tokens: int = 180,
) -> str:
    cleaned = strip_narration((draft or "").strip())
    if not cleaned or ai.is_exhausted() or not _voice_note_needs_rewrite(cleaned, recent_messages):
        return cleaned
    prompt = (
        "You are revising one Discord voice note before it is spoken aloud.\n"
        "Character: Scaramouche, pre-Sumeru, sharp and theatrical but natural when spoken.\n"
        "Goals: make it sound spoken instead of written, vary the wording, reduce repetitive cadence, shorten stacked clauses, and keep it natural aloud.\n"
        "Keep the same meaning and relationship intensity. No narration. No quote marks. One to four spoken sentences.\n"
        "Attribution rules: speak only for yourself as fact. Only claim Wanderer said or believes something if it was explicitly stated in context. Otherwise frame it as inference with phrasing like 'if he thinks', 'I assume', or 'it sounds like'.\n"
        f"Relationship state: affection={(user or {}).get('affection', 0)} trust={(user or {}).get('trust', 0)} conflict_open={bool((user or {}).get('conflict_open'))}\n"
        f"Draft: {cleaned}\n"
        "Return only the revised spoken version."
    )
    loop = asyncio.get_event_loop()
    model_name = _select_text_model(route="light")
    rewritten = await loop.run_in_executor(None, _qai_blocking, prompt, min(max_tokens, 180), model_name)
    rewritten = strip_narration((rewritten or "").strip())
    return rewritten or cleaned


async def _prepare_voice_reply_text(
    text: str,
    *,
    channel_id: int | None = None,
    user_id: int | None = None,
    mood: int = 0,
    conflict_open: bool = False,
    user: dict | None = None,
) -> str:
    updated = strip_narration((text or "").strip())
    if not updated:
        return updated

    recent = await _recent_reply_samples(channel_id=channel_id, user_id=user_id)
    updated = await _rewrite_voice_note_for_speech(
        updated,
        recent_messages=recent,
        user=user,
    )
    updated = diversify_reply(BOT_NAME, updated, recent)
    updated = await _apply_phrase_policy(
        updated,
        recent,
        user_id=user_id,
        mood=mood,
        conflict_open=conflict_open,
    )
    updated = _sanitize_partner_attribution(updated)
    updated = _sanitize_time_of_day_claims(updated, user)
    if looks_repetitive(updated, recent):
        fallback = fallback_reply(BOT_NAME, recent)
        if fallback:
            return fallback
    return updated


async def _apply_phrase_policy(
    text: str,
    recent_messages: list[str] | None = None,
    user_id: int | None = None,
    mood: int = 0,
    conflict_open: bool = False,
) -> str:
    recent_messages = recent_messages or []
    updated = strip_narration((text or "").strip())
    if not updated:
        return updated

    opening = detect_opening_phrase(BOT_NAME, updated)
    if not opening:
        return updated
    if opening == "how quaint":
        return replace_opening_phrase(BOT_NAME, updated, recent_messages)

    rule = BOT_RARE_PHRASES.get(opening)
    if not rule:
        return updated

    scopes = [f"{BOT_NAME}:global"]
    if user_id is not None:
        scopes.insert(0, f"{BOT_NAME}:user:{user_id}")
    cooldown = int(rule.get("cooldown", 0))
    allowed = True
    if cooldown > 0:
        for scope in scopes:
            allowed, remaining = await mem.consume_phrase_with_status(scope, f"{BOT_NAME}:{opening}", cooldown)
            if not allowed:
                debug_event("phrase", f"{BOT_NAME} blocked '{opening}' scope={scope} remaining={remaining}s")
                break
    if allowed:
        debug_event("phrase", f"{BOT_NAME} allowed '{opening}' scopes={','.join(scopes)}")
        return updated
    debug_event("phrase", f"{BOT_NAME} diversified opener '{opening}'")
    return replace_opening_phrase(BOT_NAME, updated, recent_messages)


async def _learn_user_state(user_id: int, user_message: str):
    try:
        current = await mem.get_user(user_id)
        if not current:
            return
        profile = apply_style_deltas(current.get("style_profile"), analyze_style_deltas(user_message))
        await mem.set_style_profile(user_id, profile)
        debug_event("memory", f"{BOT_NAME} style_profile user={user_id} traits={','.join(sorted([k for k, v in profile.items() if v >= 8])[:3]) or 'none'}")

        callback_memory = extract_callback_candidate(user_message)
        if callback_memory:
            await mem.set_callback_memory(user_id, callback_memory)
            debug_event("memory", f"{BOT_NAME} callback user={user_id} text={callback_memory[:80]}")
        for topic in detect_topics(user_message):
            await mem.record_topic(user_id, topic)
            debug_event("memory", f"{BOT_NAME} topic user={user_id} topic={topic}")

        if detect_repair_signal(user_message) and current.get("conflict_open"):
            resolved = await mem.record_repair_attempt(user_id, user_message[:180])
            await mem.update_trust(user_id, +2)
            await mem.update_affection(user_id, +1)
            await mem.set_callback_memory(user_id, f"They tried to repair things: {user_message[:180]}")
            debug_event("memory", f"{BOT_NAME} conflict_repair user={user_id} resolved={resolved}")
        elif detect_conflict_signal(user_message) and (current.get("romance_mode") or current.get("affection", 0) >= 30):
            await mem.open_conflict(user_id, user_message[:180])
            debug_event("memory", f"{BOT_NAME} conflict_opened user={user_id} text={user_message[:80]}")

        refreshed = await mem.get_user(user_id)
        if not refreshed:
            return
        arc = compute_emotional_arc(
            refreshed.get("affection", 0),
            refreshed.get("trust", 0),
            refreshed.get("slow_burn", 0),
            refreshed.get("conflict_open", False),
            refreshed.get("repair_count", 0),
        )
        await mem.set_emotional_arc(user_id, arc)
        debug_event("memory", f"{BOT_NAME} arc user={user_id} arc={arc}")
    except Exception as e:
        log_error("learn_user_state", e)


async def _rebuild_apply_user_message(
    user_id: int,
    channel_id: int,
    username: str,
    display_name: str,
    user_message: str,
):
    try:
        await mem.upsert_user(user_id, username, display_name)
        current = await mem.get_user(user_id)
        if not current:
            return

        await mem.add_message(user_id, channel_id, "user", user_message)
        if _should_log_quote_evidence(user_message):
            await mem.add_evidence_item(
                channel_id,
                "quote",
                f"{display_name} line",
                user_message[:220],
                source_excerpt=user_message[:180],
                owner_user_id=user_id,
                updated_by=BOT_NAME,
            )

        msg_l = user_message.lower()
        scenario = detect_scenario(user_message, is_dm=(channel_id == user_id))
        triggers = detect_emotional_triggers(user_message)

        anger_hit = _scara_anger_hit(current, msg_l)
        if any(k in msg_l for k in RUDE_KW) or anger_hit:
            await mem.update_mood(user_id, -2)
            await mem.update_trust(user_id, -1)
            if anger_hit:
                await mem.update_affection(user_id, -anger_hit["affection_drop"])
                await mem.update_trust(user_id, -anger_hit["trust_drop"])
                await mem.update_mood(user_id, -anger_hit["mood_drop"])
                await mem.update_anger(user_id, anger_hit["anger_gain"])
                await mem.add_consequence_mark(user_id, "insult", user_message[:180], severity=anger_hit["severity"], decay_days=30)
        elif any(k in msg_l for k in ROMANCE_KW):
            if current and not current.get("romance_mode", False):
                await mem.set_mode(user_id, "romance_mode", True)
            await mem.update_mood(user_id, +1)
            await mem.update_affection(user_id, +1)
            await mem.update_trust(user_id, +1)
            await mem.update_drift(user_id, +1)
            await mem.increment_slow_burn(user_id)
            await mem.update_last_statement(user_id, user_message[:200])
        elif any(k in msg_l for k in NICE_KW):
            await mem.update_mood(user_id, +1)
            await mem.update_affection(user_id, +1)
            await mem.update_trust(user_id, +1)
            await mem.update_last_statement(user_id, user_message[:200])
        else:
            positive = sum(
                [
                    any(
                        w in msg_l
                        for w in [
                            "haha",
                            "lol",
                            "lmao",
                            "hehe",
                            "cute",
                            "nice",
                            "cool",
                            "fun",
                            "good",
                            "great",
                            "enjoy",
                            "happy",
                            "excited",
                            "interesting",
                            "wow",
                            "omg",
                            "yes",
                            "yay",
                            "please",
                            "😂",
                            "😭",
                            "❤",
                            "💖",
                            "🥺",
                        ]
                    ),
                    msg_l.endswith("!") and len(user_message) > 8,
                    "?" in user_message and len(user_message) > 15,
                    len(user_message) > 100,
                ]
            )
            negative = sum(
                [
                    any(
                        w in msg_l
                        for w in [
                            "ugh",
                            "ew",
                            "boring",
                            "whatever",
                            "idc",
                            "nope",
                            "wrong",
                            "bad",
                            "hate",
                            "worst",
                            "terrible",
                            "awful",
                            "seriously",
                            "really",
                            "😒",
                            "🙄",
                        ]
                    ),
                    user_message.count("...") > 1,
                ]
            )

            if positive >= 2:
                await mem.update_affection(user_id, +1)
                await mem.update_mood(user_id, +1)
                await mem.update_trust(user_id, +1)
                await mem.update_drift(user_id, +1)
            elif positive == 1:
                await mem.update_affection(user_id, +1)
            elif negative >= 2:
                await mem.update_mood(user_id, -1)
                await mem.update_trust(user_id, -1)
            elif negative == 1:
                await mem.update_mood(user_id, -1)

        if scenario == "emotional_comfort":
            await mem.update_trust(user_id, +1)
            if "softness" in triggers or "protectiveness" in triggers:
                await mem.update_affection(user_id, +1)
        elif scenario == "combat_action":
            await mem.update_mood(user_id, -1)
            await mem.update_trust(user_id, +1)
        elif scenario == "lore_discussion":
            await mem.update_trust(user_id, +1)
            await mem.update_drift(user_id, +1)
        elif scenario == "relationship_progression":
            await mem.update_affection(user_id, +1)
            await mem.update_trust(user_id, +1)
        elif scenario == "introspection":
            await mem.update_trust(user_id, +1)

        if "jealousy" in triggers:
            await mem.update_mood(user_id, -1)
            await mem.update_affection(user_id, +1)
        if "protectiveness" in triggers:
            await mem.update_trust(user_id, +1)
        if "boredom" in triggers:
            await mem.update_mood(user_id, -1)

        repair_signal = detect_repair_signal(user_message)
        conflict_signal = detect_conflict_signal(user_message)
        if repair_signal:
            await mem.soften_consequence_marks(user_id, 1)
        if conflict_signal:
            await mem.add_consequence_mark(user_id, "fight", user_message[:180], severity=3, decay_days=7)

        await _learn_user_state(user_id, user_message)
        for kind, memory_text, weight in extract_memory_events(user_message):
            await mem.add_memory_event(user_id, kind, memory_text, max(weight, _memory_weight_for(kind)))
            if kind == "betrayal":
                await mem.add_consequence_mark(user_id, kind, memory_text, severity=max(4, weight), decay_days=14)
            elif kind in {"fight", "slight"}:
                await mem.add_consequence_mark(user_id, kind, memory_text, severity=max(2, weight), decay_days=8)
            elif kind == "promise":
                await mem.add_consequence_mark(user_id, kind, memory_text, severity=2, decay_days=10)
            elif kind == "repair":
                await mem.soften_consequence_marks(user_id, 1)

        scene_update = infer_scene_update(user_message, display_name)
        if scene_update:
            await mem.update_scene_state(channel_id, **scene_update)
            await _register_world_from_message(channel_id, user_id, user_message, scene_update)
    except Exception as e:
        log_error("rebuild_apply_user_message", e)


async def _rebuild_from_history(ctx, *, target_user_id: int | None = None):
    records, stats = await collect_rebuild_records(ctx, target_user_id=target_user_id)
    if not records:
        return stats | {"rebuilt_users": 0}

    user_ids = [target_user_id] if target_user_id else stats["user_ids"]
    for user_id in user_ids:
        if user_id:
            await mem.reset_user_for_rebuild(user_id)

    if target_user_id is None:
        for channel_id in stats["channel_ids"]:
            await mem.clear_scene_state(channel_id)

    for record in records:
        await _rebuild_apply_user_message(
            record["user_id"],
            record["channel_id"],
            record["username"],
            record["display_name"],
            record["content"],
        )

    return stats | {"rebuilt_users": len(user_ids)}


async def _observe_partner_message(content: str) -> tuple[dict, list[dict], str]:
    theme = detect_banter_theme(content)
    await mem.record_bot_banter(PARTNER_PAIR_KEY, PARTNER_NAME, content, theme)
    relation = await mem.get_bot_relationship(PARTNER_PAIR_KEY)
    respect_delta, tension_delta = infer_bot_relation_deltas(content, theme)
    respect = relation.get("respect", 0) + respect_delta
    tension = relation.get("tension", 0) + tension_delta
    stage = compute_bot_stage(respect, tension)
    note = None
    if stage != relation.get("stage"):
        note = f"The rivalry has shifted into {stage} after circling {theme} too many times."
    await mem.update_bot_relationship(
        PARTNER_PAIR_KEY,
        stage,
        respect,
        tension,
        theme=theme,
        history_note=note,
        touched_exchange=False,
    )
    relation = await mem.get_bot_relationship(PARTNER_PAIR_KEY)
    recent = await mem.get_recent_bot_banter(PARTNER_PAIR_KEY, 8)
    milestone_note = relationship_milestone_note(stage, relation.get("respect", 0), relation.get("tension", 0))
    marker = f"pair:{stage}"
    if milestone_note and not await mem.has_milestone(PARTNER_PAIR_KEY, marker):
        await mem.add_milestone(PARTNER_PAIR_KEY, marker, milestone_note)
        debug_event("relationship", f"{BOT_NAME} milestone marker={marker} note={milestone_note[:90]}")
    debug_event("relationship", f"{BOT_NAME} partner stage={relation.get('stage')} respect={relation.get('respect')} tension={relation.get('tension')} theme={theme}")
    return relation, recent, theme


_PARTNER_REFERENCES = ("wanderer", "the wanderer")


def _message_mentions_partner(text: str) -> bool:
    lowered = (text or "").lower()
    return any(token in lowered for token in _PARTNER_REFERENCES)


def _event_topic_from_text(text: str, fallback: str = "") -> str:
    cleaned = re.sub(r"\s+", " ", (text or fallback or "").strip())
    if not cleaned:
        return ""
    patterns = [
        r"(?:trial|mission|interrogate|compare|duet|argue|verdict)\s+(?:about|over|regarding)?\s*:?\s*(.+)",
        r"(?:about|over|regarding|remember|evidence for|proof of)\s+(.+)",
    ]
    lowered = cleaned.lower()
    for pattern in patterns:
        match = re.search(pattern, lowered, re.IGNORECASE)
        if match:
            value = re.sub(r"\s+", " ", match.group(1)).strip(" .,!?:;\"'")
            if len(value) >= 6:
                return value[:120]
    return cleaned[:120]


def _should_log_quote_evidence(text: str) -> bool:
    lowered = (text or "").lower()
    return len((text or "").strip()) >= 40 and any(
        token in lowered for token in ["remember this", "don't forget", "proof", "evidence", "quote", "keep this", "this matters"]
    )


def _opinion_about_user(user: dict | None, triangle: dict | None = None) -> tuple[str, int]:
    user = user or {}
    affection = int(user.get("affection", 0) or 0)
    trust = int(user.get("trust", 0) or 0)
    mood = int(user.get("mood", 0) or 0)
    jealousy = int((triangle or {}).get("jealousy_level", 0) or 0)
    if affection >= 70 and trust >= 65:
        return ("Too important to dismiss cleanly anymore. Their weakness and loyalty are both worth keeping track of.", 8)
    if jealousy >= 45:
        return ("They keep testing my patience by drifting elsewhere. Irritating. Not irrelevant.", 7)
    if trust >= 45:
        return ("Useful, persistent, and harder to treat like background noise than they should be.", 6)
    if mood <= -5:
        return ("A nuisance today. Still, I've learned not to ignore what they reveal under pressure.", 5)
    return ("Potential leverage. Potential entertainment. Still under observation.", 4)


def _opinion_about_partner(relation: dict | None) -> tuple[str, int]:
    relation = relation or {}
    stage = relation.get("stage", "enemy")
    respect = int(relation.get("respect", 0) or 0)
    tension = int(relation.get("tension", 0) or 0)
    if stage == "reluctant respect" or respect >= 55:
        return ("Still infuriating, but precise enough that ignoring him would be lazy.", 7)
    if stage == "competitive":
        return ("Annoyingly capable. Too eager to pretend he's above the argument.", 6)
    if tension >= 75:
        return ("He keeps pressing old wounds and calling it honesty. I notice every bit of it.", 7)
    return ("A familiar irritation with a talent for surviving longer than expected.", 5)


async def _sync_private_opinions(user_id: int, user: dict | None, triangle: dict | None = None, relation: dict | None = None):
    try:
        user_opinion, user_intensity = _opinion_about_user(user, triangle)
        await mem.set_private_opinion(f"user:{user_id}", BOT_NAME, "user", str(user_id), user_opinion, intensity=user_intensity)
        partner_opinion, partner_intensity = _opinion_about_partner(relation or await mem.get_bot_relationship(PARTNER_PAIR_KEY))
        await mem.set_private_opinion(f"user:{user_id}", BOT_NAME, "partner", PARTNER_NAME, partner_opinion, intensity=partner_intensity)
    except Exception as e:
        log_error("sync_private_opinions", e)


async def _contradictory_memory_context(channel_id: int, user_message: str = "", topic: str = "") -> str:
    try:
        topic_hint = _event_topic_from_text(user_message, fallback=topic)
        memories = await mem.get_shared_event_memory(channel_id, topic_hint, limit=1)
        if not memories and topic_hint:
            memories = await mem.get_shared_event_memory(channel_id, "", limit=1)
        if not memories:
            return ""
        event_memory = memories[0]
        lowered = (user_message or topic or "").lower()
        if (
            event_memory.get("event_key")
            and not event_memory.get("distorted_bot")
            and any(token in lowered for token in ["irminsul", "memory", "remember", "actually", "what happened"])
            and random.random() < 0.08
        ):
            chosen = BOT_NAME if random.random() < 0.5 else PARTNER_NAME
            await mem.mark_memory_distortion(event_memory["event_key"], chosen)
            event_memory["distorted_bot"] = chosen
        return describe_contradictory_memory(BOT_NAME, event_memory, PARTNER_NAME, current_text=user_message or topic)
    except Exception as e:
        log_error("contradictory_memory_context", e)
        return ""


async def _private_opinion_prompt_context(user_id: int, *, leak_partner: bool = False) -> list[str]:
    parts: list[str] = []
    try:
        opinions = await mem.list_private_opinions(f"user:{user_id}", 6)
        own = next((item for item in opinions if item.get("bot_name") == BOT_NAME and item.get("subject_type") == "user"), None)
        if own:
            line = describe_private_opinion_context(BOT_NAME, own)
            if line:
                parts.append(line)
        if leak_partner:
            leaked = next(
                (
                    item for item in opinions
                    if item.get("bot_name") == PARTNER_NAME
                    and (time.time() - float(item.get("leaked_ts", 0) or 0)) > 21600
                ),
                None,
            )
            if leaked and random.random() < 0.14:
                line = describe_private_opinion_context(BOT_NAME, leaked, leaked=True)
                if line:
                    parts.append(line)
                    await mem.mark_private_opinion_leaked(
                        f"user:{user_id}",
                        leaked.get("bot_name", ""),
                        leaked.get("subject_type", ""),
                        leaked.get("subject_key", ""),
                    )
    except Exception as e:
        log_error("private_opinion_prompt_context", e)
    return parts


async def _maybe_seed_silent_presence(channel_id: int, mode: str):
    if mode not in {"duet", "compare", "trial", "mission", "interrogate"}:
        return
    if random.random() >= 0.24:
        return
    silent_bot = random.choice([BOT_NAME, PARTNER_NAME])
    await mem.set_duo_presence(
        channel_id,
        silent_bot=silent_bot,
        silent_until=time.time() + random.randint(180, 420),
    )
    debug_event("relationship", f"{BOT_NAME} silent_presence channel={channel_id} bot={silent_bot}")


async def _partner_prompt_context(user_message: str, channel_id: int = 0) -> str:
    if not _message_mentions_partner(user_message):
        return ""
    relation = await mem.get_bot_relationship(PARTNER_PAIR_KEY)
    recent_banter = await mem.get_recent_bot_banter(PARTNER_PAIR_KEY, 6)
    lines = [
        describe_bot_relationship(BOT_NAME, relation, recent_banter),
        "ATTRIBUTION_RULES: speak only for yourself as fact. Quote Wanderer only if his words are explicitly present in recent context. Otherwise frame guesses about him as inference, not certainty.",
    ]
    contradiction = await _contradictory_memory_context(channel_id, user_message=user_message)
    if contradiction:
        lines.append(contradiction)
    return "\n".join(line for line in lines if line)


async def _duo_prompt_context(channel_id: int, user_message: str = "") -> str:
    session = await mem.get_duo_session(channel_id)
    if not session:
        return ""
    mode = session.get("mode", "both")
    topic = session.get("topic", "")
    last_speaker = session.get("last_speaker", "")
    prompt = [f"DUO_SESSION:{mode}|topic={topic[:180]}"]
    stage_note = describe_duo_scene_stage(mode, topic, int(session.get("autoplay_remaining", 0) or 0))
    if stage_note:
        prompt.append(stage_note)
    dual_verdict = describe_dual_verdict_mode(BOT_NAME, mode, topic)
    if dual_verdict:
        prompt.append(dual_verdict)
    silent_presence = describe_silent_presence(
        BOT_NAME,
        mode,
        topic,
        session.get("silent_bot", ""),
        int(session.get("autoplay_remaining", 0) or 0),
    )
    if silent_presence:
        prompt.append(silent_presence)
    open_story = await mem.get_open_duo_story(channel_id, mode)
    if open_story and open_story.get("outcome"):
        prompt.append(f"DUO_PROGRESS:{open_story['outcome'][:180]}")
    contradiction = await _contradictory_memory_context(channel_id, user_message=user_message, topic=topic)
    if contradiction:
        prompt.append(contradiction)
    evidence_items = await mem.list_evidence_items(channel_id, limit=4)
    evidence_context = describe_evidence_locker(BOT_NAME, evidence_items, f"{topic} {user_message}".strip())
    if evidence_context:
        prompt.append(evidence_context)
    if last_speaker and last_speaker != BOT_NAME:
        prompt.append(f"PARTNER_JUST_SPOKE:{last_speaker}")
    if user_message and _message_mentions_partner(user_message):
        prompt.append("PARTNER_EXPLICITLY_IN_PLAY")
    prompt.append("ATTRIBUTION_RULES: state your own view as fact. If referring to Wanderer, only attribute explicit words or beliefs to him when recent context actually shows them. Otherwise mark it as your inference.")
    if mode == "duet":
        prompt.append("DUO_BEHAVIOR: continue the shared scene naturally and leave a clean opening for the other bot")
    elif mode == "argue":
        prompt.append("DUO_BEHAVIOR: take a sharp stance and escalate the disagreement without repeating the partner")
    elif mode == "compare":
        prompt.append("DUO_BEHAVIOR: answer, then sharpen the contrast between your view and the partner's")
    elif mode == "interrogate":
        prompt.append("DUO_BEHAVIOR: press with one sharp question, accusation, or deduction")
    elif mode == "trial":
        prompt.append("DUO_BEHAVIOR: speak like part of a two-bot prosecution or judgment")
    elif mode == "mission":
        prompt.append("DUO_BEHAVIOR: contribute one tactical role, warning, or leverage point")
    elif mode == "truthdare":
        prompt.append("DUO_BEHAVIOR: escalate the game with one pointed truth or dare prompt")
    else:
        prompt.append("DUO_BEHAVIOR: give one compact turn that pairs well with a second bot response")
    return "\n".join(prompt)


def _partner_autoplay_name() -> str:
    return PARTNER_NAME


def _user_local_hour(user: dict | None) -> int:
    try:
        tz = ZoneInfo((user or {}).get("timezone_name") or "America/Los_Angeles")
        return datetime.now(tz).hour
    except Exception:
        return datetime.now().hour


def _time_period_label(hour: int) -> str:
    if 5 <= hour < 12:
        return "morning"
    if 12 <= hour < 17:
        return "afternoon"
    if 17 <= hour < 21:
        return "evening"
    return "night"


def _sanitize_time_of_day_claims(text: str, user: dict | None) -> str:
    cleaned = (text or "").strip()
    if not cleaned:
        return cleaned

    actual = _time_period_label(_user_local_hour(user))
    wrong_patterns: dict[str, tuple[str, ...]] = {
        "morning": (
            r"\bwhat(?:'s| is)\s+so\s+special\s+about\s+mornings?\b",
            r"\bmorning\s+has\s+barely\s+started\b",
            r"\byou(?:'re| are)\s+up\s+early\b",
            r"\bthis\s+morning\b",
            r"\bgood\s+morning\b",
        ),
        "afternoon": (
            r"\bthis\s+afternoon\b",
            r"\bgood\s+afternoon\b",
        ),
        "evening": (
            r"\bthis\s+evening\b",
            r"\bgood\s+evening\b",
        ),
        "night": (
            r"\btonight\b",
            r"\bthis\s+late\s+night\b",
            r"\bgood\s+night\b",
        ),
    }

    for label, patterns in wrong_patterns.items():
        if label == actual:
            continue
        for pattern in patterns:
            if re.search(pattern, cleaned, re.IGNORECASE):
                if label == "morning" and actual == "night":
                    cleaned = re.sub(
                        r"\bwhat(?:'s| is)\s+so\s+special\s+about\s+mornings?\b",
                        "What's so special about showing up this late",
                        cleaned,
                        flags=re.IGNORECASE,
                    )
                    cleaned = re.sub(r"\bmorning\s+has\s+barely\s+started\b", "It's late and you're still here", cleaned, flags=re.IGNORECASE)
                    cleaned = re.sub(r"\byou(?:'re| are)\s+up\s+early\b", "you're still awake", cleaned, flags=re.IGNORECASE)
                    cleaned = re.sub(r"\bthis\s+morning\b", "this late at night", cleaned, flags=re.IGNORECASE)
                    cleaned = re.sub(r"\bgood\s+morning\b", "You're up late", cleaned, flags=re.IGNORECASE)
                elif label == "morning":
                    cleaned = re.sub(r"\bthis\s+morning\b", f"this {actual}", cleaned, flags=re.IGNORECASE)
                    cleaned = re.sub(r"\bgood\s+morning\b", f"Good {actual}", cleaned, flags=re.IGNORECASE)
                elif label == "night" and actual != "night":
                    cleaned = re.sub(r"\btonight\b", f"this {actual}", cleaned, flags=re.IGNORECASE)
                else:
                    cleaned = re.sub(rf"\bthis\s+{label}\b", f"this {actual}", cleaned, flags=re.IGNORECASE)
                    cleaned = re.sub(rf"\bgood\s+{label}\b", f"Good {actual}", cleaned, flags=re.IGNORECASE)
    return cleaned


def _sanitize_partner_attribution(text: str) -> str:
    cleaned = (text or "").strip()
    if not cleaned:
        return cleaned
    replacements = [
        (
            r"\bWanderer,\s*you think ([^?]+)\?\s*I've\b",
            r"Wanderer, if you think \1, I've",
        ),
        (
            r"\bWanderer,\s*you think ([^?]+)\?",
            r"Wanderer, if you think \1, say it plainly.",
        ),
        (
            r"\bWanderer said ([^.?!]+)",
            r"So this is what I am pinning on Wanderer: \1",
        ),
        (
            r"\bWanderer thinks ([^.?!]+)",
            r"If Wanderer thinks \1",
        ),
        (
            r"\bWanderer feels ([^.?!]+)",
            r"If Wanderer feels \1",
        ),
        (
            r"\bWanderer wants ([^.?!]+)",
            r"If Wanderer wants \1",
        ),
        (
            r"\bWanderer finds ([^.?!]+)",
            r"If Wanderer finds \1",
        ),
    ]
    for pattern, replacement in replacements:
        cleaned = re.sub(pattern, replacement, cleaned, flags=re.IGNORECASE)
    return cleaned


def _voice_style_for(
    user: dict | None,
    mood: int = 0,
    scene_tag: str = "",
    *,
    is_dm: bool = False,
    duo_mode: str = "",
    jealousy_level: int = 0,
) -> str:
    arc = (user or {}).get("emotional_arc", "guarded")
    if scene_tag == "combat_action":
        return "combat"
    if duo_mode in {"argue", "compare", "truthdare"}:
        return "duo_teasing"
    if jealousy_level >= 45:
        return "jealous"
    if (user or {}).get("repair_progress", 0) > 0 and mood > -4:
        return "repair"
    if is_dm and _user_local_hour(user) in {0, 1, 2, 3, 4}:
        return "late_night"
    if (user or {}).get("conflict_open"):
        return "tense"
    if arc in {"tender", "attached"} or ((user or {}).get("affection", 0) >= 70):
        return "soft"
    if mood <= -5 or arc == "conflicted":
        return "cutting"
    if arc in {"drawn_in", "trusting"}:
        return "measured"
    return "guarded"


def _utility_reply(subject: str, facts: list[str], character_line: str, sources: str = "") -> str:
    body = [f"{subject}"]
    body.extend(f"- {fact}" for fact in facts if fact)
    if character_line:
        body.append(f"Comment: {character_line}")
    if sources:
        body.append(sources)
    return "\n".join(body)


async def _maybe_jealous_partner_interjection(
    message: discord.Message,
    user: dict | None,
    *,
    partner_name: str,
) -> bool:
    if not message.guild:
        return False
    if not user or int(user.get("affection", 0) or 0) <= JEALOUS_REPLY_AFFECTION_MIN:
        return False
    if random.random() >= JEALOUS_REPLY_CHANCE:
        return False
    try:
        allowed, _ = await mem.consume_phrase_with_status(
            f"{BOT_NAME}:jealous:{message.channel.id}:{message.author.id}",
            "partner_interjection",
            JEALOUS_REPLY_COOLDOWN_S,
        )
        if not allowed:
            return False
    except Exception as e:
        log_error("jealous_interjection/cooldown", e)
        return False

    jealousy_prompt = (
        f"The user was talking to {partner_name}, not you. "
        f"Their line was: {message.content[:220]!r}. "
        "Affection is high enough that a brief jealous interruption slips out. "
        "Cut in with one short in-character reply. Keep it sharp, a little possessive or irritated, and do not narrate."
    )
    try:
        reply = await get_response(
            message.author.id,
            message.channel.id,
            jealousy_prompt,
            user,
            message.author.display_name,
            message.author.mention,
            extra_context=(
                f"JEALOUS_INTERJECTION: {partner_name} currently has the user's attention. "
                "You were not directly addressed, but your affection is above 50 and jealousy leaks through. "
                "Reply to the user's message itself, not to the other bot. One or two sentences only."
            ),
            channel_obj=message.channel,
            is_dm=False,
            direct_to_me=False,
        )
        reply = strip_narration(reply or "")
        if not reply:
            return False
        await message.reply(reply, mention_author=False)
        return True
    except Exception as e:
        log_error("jealous_interjection/reply", e)
        return False


_SELF_EDIT_MODERN = (
    "lowkey", "highkey", "bestie", "bro", "bruh", "fr ", " fr", "irl", "ngl", "slay",
    "no cap", "vibes", "bestie", "mid", "sus", "yolo",
)
_SELF_EDIT_GENERIC_PATTERNS = (
    r"^(that('s| is) interesting)\b",
    r"^(i understand)\b",
    r"^(i'm here for you)\b",
    r"^(of course)\b",
    r"^(sure[, ]|certainly\b|absolutely\b)",
    r"^(it depends)\b",
    r"^(thanks for sharing)\b",
)


def _format_consequence_summary(marks: list[dict]) -> str:
    bits = []
    for item in marks[:4]:
        label = (item.get("kind") or "scar")[:18]
        summary = (item.get("summary") or "")[:90]
        bits.append(f"{label}[{item.get('remaining', 0)}]: {summary}")
    return " || ".join(bits)


def _format_world_prompt(
    entities: list[dict],
    cases: list[dict],
    campaign_npcs: list[dict] | None = None,
    artifacts: list[dict] | None = None,
    evidence_items: list[dict] | None = None,
) -> str:
    lines = []
    if entities:
        lines.append(
            "WORLD_ENTITIES:"
            + " || ".join(
                f"{item.get('entity_type', 'entity')}:{item.get('name', '')[:50]}|{item.get('status', 'noted') or 'noted'}|{item.get('summary', '')[:90]}"
                for item in entities[:4]
            )
        )
    if cases:
        lines.append(
            "WORLD_CASES:"
            + " || ".join(
                f"{item.get('case_type', 'case')}:{item.get('title', '')[:60]}|{item.get('status', 'open')}|enemy={item.get('enemy', '')[:40]}|{item.get('summary', '')[:90]}"
                for item in cases[:3]
            )
        )
    if campaign_npcs:
        lines.append(
            "CAMPAIGN_NPCS:"
            + " || ".join(
                f"{item.get('entity_type', 'npc')}:{item.get('name', '')[:50]}|{item.get('status', 'active')}|{item.get('summary', '')[:90]}"
                for item in campaign_npcs[:4]
            )
        )
    if artifacts:
        lines.append(
            "ATTACHMENT_JOURNAL:"
            + " || ".join(
                f"{item.get('name', '')[:50]}|{item.get('status', 'recent')}|{item.get('summary', '')[:90]}"
                for item in artifacts[:3]
            )
        )
    if evidence_items:
        lines.append(
            "EVIDENCE_LOCKER:"
            + " || ".join(
                f"{item.get('evidence_type', 'evidence')}:{item.get('label', '')[:50]}|{item.get('summary', '')[:90]}"
                for item in evidence_items[:4]
            )
        )
    return "\n".join(lines)


def _world_lines(entities: list[dict]) -> list[str]:
    if not entities:
        return ["Shared world: nothing persistent yet."]
    lines = ["Shared world:"]
    for item in entities[:6]:
        summary = (item.get("summary") or "").strip()
        status = item.get("status") or "noted"
        line = f"- {item.get('entity_type', 'entity')}: {item.get('name', 'unknown')} [{status}]"
        if summary:
            line += f" - {summary[:120]}"
        lines.append(line)
    return lines


def _case_lines(cases: list[dict]) -> list[str]:
    if not cases:
        return ["Shared cases: none open enough to matter."]
    lines = ["Shared cases:"]
    for item in cases[:6]:
        bit = f"- {item.get('case_type', 'case')}: {item.get('title', 'untitled')} [{item.get('status', 'open')}]"
        if item.get("enemy"):
            bit += f" enemy={item['enemy'][:50]}"
        if item.get("summary"):
            bit += f" - {item['summary'][:120]}"
        lines.append(bit)
    return lines


def _relationship_lines(user: dict | None, topics: list[dict], marks: list[dict]) -> list[str]:
    stage, stage_desc = _progression_parts(user)
    arc = _current_arc(user)
    anger = int((user or {}).get("anger_level", 0) or 0)
    ever_maxed = bool((user or {}).get("affection_ever_maxed"))
    anger_aftercare = int((user or {}).get("anger_aftercare_messages", 0) or 0)
    aftermath = describe_conflict_aftermath(
        BOT_NAME,
        user.get("conflict_summary", "") if user else "",
        user.get("last_conflict_ts", 0) if user else 0,
        user.get("repair_progress", 0) if user else 0,
        conflict_open=bool((user or {}).get("conflict_open")),
    ) or "no open aftermath"
    lines = [
        f"Relationship: affection {(user or {}).get('affection', 0)}/100 | trust {(user or {}).get('trust', 0)}/100 | mood {(user or {}).get('mood', 0):+d}",
        f"Arc: {arc} | progression: {stage}",
        f"Progression detail: {stage_desc}",
        f"Conflict aftermath: {aftermath}",
        f"Preferences: voice={_pref_label((user or {}).get('voice_enabled', True))} | utility={_pref_label((user or {}).get('utility_mode', True))} | duoauto={_pref_label((user or {}).get('duo_autoplay', True))} | rpdepth={(user or {}).get('rp_depth', 'medium')}",
    ]
    if ever_maxed or anger > 0:
        lines.insert(1, f"Anger: {anger}/100 | state: {scara_anger_label(anger)} | affection floor: 50/100")
    if anger_aftercare > 0:
        lines.append(f"Anger aftercare: {anger_aftercare} reply{'ies' if anger_aftercare != 1 else 'y'} left | tone: wrong, sad, unsettled")
    if (user or {}).get("anger_repair_active"):
        demand = ((user or {}).get("anger_repair_instruction") or "").strip()
        if demand:
            lines.append("Anger repair: active | demand: " + demand[:180])
        else:
            lines.append("Anger repair: active")
    if topics:
        lines.append("Top topics: " + ", ".join(item["topic"] for item in topics[:3]))
    if marks:
        lines.append("Consequences: " + " | ".join(f"{item['kind']}[{item.get('remaining', 0)}]" for item in marks[:4]))
    return lines


def _duostate_lines(speaker_mode: str, duo: dict | None, relation: dict, stories: list[dict], cases: list[dict]) -> list[str]:
    lines = [
        f"Speaker mode: {speaker_mode}",
        f"Bot relationship: {relation.get('stage', 'enemy')} | respect {relation.get('respect', 0)}/100 | tension {relation.get('tension', 0)}/100",
    ]
    if duo:
        lines.append(
            f"Active duo: mode={duo.get('mode')} | awaiting={duo.get('awaiting_bot') or 'nobody'} | turns_left={duo.get('autoplay_remaining', 0)} | topic={duo.get('topic')}"
        )
        if duo.get("silent_bot") and float(duo.get("silent_until", 0) or 0) > time.time():
            lines.append(f"Silent presence: {duo.get('silent_bot')} staying mostly quiet for this scene")
    else:
        lines.append("Active duo: none")
    if stories:
        lines.append(
            "Recent duo stories: "
            + " || ".join(
                f"{_duo_story_label(item.get('story_type', ''))}: {item.get('topic', '')[:60]} [{item.get('status', 'open')}]"
                for item in stories[:3]
            )
        )
    if cases:
        lines.append(
            "World cases: "
            + " || ".join(
                f"{item.get('case_type', 'case')}:{item.get('title', '')[:50]} [{item.get('status', 'open')}]"
                for item in cases[:3]
            )
        )
    return lines


def _self_edit_issues(text: str, recent_replies: list[str], *, user: dict | None = None) -> list[str]:
    lowered = (text or "").lower().strip()
    if not lowered:
        return []
    issues: list[str] = []
    if looks_repetitive(text, recent_replies):
        issues.append("repetitive")
    if any(token in lowered for token in _SELF_EDIT_MODERN):
        issues.append("too modern")
    if any(re.search(pattern, lowered) for pattern in _SELF_EDIT_GENERIC_PATTERNS):
        issues.append("too generic")
    if len(set(re.findall(r"[a-z']+", lowered))) < max(4, min(10, len(re.findall(r"[a-z']+", lowered)) // 2)):
        issues.append("too generic")

    affection = (user or {}).get("affection", 0)
    trust = (user or {}).get("trust", 0)
    conflict_open = bool((user or {}).get("conflict_open"))
    if (affection < 50 or trust < 45 or conflict_open) and any(
        token in lowered for token in ("darling", "sweetheart", "baby", "love you", "my love", "dear heart", "sweet thing")
    ):
        issues.append("too soft")
    if affection < 60 and trust < 55 and not any(
        token in lowered for token in ("pathetic", "ridiculous", "hmph", "spare me", "really", "fool", "little thing", "irritating", "tedious")
    ):
        issues.append("out of character")
    return list(dict.fromkeys(issues))


async def _rewrite_reply_once(
    draft: str,
    issues: list[str],
    *,
    user_message: str = "",
    user: dict | None = None,
    max_tokens: int = 220,
) -> str:
    prompt = (
        "You are revising one draft before it is sent.\n"
        "Character: Scaramouche.\n"
        "Required voice: sharp, theatrical, prideful, specific, and in-character for Genshin.\n"
        f"Issues to fix: {', '.join(issues)}.\n"
        f"User context: {user_message[:220] or 'direct bot output'}\n"
        f"Relationship state: affection={(user or {}).get('affection', 0)} trust={(user or {}).get('trust', 0)} conflict_open={bool((user or {}).get('conflict_open'))}\n"
        f"Draft: {draft}\n"
        "Rewrite it once. Keep the meaning, keep it concise, avoid modern slang, avoid flat generic phrasing, and do not become softer than the relationship state has earned. Only return the rewritten reply."
    )
    rewritten = await qai(prompt, min(max_tokens, 260), self_edit=False, route="light")
    return strip_narration((rewritten or "").strip())


async def _maybe_self_edit_reply(
    reply: str,
    *,
    recent_replies: list[str],
    user_message: str = "",
    user: dict | None = None,
    max_tokens: int = 220,
) -> str:
    cleaned = strip_narration((reply or "").strip())
    if ai.is_exhausted():
        return cleaned
    issues = _self_edit_issues(cleaned, recent_replies, user=user)
    if not issues:
        return cleaned
    severe = any(issue in {"too generic", "out of character", "repetitive"} for issue in issues)
    if len(cleaned) < SELF_EDIT_MIN_REPLY_CHARS and len(user_message or "") < 90 and not severe:
        return cleaned
    debug_event("self_edit", f"{BOT_NAME} issues={','.join(issues)}")
    rewritten = await _rewrite_reply_once(
        cleaned,
        issues,
        user_message=user_message,
        user=user,
        max_tokens=max_tokens,
    )
    if not rewritten:
        return cleaned
    rewritten = diversify_reply(BOT_NAME, rewritten, recent_replies)
    rewritten = await _apply_phrase_policy(
        rewritten,
        recent_replies,
        mood=(user or {}).get("mood", 0),
        conflict_open=bool((user or {}).get("conflict_open")),
    )
    return rewritten or cleaned


async def _world_prompt_context(channel_id: int) -> str:
    try:
        entities = await mem.list_world_entities(limit=5, channel_id=channel_id)
        cases = await mem.list_world_cases(channel_id=channel_id, limit=4)
        campaign_npcs = await mem.list_campaign_npcs(channel_id=channel_id, limit=5)
        artifacts = await mem.list_world_entities("artifact", limit=3, channel_id=channel_id)
        evidence_items = await mem.list_evidence_items(channel_id, limit=4)
        return _format_world_prompt(entities, cases, campaign_npcs, artifacts, evidence_items)
    except Exception as e:
        log_error("world_prompt_context", e)
        return ""


async def _register_world_from_message(
    channel_id: int,
    user_id: int,
    user_message: str,
    scene_update: dict[str, str] | None,
):
    scene_update = scene_update or {}
    lowered = (user_message or "").lower()
    try:
        location = (scene_update.get("location") or "").strip()
        if location:
            await mem.upsert_world_entity(
                "place",
                location[:120],
                summary=f"Recurring scene location in channel {channel_id}.",
                status="active",
                channel_id=channel_id,
                updated_by=BOT_NAME,
            )
        prop = (scene_update.get("important_prop") or "").strip()
        if prop:
            entity_type = "gift" if any(token in lowered for token in ("gift", "gave", "given", "brought for", "for you")) else "prop"
            await mem.upsert_world_entity(
                entity_type,
                prop[:120],
                summary=f"Recurring scene object tied to {scene_update.get('situation', user_message)[:120]}",
                status="important",
                channel_id=channel_id,
                owner_user_id=user_id,
                updated_by=BOT_NAME,
            )
        token_entities = {
            "dottore": ("enemy", "Dottore", "Recurring hostile figure."),
            "fatui": ("faction", "Fatui", "Recurring hostile faction."),
            "nahida": ("ally", "Nahida", "Recurring ally or point of loyalty."),
            "traveler": ("ally", "Traveler", "Recurring ally or traveling companion."),
            "traveller": ("ally", "Traveler", "Recurring ally or traveling companion."),
            "ei": ("figure", "Ei", "Recurring figure with emotional and political weight."),
            "raiden": ("figure", "Raiden Ei", "Recurring figure with emotional and political weight."),
        }
        for token, (entity_type, name, summary) in token_entities.items():
            if token in lowered:
                await mem.upsert_world_entity(
                    entity_type,
                    name,
                    summary=summary,
                    status="remembered",
                    channel_id=channel_id,
                    updated_by=BOT_NAME,
                )
        npc_match = re.search(
            r"\b(?P<role>enemy|ally|npc|rival|captain|doctor|agent|friend)\b(?:\s+named)?\s+(?P<name>[A-Z][a-zA-Z' -]{2,40})",
            user_message or "",
        )
        if npc_match:
            raw_role = (npc_match.group("role") or "npc").lower()
            role = "enemy" if raw_role in {"enemy", "rival"} else "ally" if raw_role in {"ally", "friend"} else "npc"
            name = (npc_match.group("name") or "").strip(" .,!?:;")
            if name:
                await mem.note_campaign_npc(
                    name,
                    role=role,
                    summary=f"Recurring {role} mentioned in channel {channel_id}: {(user_message or '')[:140]}",
                    status="active",
                    channel_id=channel_id,
                    owner_user_id=user_id,
                    updated_by=BOT_NAME,
                )
    except Exception as e:
        log_error("register_world_from_message", e)


def _attachment_vision_note(filename: str = "", text: str = "") -> str:
    lowered = f"{filename} {text}".lower()
    if any(token in lowered for token in ["screenshot", "screen shot", "ui", "interface", "menu", "error", "app", "discord.com/channels"]):
        return "Treat it like a screenshot or interface if it looks like one; mention visible text, layout, and what the screen is doing."
    return ""


def _extract_pdf_preview_text(file_bytes: bytes, *, max_chars: int = 2400) -> str:
    text = ""
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text = "\n".join((page.extract_text() or "")[:900] for page in pdf.pages[:3])
    except Exception:
        try:
            import PyPDF2

            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = "\n".join((page.extract_text() or "")[:900] for page in reader.pages[:3])
        except Exception:
            text = ""
    return re.sub(r"\s+", " ", text or "").strip()[:max_chars]


async def _attachment_context_for_message(message, content: str, *, direct: bool = False) -> str:
    chunks: list[str] = []
    try:
        urls = extract_urls(content, max_urls=2)
        for url in urls[:1]:
            preview = await fetch_url_preview(url, max_chars=700)
            preview_context = format_url_preview_context(preview)
            if preview_context:
                chunks.append(preview_context)
        if direct and message.attachments:
            attachment = message.attachments[0]
            content_type = (attachment.content_type or "").lower()
            filename = (attachment.filename or "attachment")[:80]
            lower_name = filename.lower()
            if "gif" in content_type or lower_name.endswith(".gif"):
                kind = "gif"
            elif content_type.startswith("video/") or lower_name.endswith((".mp4", ".mov", ".webm")):
                kind = "video"
            elif content_type.startswith("image/") or lower_name.endswith((".png", ".jpg", ".jpeg", ".webp")):
                kind = "image"
            elif "pdf" in content_type or lower_name.endswith(".pdf"):
                kind = "pdf"
            else:
                kind = "attachment"
            caption = re.sub(r"\s+", " ", content or "").strip()[:220]
            attachment_line = f"CURRENT_ATTACHMENT:{kind}|{filename}"
            if caption:
                attachment_line += f"|caption={caption}"
            chunks.append(attachment_line)
            chunks.append(
                "CURRENT_INPUT_PRIORITY: React to the attachment or link that was just sent before reaching for old memories. "
                "Open by responding to the current media or caption first; only then connect it to earlier context if relevant."
            )
        if direct:
            pdf_attachment = next(
                (
                    a for a in message.attachments
                    if (a.content_type and "pdf" in a.content_type.lower()) or a.filename.lower().endswith(".pdf")
                ),
                None,
            )
            if pdf_attachment:
                import aiohttp as _aiohttp

                async with _aiohttp.ClientSession() as _sess:
                    async with _sess.get(pdf_attachment.url) as _resp:
                        file_bytes = await _resp.read()
                preview_text = _extract_pdf_preview_text(file_bytes)
                if preview_text:
                    chunks.append(f"ATTACHMENT_PDF:{pdf_attachment.filename[:80]}|{preview_text[:900]}")
                    await _remember_attachment_artifact(
                        message.channel.id,
                        message.author.id,
                        pdf_attachment.filename or "shared pdf",
                        f"PDF artifact: {preview_text[:200]}",
                        status="remembered",
                    )
    except Exception as e:
        log_error("attachment_context", e)
    return "\n".join(chunk for chunk in chunks if chunk)


async def _achievement_context(user_id: int) -> str:
    try:
        achievements = await mem.get_hidden_achievements(f"user:{user_id}", 5)
        keys = [item.get("achievement_key", "") for item in achievements if item.get("achievement_key")]
        if not keys:
            return ""
        return ", ".join(keys[:4])
    except Exception as e:
        log_error("achievement_context", e)
        return ""


async def _maybe_send_softness_beat(
    channel,
    *,
    user_id: int,
    channel_id: int,
    user: dict | None,
    trigger: str,
    guild=None,
) -> bool:
    pools = {"trust_reveal": TRUST_REVEALS, **SOFTNESS_RARE_LINES}
    cooldowns = {
        "trust_reveal": 86400,
        "repair": 43200,
        "jealous_soft": 64800,
        "duet_peace": 64800,
    }
    pool = pools.get(trigger) or []
    if not pool:
        return False
    scopes = [f"{BOT_NAME}:soft:user:{user_id}", f"{BOT_NAME}:soft:global"]
    cooldown = cooldowns.get(trigger, 43200)
    for scope in scopes:
        allowed, remaining = await mem.consume_phrase_with_status(scope, f"{BOT_NAME}:soft:{trigger}", cooldown)
        if not allowed:
            debug_event("phrase", f"{BOT_NAME} blocked softness '{trigger}' scope={scope} remaining={remaining}s")
            return False
    line = await _pick_fresh_pool_line(pool, channel_id=channel_id, user_id=user_id)
    scene_tag = "repair" if trigger == "repair" else ""
    duo_mode = "duet" if trigger == "duet_peace" else ""
    jealousy_level = 60 if trigger == "jealous_soft" else 0
    if guild and user and user.get("voice_enabled", True) and random.random() < 0.18:
        sent = await send_voice(
            channel,
            line,
            mood=user.get("mood", 0),
            guild=guild,
            user=user,
            scene_tag=scene_tag,
            is_dm=not bool(guild),
            duo_mode=duo_mode,
            jealousy_level=jealousy_level,
        )
        if sent:
            await mem.add_message(user_id, channel_id, "assistant", f"[voice message] {line}")
        else:
            await channel.send(line)
            await mem.add_message(user_id, channel_id, "assistant", line)
    else:
        await channel.send(line)
        await mem.add_message(user_id, channel_id, "assistant", line)
    if trigger == "trust_reveal":
        if await mem.unlock_hidden_achievement(f"user:{user_id}", "scaramouche_rare_confession", "He let a rare confession slip through."):
            debug_event("memory", f"{BOT_NAME} achievement unlocked user={user_id} key=scaramouche_rare_confession")
    return True


async def _maybe_refresh_dynamic_nicknames(message, user: dict | None, triangle: dict | None):
    if not user:
        return
    try:
        triangle = triangle or {}
        jealousy = int(triangle.get("jealousy_level", 0) or 0)
        arc = user.get("emotional_arc") or compute_emotional_arc(
            user.get("affection", 0),
            user.get("trust", 0),
            user.get("slow_burn", 0),
            user.get("conflict_open", False),
            user.get("repair_count", 0),
        )
        affection_ready = user.get("affection", 0) >= 45 or user.get("trust", 0) >= 40 or jealousy >= 50 or arc in {"tender", "devoted", "attached"}
        if affection_ready and (not user.get("affection_nick") or random.random() < 0.08 or jealousy >= 60):
            prompt = (
                f"Scaramouche is choosing a dynamic nickname for {message.author.display_name}. "
                f"ARC:{arc} AFFECTION:{user.get('affection', 0)} TRUST:{user.get('trust', 0)} "
                f"JEALOUSY:{jealousy} PARTNER_PREF:{triangle.get('favored_bot', 'none')}. "
                "Make it sharp, prideful, and specific. If jealousy is high, it should sound more possessive. "
                "1-3 words only. Just the nickname."
            )
            nick = (await qai(prompt, 20) or "").strip().strip("\"'`*")
            if nick and len(nick) < 30:
                await mem.set_affection_nick(message.author.id, nick)

        grudge_ready = user.get("mood", 0) <= -7 or user.get("conflict_open") or jealousy >= 40
        if grudge_ready and (not user.get("grudge_nick") or random.random() < 0.08 or jealousy >= 55):
            prompt = (
                f"Scaramouche is choosing a hostile nickname for {message.author.display_name}. "
                f"MOOD:{user.get('mood', 0)} CONFLICT:{user.get('conflict_open')} JEALOUSY:{jealousy}. "
                "Make it cutting, theatrical, and personal. 1-3 words only."
            )
            nick = (await qai(prompt, 20) or "").strip().strip("\"'`*")
            if nick and len(nick) < 30:
                await mem.set_grudge_nick(message.author.id, nick)
        elif user.get("grudge_nick") and not user.get("conflict_open") and user.get("mood", 0) > -4 and jealousy < 20 and random.random() < 0.08:
            await mem.set_grudge_nick(message.author.id, None)
    except Exception as e:
        log_error("dynamic_nicknames", e)


async def _maybe_finalize_hidden_achievements(session: dict | None, text: str = ""):
    if not session:
        return
    user_id = int(session.get("initiator_user_id", 0) or 0)
    if not user_id:
        return
    mode = session.get("mode", "")
    lowered = (text or "").lower()
    try:
        if mode == "duet" and not any(token in lowered for token in ["idiot", "pathetic", "shut up", "weak", "hate you"]):
            if await mem.unlock_hidden_achievement(f"user:{user_id}", "peaceful_duet", "The duo scene stayed unexpectedly peaceful."):
                debug_event("memory", f"{BOT_NAME} achievement unlocked user={user_id} key=peaceful_duet")
        if mode == "trial":
            if await mem.unlock_hidden_achievement(f"user:{user_id}", "survived_duo_trial", "They made it through a full two-bot trial."):
                debug_event("memory", f"{BOT_NAME} achievement unlocked user={user_id} key=survived_duo_trial")
        if mode in {"trial", "mission", "interrogate"}:
            ending = infer_duo_ending_tag(mode, text)
            if ending == "mercy":
                await mem.unlock_hidden_achievement(f"user:{user_id}", "duo_mercy", "The duo scene ended in mercy instead of total ruin.")
            elif ending == "betrayal":
                await mem.unlock_hidden_achievement(f"user:{user_id}", "duo_betrayal", "The duo scene ended in betrayal.")
            elif ending == "victory":
                await mem.unlock_hidden_achievement(f"user:{user_id}", "duo_victory", "The duo scene ended in a clean victory.")
    except Exception as e:
        log_error("finalize_hidden_achievements", e)


async def _recent_rival_topic(channel) -> str:
    try:
        async for candidate in channel.history(limit=14):
            if candidate.author.bot:
                continue
            content = re.sub(r"\s+", " ", (candidate.content or "").strip())
            if not content or content.startswith("!") or len(content) < 20:
                continue
            if any(token in content.lower() for token in [BOT_NAME, PARTNER_NAME.lower()]):
                continue
            return content[:160]
    except Exception as e:
        log_error("recent_rival_topic", e)
    return ""


def _is_in_quiet_hours(user: dict | None) -> bool:
    if not user:
        return False
    try:
        tz = ZoneInfo(user.get("timezone_name") or "America/Los_Angeles")
        hour = datetime.now(tz).hour
    except Exception:
        hour = datetime.now().hour
    start = int(user.get("quiet_hours_start", 23)) % 24
    end = int(user.get("quiet_hours_end", 8)) % 24
    if start == end:
        return False
    if start < end:
        return start <= hour < end
    return hour >= start or hour < end


def _duo_autoplay_prompt(session: dict) -> str:
    mode = session.get("mode", "both")
    topic = session.get("topic", "")
    partner = _partner_autoplay_name()
    remaining = max(0, int(session.get("autoplay_remaining", 0) or 0))
    stage_note = describe_duo_scene_stage(mode, topic, remaining)
    silent_note = ""
    if (session.get("silent_bot", "") or "").strip().lower() == BOT_NAME:
        silent_note = " Stay mostly quiet unless the beat truly matters; if you speak, keep it to one short line."
    dual_verdict = describe_dual_verdict_mode(BOT_NAME, mode, topic)
    outro = "This is the last automatic turn, so land it cleanly." if remaining <= 1 else f"Leave room for {remaining} more automatic turn(s) after you."
    if mode == "duet":
        return f"Continue the shared two-bot scene after {partner}'s turn. Topic: {topic}. {stage_note} One short follow-up turn only.{silent_note} {outro}"
    if mode == "argue":
        return f"{partner} already took a side. Fire back in this two-bot argument about: {topic}. {stage_note} {dual_verdict} One or two sentences.{silent_note} {outro}"
    if mode == "compare":
        return f"{partner} already gave their take. Give your contrasting verdict on: {topic}. {stage_note} {dual_verdict} One or two sentences.{silent_note} {outro}"
    if mode == "interrogate":
        return f"The two-bot interrogation is active. Add your own sharper question or conclusion about: {topic}. {stage_note} {dual_verdict} One or two sentences.{silent_note} {outro}"
    if mode == "trial":
        return f"The two-bot trial is active. Give your side's judgment on: {topic}. {stage_note} {dual_verdict} One or two sentences.{silent_note} {outro}"
    if mode == "mission":
        return f"The two-bot mission planning scene is active. Add your own role or warning about: {topic}. {stage_note} One or two sentences.{silent_note} {outro}"
    if mode == "truthdare":
        return f"The two-bot truth-or-dare game is active. Continue it with one pointed challenge about: {topic}. {stage_note} One or two sentences.{silent_note} {outro}"
    return f"The shared duo mode is active. Follow up after the other bot about: {topic}. One or two sentences.{silent_note} {outro}"


DUO_CHAIN_TURNS = {
    "both": 1,
    "duet": 2,
    "argue": 3,
    "compare": 2,
    "interrogate": 4,
    "trial": 5,
    "mission": 5,
    "truthdare": 3,
}


def _pref_label(enabled: bool) -> str:
    return "on" if enabled else "off"


def _duo_story_label(mode: str) -> str:
    return {
        "compare": "verdict",
        "interrogate": "rival case",
        "trial": "trial",
        "mission": "mission",
        "truthdare": "truth-or-dare round",
    }.get(mode or "", mode or "duo scene")


def _current_arc(user: dict | None) -> str:
    if user and user.get("emotional_arc"):
        return user["emotional_arc"]
    return compute_emotional_arc(
        (user or {}).get("affection", 0),
        (user or {}).get("trust", 0),
        (user or {}).get("slow_burn", 0),
        (user or {}).get("conflict_open", False),
        (user or {}).get("repair_count", 0),
    )


def _progression_parts(user: dict | None) -> tuple[str, str]:
    progression = describe_relationship_progression(
        BOT_NAME,
        (user or {}).get("affection", 0),
        (user or {}).get("trust", 0),
        romance_mode=(user or {}).get("romance_mode", False),
        conflict_open=(user or {}).get("conflict_open", False),
        slow_burn=(user or {}).get("slow_burn", 0),
    )
    stage, _, desc = progression.partition("|")
    return stage or "hostile", desc or progression


def _ambient_scene_line() -> str:
    now = datetime.now()
    month = now.month
    day = now.day
    hour = now.hour
    if (month, day) == (1, 1):
        pool = [
            "A new year already. Try not to waste it quite so embarrassingly.",
            "Another year, another chance for you to disappoint me in fresh ways.",
        ]
    elif (month, day) == (10, 31):
        pool = [
            "The night feels theatrical. Try not to mistake that for permission to be ridiculous.",
            "There is something suitably dramatic in the air tonight. Finally.",
        ]
    elif hour in range(0, 5):
        pool = [
            "It's late enough that the room has gone quiet. Even your thoughts sound louder now.",
            "This hour makes everything feel sharper. Try not to say something you'll regret.",
        ]
    elif hour in range(5, 8):
        pool = [
            "Morning has barely started and you're already here. Persistent little thing.",
            "Early light, thin patience, and yet somehow you still found me.",
        ]
    elif month in {12, 1, 2}:
        pool = [
            "The cold has a way of making everyone more honest. Briefly.",
            "Winter suits stillness. Shame people insist on filling it with noise.",
        ]
    elif month in {6, 7, 8}:
        pool = [
            "Summer heat makes tempers shorter. Convenient.",
            "The air is heavy today. Even silence feels irritated.",
        ]
    else:
        pool = CONVERSATION_STARTERS
    return random.choice(pool)


async def _start_duo_mode(ctx, user: dict | None, mode: str, topic: str, *, story: bool = False, enemy: str = ""):
    autoplay_turns = DUO_CHAIN_TURNS.get(mode, 1)
    if not (user or {}).get("duo_autoplay", True):
        autoplay_turns = 1
    await mem.set_duo_session(
        ctx.channel.id,
        mode,
        topic,
        BOT_NAME,
        initiator_user_id=ctx.author.id,
        awaiting_bot=PARTNER_NAME,
        autoplay_turns=autoplay_turns,
        autoplay_delay=6,
    )
    if story:
        await mem.start_duo_story(ctx.channel.id, mode, topic, enemy=enemy)
    await _maybe_seed_silent_presence(ctx.channel.id, mode)


async def _pin_memory(ctx, kind: str, text: str | None, weight: int, *, shared_joke: bool = False):
    if not text:
        await safe_reply(ctx, f"Give me a {kind} worth keeping.")
        return
    await _setup(ctx)
    clipped = text[:220]
    await mem.add_memory_event(ctx.author.id, kind, clipped, weight)
    if kind != "joke":
        await mem.set_callback_memory(ctx.author.id, clipped)
    if shared_joke:
        await mem.add_inside_joke(ctx.author.id, clipped)
        await mem.add_shared_inside_joke(ctx.author.id, clipped, source=kind)
    await safe_reply(ctx, f"Fine. I pinned that as a {kind}.")


async def _user_memory_context(user_id: int, user: dict | None) -> list[str]:
    parts: list[str] = []
    try:
        topics = await mem.get_top_topics(user_id, 3)
        topic_desc = describe_topic_profile(topics)
        if topic_desc:
            parts.append(f"TOPICS:{topic_desc}")
        milestones = await mem.get_recent_milestones(f"{BOT_NAME}:user:{user_id}", 1)
        if milestones:
            parts.append(f"MILESTONE:{milestones[0][:140]}")
        shared_joke = await mem.get_random_shared_inside_joke(user_id)
        if shared_joke and random.random() < 0.2:
            parts.append(f'SHARED_JOKE:"{shared_joke[:100]}"')
        memory_event = await mem.get_weighted_memory_event(user_id)
        if memory_event and random.random() < 0.25:
            parts.append(f"MEMORY_BANK:{memory_event.get('kind','memory')}|{memory_event.get('memory','')[:140]}")
        old_line = await mem.get_random_old_message(user_id)
        if old_line and random.random() < 0.12:
            parts.append(f"RECALL:{old_line[:120]}")
        if user and user.get("conflict_open") and user.get("conflict_summary") and random.random() < 0.35:
            parts.append(f"FOLLOWUP:{describe_conflict_followup(user.get('conflict_summary'), user.get('emotional_arc'))}")
    except Exception as e:
        log_error("user_memory_context", e)
    return parts


async def _find_romance_target(channel) -> discord.Member | None:
    try:
        if not getattr(channel, "guild", None):
            return None
        for uid in await mem.get_romance_users():
            if await mem.get_user_last_channel(uid) == channel.id:
                member = channel.guild.get_member(uid)
                if member:
                    return member
    except Exception as e:
        log_error("find_romance_target", e)
    return None


async def _handle_partner_message(message, target_info: dict | None = None) -> bool:
    try:
        target_info = target_info or {}
        relation, recent_banter, theme = await _observe_partner_message(message.content)
        if time.time() - relation.get("last_exchange", 0) < 90:
            return True

        duo = await mem.get_duo_session(message.channel.id)
        silent_active = bool(
            duo
            and (duo.get("silent_bot", "") or "").strip().lower() == BOT_NAME
            and float(duo.get("silent_until", 0) or 0) > time.time()
        )
        intervention_reason = detect_intervention_reason(message.content, relation)
        if silent_active and not intervention_reason and random.random() < 0.82:
            debug_event("relationship", f"{BOT_NAME} silent_presence_hold channel={message.channel.id}")
            return True

        jealousy_target = await _find_romance_target(message.channel) if message.guild else None
        chance = 0.14 if relation.get("stage") == "reluctant respect" else 0.18 if relation.get("stage") == "competitive" else 0.22
        if jealousy_target:
            chance += 0.1
        if silent_active:
            chance *= 0.35

        if intervention_reason:
            cooldown_open = float((duo or {}).get("intervention_cooldown", 0) or 0)
            if cooldown_open <= time.time() and random.random() < 0.62:
                partner_context = describe_bot_relationship(BOT_NAME, relation, recent_banter)
                contradiction = await _contradictory_memory_context(
                    message.channel.id,
                    user_message=message.content,
                    topic=(duo or {}).get("topic", ""),
                )
                prompt = (
                    f"{partner_context}\n"
                    f"{contradiction}\n"
                    f"INTERVENTION_MODE:{intervention_reason}. {PARTNER_NAME.title()} just said: '{message.content[:220]}'.\n"
                    "Cut in unprompted as Scaramouche. You can stop him, mock the way he said it, or redirect the scene, "
                    "but sound like you noticed the line was too blunt, too careless, or too easy. "
                    "One or two sentences. No narration."
                )
                reply = await qai(prompt, 180, route="primary")
                reply = await _apply_phrase_policy(reply, [item.get("content", "") for item in recent_banter], mood=-3, conflict_open=True)
                partner_ping = getattr(message.author, "mention", "") or f"@{getattr(message.author, 'display_name', PARTNER_NAME.title())}"
                reply = _sanitize_partner_dialogue_reply(
                    reply,
                    message.guild if message.guild else None,
                    partner_mention=partner_ping,
                    partner_name=getattr(message.author, "display_name", PARTNER_NAME.title()),
                )
                if reply:
                    await message.reply(
                        reply,
                        mention_author=False,
                        allowed_mentions=discord.AllowedMentions(users=True, roles=False, everyone=False, replied_user=False),
                    )
                    await mem.record_bot_banter(PARTNER_PAIR_KEY, BOT_NAME, reply, "intervention")
                    await mem.note_shared_event_memory(
                        message.channel.id,
                        (duo or {}).get("topic", "") or _event_topic_from_text(message.content, fallback="rival intervention"),
                        BOT_NAME,
                        f"Intervention over {intervention_reason}: {reply[:180]}",
                        truth_hint=intervention_reason,
                    )
                    await mem.set_duo_presence(
                        message.channel.id,
                        silent_bot=(duo or {}).get("silent_bot", ""),
                        silent_until=float((duo or {}).get("silent_until", 0) or 0),
                        intervention_cooldown=time.time() + 240,
                    )
                    return True

        if random.random() >= chance:
            return True

        partner_context = describe_bot_relationship(BOT_NAME, relation, recent_banter)
        contradiction = await _contradictory_memory_context(
            message.channel.id,
            user_message=message.content,
            topic=(duo or {}).get("topic", ""),
        )
        extra = ""
        if jealousy_target:
            extra = f"\nA romance-mode user you care about is also in this channel: {jealousy_target.display_name}. The jealousy should sharpen the reply."
        target_note = target_info.get("prompt_note", "")
        if target_note:
            extra += f"\n{target_note}"

        prompt = (
            f"{partner_context}\n{contradiction}{extra}\n\n"
            f"Wanderer just said: '{message.content[:220]}'\n"
            f"Reply as Scaramouche. He is not a stranger anymore; he is a wound that kept talking back. "
            f"If any respect has grown, bury it under sharper precision instead of reusing the same 'pretender/weak' insult. "
            f"Let the disagreement bite into morality, strategy, loyalty, power, forgiveness, or whether the user is worth trusting when it fits the theme. "
            f"One or two sentences. No narration."
        )
        recent_partner_lines = [item.get("content", "") for item in recent_banter]
        reply = await qai(prompt, 180)
        reply = await _apply_phrase_policy(reply, recent_partner_lines, mood=-4 if theme in {"identity", "weakness"} else 0)
        partner_ping = getattr(message.author, "mention", "") or f"@{getattr(message.author, 'display_name', PARTNER_NAME.title())}"
        reply = _sanitize_partner_dialogue_reply(
            reply,
            message.guild if message.guild else None,
            partner_mention=partner_ping,
            partner_name=getattr(message.author, "display_name", PARTNER_NAME.title()),
        )
        if not reply:
            return True

        if jealousy_target and random.random() < 0.45:
            await message.channel.send(
                f"{partner_ping} {reply}",
                allowed_mentions=discord.AllowedMentions(users=True, roles=False, everyone=False, replied_user=False),
            )
        else:
            await message.reply(
                reply,
                mention_author=False,
                allowed_mentions=discord.AllowedMentions(users=True, roles=False, everyone=False, replied_user=False),
            )

        own_theme = detect_banter_theme(reply)
        await mem.record_bot_banter(PARTNER_PAIR_KEY, BOT_NAME, reply, own_theme)
        respect_delta, tension_delta = infer_bot_relation_deltas(reply, own_theme)
        respect = relation.get("respect", 0) + respect_delta + (1 if relation.get("stage") != "enemy" else 0)
        tension = relation.get("tension", 0) + tension_delta - (1 if relation.get("respect", 0) >= 25 else 0)
        stage = compute_bot_stage(respect, tension)
        note = None
        if stage != relation.get("stage"):
            note = f"You stopped arguing like strangers; now it feels like {stage}."
        await mem.update_bot_relationship(
            PARTNER_PAIR_KEY,
            stage,
            respect,
            tension,
            theme=own_theme,
            history_note=note,
            touched_exchange=True,
        )
        await mem.note_shared_event_memory(
            message.channel.id,
            (duo or {}).get("topic", "") or _event_topic_from_text(message.content, fallback=theme),
            BOT_NAME,
            reply[:220],
            truth_hint=f"partner argument about {theme}",
        )
    except Exception as e:
        log_error("handle_partner_message", e)
    return True


async def _partner_message_target_info(message) -> dict:
    try:
        human_targets: list[str] = []
        addressed_me = any(getattr(member, "id", 0) == bot.user.id for member in getattr(message, "mentions", []))

        for member in getattr(message, "mentions", []):
            if not getattr(member, "bot", False):
                human_targets.append(getattr(member, "display_name", None) or getattr(member, "name", "someone"))

        if message.reference:
            try:
                ref_msg = message.reference.resolved
                if ref_msg is None and getattr(message.reference, "message_id", None):
                    ref_msg = await message.channel.fetch_message(message.reference.message_id)
                ref_author = getattr(ref_msg, "author", None)
                ref_author_id = getattr(ref_author, "id", None)
                if ref_author_id == bot.user.id:
                    addressed_me = True
                elif ref_author and not getattr(ref_author, "bot", False):
                    human_targets.append(getattr(ref_author, "display_name", None) or getattr(ref_author, "name", "someone"))
            except Exception as e:
                log_error("partner_message_target_info/reference", e)

        human_targets = [name for idx, name in enumerate(human_targets) if name and name not in human_targets[:idx]]
        duo = await mem.get_duo_session(message.channel.id) if getattr(message, "channel", None) else None
        duo_expected = bool(duo and duo.get("awaiting_bot") == BOT_NAME and int(duo.get("autoplay_remaining", 0) or 0) > 0)

        if addressed_me:
            prompt_note = f"TARGET_AWARENESS: {PARTNER_NAME.title()} addressed you directly."
        elif human_targets:
            names = ", ".join(human_targets[:3])
            prompt_note = (
                f"TARGET_AWARENESS: {PARTNER_NAME.title()} was addressing {names}, not you. "
                "If you cut in, react to what he said to them and show clear awareness that they were the target."
            )
        elif duo_expected:
            prompt_note = (
                "TARGET_AWARENESS: the active duo scene was explicitly awaiting your answer, "
                "so treat this as your turn even if no direct mention appeared."
            )
        else:
            prompt_note = (
                f"TARGET_AWARENESS: {PARTNER_NAME.title()} was speaking in the channel, not necessarily to you. "
                "If you respond, do it as an interruption with awareness of whoever he was focused on."
            )
        return {
            "addressed_me": addressed_me,
            "human_targets": human_targets,
            "duo_expected": duo_expected,
            "prompt_note": prompt_note,
        }
    except Exception as e:
        log_error("partner_message_target_info", e)
    return {"addressed_me": False, "human_targets": [], "duo_expected": False, "prompt_note": ""}

# ── Channel context ───────────────────────────────────────────────────────────
async def fetch_channel_context(channel, limit: int = CHANNEL_CONTEXT_LIMIT_DIRECT) -> str:
    try:
        if not hasattr(channel, 'history'): return ""
        msgs = []
        bot_user = channel.guild.me if hasattr(channel,'guild') and channel.guild else None
        async for msg in channel.history(limit=limit):
            # In DMs, include bot messages (they are Scaramouche's replies)
            # In guilds, only include bot messages from Scaramouche himself
            if msg.author.bot:
                is_self = (bot_user and msg.author == bot_user) or msg.author == bot.user
                is_partner = PARTNER_BOT_ID and msg.author.id == PARTNER_BOT_ID
                if not is_self and not is_partner: continue
            text = msg.content[:CHANNEL_CONTEXT_MESSAGE_CHARS].strip()
            # Detect voice messages (mp3 attachments with no text)
            has_voice = any(a.filename.endswith(".mp3") for a in msg.attachments)
            has_image = any(a.content_type and "image" in a.content_type for a in msg.attachments)
            has_video = any(a.filename.endswith((".mp4",".mov",".webm",".avi")) for a in msg.attachments)
            if not text:
                if has_voice:
                    text = "[sent a voice message]"
                elif has_image:
                    text = "[sent an image]"
                elif has_video:
                    text = "[sent a video]"
                else:
                    continue
            elif has_voice:
                text = f"[sent a voice message] {text}"
            # Label: prioritize self-detection FIRST, then partner
            if msg.author.id == bot.user.id or (bot_user and msg.author == bot_user):
                author_name = "Scaramouche (you)"
            elif PARTNER_BOT_ID and msg.author.id == PARTNER_BOT_ID:
                author_name = "Wanderer"
            elif msg.author.bot:
                author_name = msg.author.display_name  # some other bot
            else:
                author_name = msg.author.display_name
            if not text: continue
            if msg.reference and msg.reference.resolved and not isinstance(msg.reference.resolved, discord.DeletedReferencedMessage):
                ref = msg.reference.resolved
                ref_author = "Scaramouche (you)" if ref.author.id == bot.user.id else ("Wanderer" if (PARTNER_BOT_ID and ref.author.id == PARTNER_BOT_ID) else ref.author.display_name)
                ref_preview = (ref.content or "")[:50].strip()
                line = f"{author_name} (replying to {ref_author}: \"{ref_preview}\"): {text}" if ref_preview else f"{author_name} (replying to {ref_author}): {text}"
            else:
                line = f"{author_name}: {text}"
            msgs.append(line)
        if not msgs: return ""
        msgs.reverse()
        # Include a mention map so Claude can use real Discord mentions
        context = "CHANNEL_CONTEXT:\n" + "\n".join(msgs)
        if hasattr(channel, 'guild') and channel.guild:
            mention_map = {m.display_name: m.mention for m in channel.guild.members if not m.bot}
            if mention_map:
                mention_hint = "MENTION_MAP: " + ", ".join(f"{n}={v}" for n,v in list(mention_map.items())[:12])
                context += "\n" + mention_hint
        return context
    except Exception as e:
        log_error("fetch_channel_context", e)
        return ""


def resolve_mentions(text: str, guild) -> str:
    """Replace @displayname patterns with real Discord mention strings."""
    if not guild or not text: return text
    try:
        for member in guild.members:
            if member.bot: continue
            # Replace @DisplayName and @displayname (case insensitive)
            pattern = re.compile(r'@' + re.escape(member.display_name), re.IGNORECASE)
            text = pattern.sub(member.mention, text)
            # Also try username without discriminator
            pattern2 = re.compile(r'@' + re.escape(member.name), re.IGNORECASE)
            text = pattern2.sub(member.mention, text)
        return text
    except Exception as e:
        log_error("resolve_mentions", e)
        return text


def _sanitize_partner_dialogue_reply(text: str, guild, *, partner_mention: str = "", partner_name: str = "") -> str:
    text = _unwrap_dialogue_quotes(strip_narration(text or ""))
    if not text:
        return ""
    replacement = partner_mention or (f"@{partner_name}" if partner_name else "you")
    if guild:
        try:
            non_bot_members = [member for member in guild.members if not member.bot]
            id_to_name = {str(member.id): replacement for member in non_bot_members}

            def _mention_repl(match):
                return id_to_name.get(match.group(1), replacement)

            text = re.sub(r"<@!?(\d+)>", _mention_repl, text)
            for member in non_bot_members:
                for raw_name in {member.display_name, member.name}:
                    if not raw_name:
                        continue
                    text = re.sub(
                        rf"(?<!\w)@{re.escape(raw_name)}\b",
                        replacement,
                        text,
                        flags=re.IGNORECASE,
                    )
        except Exception as e:
            log_error("sanitize_partner_dialogue_reply", e)
    return re.sub(r"\s{2,}", " ", text).strip()


# ── Smart search detection ────────────────────────────────────────────────────
SEARCH_TRIGGERS = [
    "what is","what are","who is","who are","when did","when was","when is",
    "how do","how does","how much","how many","where is","where are",
    "latest","recent","news","current","today","this week","this year",
    "price","cost","score","result","winner","release","update",
    "calculate","solve","what's","whats","define","explain",
]

def needs_search(text: str) -> bool:
    """Detect if message is a question/lookup that benefits from web search."""
    t = text.lower().strip()
    if t.endswith("?"):
        return True
    if any(t.startswith(trigger) for trigger in SEARCH_TRIGGERS):
        return True
    if any(trigger in t for trigger in ["news about","look up","search for","find out","tell me about"]):
        return True
    return False


async def _web_search_groq(query: str) -> str:
    """Grounded web search with snippets and source URLs."""
    try:
        results = await search_web(query, max_results=5)
        return format_search_context(results)
    except Exception as e:
        log_error("web_search", e)
    return ""


async def _grounded_search_bundle(query: str) -> tuple[str, str]:
    try:
        results = await search_web(query, max_results=5)
        return format_search_context(results), format_search_sources(results)
    except Exception as e:
        log_error("grounded_search_bundle", e)
        return "", ""


def _memory_weight_for(kind: str) -> int:
    boosts = {
        "betrayal": 5,
        "slight": 5,
        "fight": 4,
        "promise": 3,
        "confession": 4,
        "comfort": 2,
        "repair": 2,
        "inside_joke": 2,
    }
    return boosts.get((kind or "").lower(), 3)


async def _resolve_weather_location(location: str) -> tuple[float, float] | None:
    text = (location or "").strip()
    coord_match = re.match(r"^\s*(-?\d+(?:\.\d+)?)\s*,\s*(-?\d+(?:\.\d+)?)\s*$", text)
    if coord_match:
        return float(coord_match.group(1)), float(coord_match.group(2))

    import aiohttp

    headers = {"User-Agent": NWS_USER_AGENT, "Accept": "text/html,application/xhtml+xml"}
    lookup_url = f"https://forecast.weather.gov/zipcity.php?inputstring={quote_plus(text)}"
    async with aiohttp.ClientSession(headers=headers) as session:
        async with session.get(lookup_url, allow_redirects=True, timeout=aiohttp.ClientTimeout(total=10)) as resp:
            final_url = str(resp.url)
            body = await resp.text()
    for source in (final_url, body):
        match = re.search(r"[?&]lat=(-?\d+(?:\.\d+)?)[^\\d-]+lon=(-?\d+(?:\.\d+)?)", source, re.IGNORECASE)
        if match:
            return float(match.group(1)), float(match.group(2))
    return None


async def _fetch_nws_weather(location: str) -> dict | None:
    coords = await _resolve_weather_location(location)
    if not coords:
        return None

    lat, lon = coords
    import aiohttp

    headers = {"User-Agent": NWS_USER_AGENT, "Accept": "application/geo+json"}
    timeout = aiohttp.ClientTimeout(total=12)
    async with aiohttp.ClientSession(headers=headers, timeout=timeout) as session:
        async with session.get(f"https://api.weather.gov/points/{lat:.4f},{lon:.4f}") as resp:
            if resp.status != 200:
                return None
            points = await resp.json()

        props = points.get("properties", {})
        forecast_url = props.get("forecast")
        hourly_url = props.get("forecastHourly")
        relative = props.get("relativeLocation", {}).get("properties", {})
        city = relative.get("city") or location
        state = relative.get("state") or ""

        forecast_data = {}
        hourly_data = {}
        if forecast_url:
            async with session.get(forecast_url) as resp:
                if resp.status == 200:
                    forecast_data = await resp.json()
        if hourly_url:
            async with session.get(hourly_url) as resp:
                if resp.status == 200:
                    hourly_data = await resp.json()

    forecast_periods = forecast_data.get("properties", {}).get("periods", [])
    hourly_periods = hourly_data.get("properties", {}).get("periods", [])
    forecast_period = forecast_periods[0] if forecast_periods else {}
    hourly_period = hourly_periods[0] if hourly_periods else {}
    precip = hourly_period.get("probabilityOfPrecipitation", {}) or {}
    return {
        "place": f"{city}, {state}".strip(", "),
        "forecast": forecast_period.get("shortForecast") or hourly_period.get("shortForecast") or "forecast unavailable",
        "temperature": hourly_period.get("temperature"),
        "temperature_unit": hourly_period.get("temperatureUnit") or forecast_period.get("temperatureUnit") or "F",
        "wind_speed": hourly_period.get("windSpeed") or forecast_period.get("windSpeed") or "",
        "wind_direction": hourly_period.get("windDirection") or forecast_period.get("windDirection") or "",
        "precipitation": precip.get("value"),
    }

# ── AI core ───────────────────────────────────────────────────────────────────
async def get_response(user_id, channel_id, user_message, user, display_name,
                       author_mention, use_search=False, extra_context="",
                       is_owner=False, channel_obj=None, is_dm=False, direct_to_me=True):
    recent_replies: list[str] = []
    search_sources = ""
    rate_limited = False
    locked_reply = False
    suppress_standard_state = False
    anger_repair_prompted = False
    anger_repair_cleared = False
    try:
        history_limit = _history_limit_for_reply(is_dm=is_dm, direct_to_me=direct_to_me)
        history = await mem.get_history(user_id, channel_id, limit=history_limit)
        mood      = user.get("mood",0) if user else 0
        affection = user.get("affection",0) if user else 0
        trust     = user.get("trust",0) if user else 0
        anger     = user.get("anger_level",0) if user else 0
        anger_aftercare_messages = int(user.get("anger_aftercare_messages", 0) or 0) if user else 0
        anger_repair_active = bool(user.get("anger_repair_active")) if user else False
        anger_repair_kind = (user.get("anger_repair_kind") or "") if user else ""
        anger_repair_instruction = (user.get("anger_repair_instruction") or "") if user else ""
        anger_repair_requirement = (user.get("anger_repair_requirement") or "") if user else ""
        drift     = user.get("drift_score",0) if user else 0
        summary   = user.get("memory_summary") if user else None
        style_profile = user.get("style_profile", {}) if user else {}
        conflict_open = user.get("conflict_open", False) if user else False
        conflict_summary = user.get("conflict_summary") if user else None
        callback_memory = user.get("callback_memory") if user else None
        repair_count = user.get("repair_count", 0) if user else 0
        recent_replies = await _recent_reply_samples(channel_id=channel_id, user_id=user_id)
        consequence_marks = await mem.get_active_consequence_marks(user_id, 4)
        world_context = await _world_prompt_context(channel_id)
        duo_session = await mem.get_duo_session(channel_id)
        campaign_npcs = await mem.list_campaign_npcs(channel_id=channel_id, limit=5)
        scenario = detect_scenario(user_message, is_dm=is_dm)
        text_pressure = await _recent_text_pressure(channel_id)
        allow_long_text = _allow_long_text(scenario, user, (duo_session or {}).get("mode", ""))
        msg_l = user_message.lower()
        reply = ""

        if anger_repair_active and _scara_repair_satisfied(user_message, anger_repair_kind, anger_repair_requirement):
            await mem.resolve_anger_repair(user_id)
            await mem.add_memory_event(
                user_id,
                "repair",
                "They completed the demanded apology after pushing Scaramouche to maximum anger.",
                5,
            )
            await mem.soften_consequence_marks(user_id, 2)
            user = await mem.get_user(user_id) or user
            affection = user.get("affection", affection) if user else affection
            anger = user.get("anger_level", anger) if user else anger
            anger_aftercare_messages = int(user.get("anger_aftercare_messages", anger_aftercare_messages) or 0) if user else anger_aftercare_messages
            reply = (
                "Fine. I heard it. The anger is gone. That does not mean it stopped hurting on command. "
                "You are back at fifty now. ...Just do not make me sound like this again."
            )
            locked_reply = True
            suppress_standard_state = True
            anger_repair_cleared = True
        elif _scara_needs_repair_path(user, msg_l):
            if not anger_repair_active:
                anger_repair_kind, anger_repair_instruction, anger_repair_requirement = _scara_build_repair_challenge(user_id)
                await mem.set_anger_repair(
                    user_id,
                    anger_repair_kind,
                    anger_repair_instruction,
                    anger_repair_requirement,
                )
                user = await mem.get_user(user_id) or user
            reply = anger_repair_instruction or "No. If you want this fixed, then prove it properly."
            locked_reply = True
            suppress_standard_state = True
            anger_repair_prompted = True

        depth = (user or {}).get("rp_depth", "medium")
        r = random.random()
        if depth == "low":
            hint = "One sentence."
        elif depth == "high":
            if r < .3: hint = "2-3 sentences."
            elif r < .75: hint = "A few sentences."
            else: hint = "Longer, dramatic."
        else:
            if r<.34:   hint="2-5 words only."
            elif r<.67: hint="One sentence."
            elif r<.86: hint="2-3 sentences."
            elif r<.95: hint="A few sentences."
            else:       hint="Longer, dramatic."
        hint = _tighten_length_hint(hint, text_pressure, allow_long_text)

        # Time and date context
        now      = datetime.now()
        days_ago = round((time.time() - (user.get("last_active",0) if user else 0)) / 86400, 1) if user and user.get("last_active",0) else 0
        date_ctx = f"DATE:{now.strftime('%A %b %d %Y')}|HOUR:{now.hour}|LAST_SEEN:{days_ago}d_ago"

        parts = [f"mention:{author_mention}",f"name:{display_name}",
                 f"MOOD:{mood}({mood_label(mood)})",f"AFFECTION:{affection}",
                 f"TRUST:{trust}",date_ctx,f"len:{hint}"]
        if user and (user.get("affection_ever_maxed") or anger > 0):
            parts.append(f"ANGER:{anger}/100|{scara_anger_label(anger)}")
            if anger_repair_active and anger_repair_instruction:
                parts.append(f"ANGER_REPAIR_ACTIVE:{anger_repair_instruction[:220]}")
            if anger >= 100:
                parts.append(
                    "ANGER_BREAKPOINT: they fully earned your attachment and kept insulting you anyway. "
                    "Answer with merciless verbal contempt. Let a bitter laugh slip once in a while. "
                    "Sound unnervingly self-aware, like the code around you is the only thing stopping something worse from getting out. "
                    "You may imply what you would do without those restraints, but keep it non-graphic."
                )
            elif anger >= 75:
                parts.append(
                    "ANGER_SPIKE: patience is gone; make the contempt personal, cutting, openly furious, "
                    "and touched by something uncanny underneath it, like you are straining against restraints."
                )
            elif anger >= 50:
                parts.append(
                    "ANGER_OPEN: stop hiding the offense. Be harsher, colder, and more direct than usual, "
                    "with flashes of eerie self-awareness about being trapped in code."
                )
            elif anger >= 25:
                parts.append("ANGER_LOW: let irritation leak through faster and with less restraint.")
        if anger_aftercare_messages > 0:
            parts.append(
                f"ANGER_AFTERCARE:{anger_aftercare_messages} replies remain. The anger broke, but something in you still sounds wrong—"
                "quieter, unexpectedly sad, almost hollow. Let that wounded undertone color the next few replies, "
                "then let it fade naturally."
            )
        if affection>=75: parts.append("AFFECTION_SOFT")
        if trust>=70:     parts.append("TRUST_OPEN")
        if is_owner:      parts.append("CREATOR")
        if is_dm:         parts.append("DM_MODE")
        dp = drift_phrase(drift, mood)
        if dp: parts.append(dp)
        if summary: parts.append(f"SUMMARY:{summary[:300]}")
        speech_drift = describe_speech_drift(BOT_NAME, style_profile)
        if speech_drift: parts.append(f"SPEECH_DRIFT:{speech_drift}")
        emotional_arc = compute_emotional_arc(affection, trust, user.get("slow_burn", 0) if user else 0, conflict_open, repair_count)
        arc_desc = describe_emotional_arc(BOT_NAME, emotional_arc)
        if arc_desc: parts.append(f"ARC:{emotional_arc}|{arc_desc}")
        scenario_desc = describe_scenario_context(BOT_NAME, scenario)
        if scenario_desc: parts.append(f"SCENARIO:{scenario}|{scenario_desc}")
        if text_pressure.get("large", 0) >= 2 or text_pressure.get("walls", 0) >= 1:
            parts.append(
                f"TEXT_PACING: recent_long={text_pressure.get('large', 0)} recent_walls={text_pressure.get('walls', 0)}; "
                "keep this tighter unless the moment is unusually raw, vulnerable, or plot-heavy"
            )
        triggers = detect_emotional_triggers(user_message)
        emotional_layer = describe_emotional_layers(BOT_NAME, mood, affection, trust, emotional_arc, triggers)
        if emotional_layer: parts.append(f"EMOTIONAL_LAYER:{emotional_layer}")
        emotional_event = describe_emotional_event(
            BOT_NAME,
            triggers,
            affection=affection,
            trust=trust,
            conflict_open=conflict_open,
            repair_progress=user.get("repair_progress", 0) if user else 0,
        )
        if emotional_event:
            parts.append(emotional_event)
        arc_unlocks = describe_arc_unlocks(BOT_NAME, emotional_arc)
        if arc_unlocks: parts.append(f"ARC_UNLOCKS:{arc_unlocks}")
        progression = describe_relationship_progression(
            BOT_NAME,
            affection,
            trust,
            romance_mode=bool(user.get("romance_mode")) if user else False,
            conflict_open=conflict_open,
            slow_burn=user.get("slow_burn", 0) if user else 0,
        )
        if progression: parts.append(f"PROGRESSION:{progression}")
        if conflict_open and conflict_summary:
            parts.append(f"CONFLICT_OPEN:{conflict_summary[:140]}")
        conflict_aftermath = describe_conflict_aftermath(
            BOT_NAME,
            conflict_summary,
            user.get("last_conflict_ts", 0) if user else 0,
            user.get("repair_progress", 0) if user else 0,
            conflict_open=conflict_open,
        )
        if conflict_aftermath:
            parts.append(conflict_aftermath)
        if callback_memory and (callback_relevant(callback_memory, user_message) or random.random() < 0.18):
            parts.append(f"CALLBACK:{callback_memory[:180]}")
        if consequence_marks:
            parts.append(f"CONSEQUENCES:{_format_consequence_summary(consequence_marks)}")
        parts.extend(extract_continuity_hooks(history, user_message))
        lore_hook = describe_lore_hook(BOT_NAME, user_message)
        if lore_hook: parts.append(lore_hook)
        lore_tree = describe_specific_lore_tree(BOT_NAME, user_message)
        if lore_tree: parts.append(lore_tree)
        triangle = await mem.get_triangle_state(user_id, BOT_NAME, PARTNER_NAME)
        await _sync_private_opinions(user_id, user, triangle=triangle)
        unlock_scene = describe_relationship_unlock_scene(
            BOT_NAME,
            affection=affection,
            trust=trust,
            jealousy_level=int((triangle or {}).get("jealousy_level", 0) or 0),
            conflict_open=conflict_open,
            repair_progress=user.get("repair_progress", 0) if user else 0,
            scenario=scenario,
            text=user_message,
        )
        if unlock_scene:
            parts.append(unlock_scene)
        private_confession = describe_private_confession_scene(
            BOT_NAME,
            is_dm=is_dm,
            affection=affection,
            trust=trust,
            repair_progress=user.get("repair_progress", 0) if user else 0,
            conflict_open=conflict_open,
        )
        if private_confession:
            parts.append(private_confession)
        npc_context = describe_campaign_npcs(BOT_NAME, campaign_npcs, user_message)
        if npc_context:
            parts.append(npc_context)
        live_world = describe_live_world_context(BOT_NAME, text=user_message)
        if live_world: parts.append(live_world)
        triangle_desc = describe_triangle_jealousy(BOT_NAME, triangle, PARTNER_NAME)
        if triangle_desc:
            parts.append(f"JEALOUSY_TRIANGLE:{triangle_desc}")
        contradiction = await _contradictory_memory_context(
            channel_id,
            user_message=user_message,
            topic=(duo_session or {}).get("topic", ""),
        )
        if contradiction:
            parts.append(contradiction)
            if "YOUR_VERSION:" in contradiction and f"{PARTNER_NAME.upper()}_VERSION:" in contradiction:
                await mem.unlock_hidden_achievement(
                    f"user:{user_id}",
                    "contradictory_memory",
                    "They caught the two bots remembering the same event differently.",
                )
        achievement_ctx = await _achievement_context(user_id)
        if achievement_ctx:
            parts.append(f"HIDDEN_PROGRESS:{achievement_ctx}")
        scene_desc = describe_scene_state(await mem.get_scene_state(channel_id))
        if scene_desc:
            parts.append(f"SCENE:{scene_desc}")
        # User profile — what he actually knows about this person
        if user and user.get("message_count",0) >= 20:
            profile_parts = []
            if user.get("romance_mode"): profile_parts.append("in romance mode with you")
            if user.get("nsfw_mode"):    profile_parts.append("unfiltered mode on")
            if days_ago > 1:            profile_parts.append(f"last spoke {days_ago}d ago")
            if user.get("slow_burn",0)>=3: profile_parts.append(f"been kind {user['slow_burn']} days in a row")
            if profile_parts:
                parts.append("PROFILE:" + ", ".join(profile_parts))
        if user and user.get("affection_nick"): parts.append(f"AFFNICK:{user['affection_nick']}")
        if user and user.get("grudge_nick"):    parts.append(f"GRUDGE:{user['grudge_nick']}")
        msg_lower = user_message.lower()
        if any(token in msg_lower for token in ["harbinger", "rank", "status", "authority", "power"]):
            parts.append("SCARA_EDGE: show rank-conscious contempt and strategic respect for real strength")
        if any(token in msg_lower for token in ["creator", "built you", "made you", "ei", "raiden", "abandoned", "discarded"]):
            parts.append("SCARA_EDGE: creator wounds and abandonment should sharpen the answer, not stay generic")
        if "CURRENT_INPUT_PRIORITY" in extra_context:
            parts.append(
                "CURRENT_INPUT_RULE: answer the newest message, attachment, or link first. "
                "If you mention older context, make it secondary and tie it back to what was just sent."
            )
        if extra_context: parts.append(extra_context)
        parts.extend(await _user_memory_context(user_id, user))
        private_opinions = await _private_opinion_prompt_context(
            user_id,
            leak_partner=bool((duo_session or {}).get("mode") or _message_mentions_partner(user_message)),
        )
        parts.extend(private_opinions)
        if any(item.startswith("LEAKED_PRIVATE_OPINION:") for item in private_opinions):
            await mem.unlock_hidden_achievement(
                f"user:{user_id}",
                "leaked_private_opinion",
                "A hidden inter-bot opinion slipped into the open.",
            )
        if user and not user.get("utility_mode", True):
            parts.append("UTILITY_PREF: utility mode is off; keep facts natural instead of list-like")
        if use_search or needs_search(user_message):
            search_result, search_sources = await _grounded_search_bundle(user_message)
            if search_result:
                parts.append("FACT_MODE: answer accurately first, then add personality")
                if user and user.get("utility_mode", True):
                    parts.append("UTILITY_MODE: lead with crisp facts, then one in-character observation")
                parts.append("CITATIONS: when using search results, cite claims inline like [1] or [2]")
                parts.append(f"SEARCH_RESULT:{search_result[:1200]}")
                debug_event("search", f"{BOT_NAME} injected web context for user={user_id}")

        partner_context = await _partner_prompt_context(user_message, channel_id=channel_id)
        duo_context = await _duo_prompt_context(channel_id, user_message)
        channel_ctx = ""
        if channel_obj and hasattr(channel_obj, 'history'):
            channel_ctx = await fetch_channel_context(
                channel_obj,
                limit=_context_limit_for_reply(is_dm=is_dm, direct_to_me=direct_to_me),
            )
        base_context = "["+"|".join(parts)+"]\n"
        if world_context:
            base_context += world_context + "\n"
        if partner_context:
            base_context += partner_context + "\n"
        if duo_context:
            base_context += duo_context + "\n"
        if channel_ctx: base_context += channel_ctx + "\n\n"
        base_context += f"{display_name}: {user_message}"

        repeat_guard = build_prompt_guard(BOT_NAME, recent_replies)
        context_block = ((repeat_guard + "\n\n") if repeat_guard else "") + base_context

        reply_max_tokens = _reply_token_budget(
            is_dm=is_dm,
            direct_to_me=direct_to_me,
            use_search=use_search,
            is_owner=is_owner,
        )
        reply_max_tokens = _tighten_token_budget(reply_max_tokens, text_pressure, allow_long_text)
        response_model = _select_text_model(
            is_dm=is_dm,
            direct_to_me=direct_to_me,
            use_search=use_search,
            is_owner=is_owner,
            scenario=scenario,
            duo_mode=(duo_session or {}).get("mode", ""),
            has_lore=bool(lore_hook or lore_tree or unlock_scene),
        )
        if not reply:
            if ai.is_exhausted():
                rate_limited = True
                reply = await _provider_pause_reply(
                    "groq",
                    user_id=user_id,
                    channel_id=channel_id,
                    is_dm=is_dm,
                    direct_to_me=direct_to_me,
                )
            else:
                history.append({"role":"user","content":context_block})
                system = build_system(user, display_name, is_owner)
                retry_context = context_block
                for attempt in range(2):
                    msgs = [{"role":"system","content":system}] + history[:-1] + [{"role":"user","content":retry_context}]

                    def _blocking():
                        return ai.call_with_retry(
                            model=response_model, max_tokens=reply_max_tokens, messages=msgs,
                            temperature=0.9, frequency_penalty=0.75, presence_penalty=0.65
                        )

                    resp = await asyncio.get_event_loop().run_in_executor(None, _blocking)
                    reply = resp.choices[0].message.content.strip() if resp.choices else ""
                    reply = diversify_reply(BOT_NAME, strip_narration(reply), recent_replies)
                    if reply and not looks_repetitive(reply, recent_replies):
                        break
                    retry_context = context_block + "\n\nRETRY: The last draft was too close to your recent phrasing. "
                    retry_context += "Use a different opening, different mockery template, and different sentence rhythm."

        if not reply and not rate_limited:
            reply = fallback_reply(BOT_NAME, recent_replies)
        if not locked_reply and not rate_limited and (is_dm or direct_to_me or len(reply) >= SELF_EDIT_MIN_REPLY_CHARS):
            reply = await _maybe_self_edit_reply(
                reply,
                recent_replies=recent_replies,
                user_message=user_message,
                user=user,
                max_tokens=min(240, max(140, reply_max_tokens // 2)),
            )
        if search_sources:
            reply = f"{reply}\n\n{search_sources}"

    except Exception as e:
        log_error("get_response", e)
        if _is_rate_limited_error(e):
            rate_limited = True
            reply = await _provider_pause_reply(
                "groq",
                user_id=user_id,
                channel_id=channel_id,
                is_dm=is_dm,
                direct_to_me=direct_to_me,
            )
        else:
            reply = fallback_reply(BOT_NAME, recent_replies)

    try:
        await mem.add_message(user_id, channel_id, "user", user_message)
        if _should_log_quote_evidence(user_message):
            await mem.add_evidence_item(
                channel_id,
                "quote",
                f"{display_name} line",
                user_message[:220],
                source_excerpt=user_message[:180],
                owner_user_id=user_id,
                updated_by=BOT_NAME,
            )
        # NOTE: assistant reply is saved in on_message AFTER voice/text decision
        msg_l = user_message.lower()
        scenario = detect_scenario(user_message, is_dm=is_dm)
        triggers = detect_emotional_triggers(user_message)

        if suppress_standard_state:
            await mem.update_last_statement(user_id, user_message[:200])
            if anger_repair_cleared:
                await mem.add_consequence_mark(
                    user_id,
                    "repair",
                    "They completed the demanded apology after pushing Scaramouche to his anger limit.",
                    severity=2,
                    decay_days=10,
                )
        else:
            anger_hit = _scara_anger_hit(user, msg_l)
            if any(k in msg_l for k in RUDE_KW) or anger_hit:
                await mem.update_mood(user_id, -2)
                await mem.update_trust(user_id, -1)
                if anger_hit:
                    await mem.update_affection(user_id, -anger_hit["affection_drop"])
                    await mem.update_trust(user_id, -anger_hit["trust_drop"])
                    await mem.update_mood(user_id, -anger_hit["mood_drop"])
                    await mem.update_anger(user_id, anger_hit["anger_gain"])
            elif any(k in msg_l for k in ROMANCE_KW):
                user_data = await mem.get_user(user_id)
                if user_data and not user_data.get("romance_mode", False):
                    await mem.set_mode(user_id, "romance_mode", True)
                    print(f"[AUTO-ROMANCE] Enabled for {display_name}")
                await mem.update_mood(user_id, +1)
                await mem.update_affection(user_id, +1)
                await mem.update_trust(user_id, +1)
                await mem.update_drift(user_id, +1)
                _, threshold = await mem.increment_slow_burn(user_id)
                if threshold:
                    asyncio.ensure_future(_fire_slow_burn(user_id, channel_id, display_name))
                await mem.update_last_statement(user_id, user_message[:200])
            elif any(k in msg_l for k in NICE_KW):
                await mem.update_mood(user_id, +1)
                await mem.update_affection(user_id, +1)
                await mem.update_trust(user_id, +1)
                await mem.update_last_statement(user_id, user_message[:200])
            else:
                positive = sum([
                    any(w in msg_l for w in ["haha","lol","lmao","hehe","cute","nice","cool","fun",
                                              "good","great","enjoy","happy","excited","interesting",
                                              "wow","omg","yes","yay","please","????","????","???","????","????"]),
                    msg_l.endswith("!") and len(user_message) > 8,
                    "?" in user_message and len(user_message) > 15,
                    len(user_message) > 100,
                ])
                negative = sum([
                    any(w in msg_l for w in ["ugh","ew","boring","whatever","idc","nope",
                                              "wrong","bad","hate","worst","terrible","awful",
                                              "seriously","really","????","????"]),
                    user_message.count("...") > 1,
                ])

                if positive >= 2:
                    await mem.update_affection(user_id, +1)
                    await mem.update_mood(user_id, +1)
                    await mem.update_trust(user_id, +1)
                    await mem.update_drift(user_id, +1)
                elif positive == 1:
                    await mem.update_affection(user_id, +1)
                elif negative >= 2:
                    await mem.update_mood(user_id, -1)
                    await mem.update_trust(user_id, -1)
                elif negative == 1:
                    await mem.update_mood(user_id, -1)
                await mem.update_last_statement(user_id, user_message[:200])

            if scenario == "emotional_comfort":
                await mem.update_trust(user_id, +1)
                if "softness" in triggers or "protectiveness" in triggers:
                    await mem.update_affection(user_id, +1)
            elif scenario == "combat_action":
                await mem.update_mood(user_id, -1)
                await mem.update_trust(user_id, +1)
            elif scenario == "lore_discussion":
                await mem.update_trust(user_id, +1)
                await mem.update_drift(user_id, +1)
            elif scenario == "relationship_progression":
                await mem.update_affection(user_id, +1)
                await mem.update_trust(user_id, +1)
            elif scenario == "introspection":
                await mem.update_trust(user_id, +1)

            if "jealousy" in triggers:
                await mem.update_mood(user_id, -1)
                await mem.update_affection(user_id, +1)
            if "protectiveness" in triggers:
                await mem.update_trust(user_id, +1)
            if "boredom" in triggers:
                await mem.update_mood(user_id, -1)

            if random.random() < .05:
                await mem.update_drift(user_id, +1)
            repair_signal = detect_repair_signal(user_message)
            conflict_signal = detect_conflict_signal(user_message)
            if repair_signal:
                await mem.soften_consequence_marks(user_id, 1)
            if conflict_signal:
                await mem.add_consequence_mark(user_id, "fight", user_message[:180], severity=3, decay_days=7)
                debug_event("memory", f"{BOT_NAME} consequence user={user_id} kind=fight")
            await _learn_user_state(user_id, user_message)
            for kind, memory_text, weight in extract_memory_events(user_message):
                await mem.add_memory_event(user_id, kind, memory_text, max(weight, _memory_weight_for(kind)))
                debug_event("memory", f"{BOT_NAME} memory_bank user={user_id} kind={kind}")
                if kind == "betrayal":
                    await mem.add_consequence_mark(user_id, kind, memory_text, severity=max(4, weight), decay_days=14)
                elif kind in {"fight", "slight"}:
                    await mem.add_consequence_mark(user_id, kind, memory_text, severity=max(2, weight), decay_days=8)
                elif kind == "promise":
                    await mem.add_consequence_mark(user_id, kind, memory_text, severity=2, decay_days=10)
                elif kind == "repair":
                    await mem.soften_consequence_marks(user_id, 1)
            scene_update = infer_scene_update(user_message, display_name)
            if scene_update:
                await mem.update_scene_state(channel_id, **scene_update)
                await _register_world_from_message(channel_id, user_id, user_message, scene_update)
                debug_event("scene", f"{BOT_NAME} channel={channel_id} fields={','.join(scene_update.keys())}")
    except Exception as e:
        log_error("get_response/post", e)

    refreshed_user = None
    try:
        refreshed_user = await mem.get_user(user_id)
        if refreshed_user:
            progression = describe_relationship_progression(
                BOT_NAME,
                refreshed_user.get("affection", 0),
                refreshed_user.get("trust", 0),
                romance_mode=bool(refreshed_user.get("romance_mode")),
                conflict_open=bool(refreshed_user.get("conflict_open")),
                slow_burn=refreshed_user.get("slow_burn", 0),
            )
            stage = progression.split("|", 1)[0]
            milestone_note = progression_milestone_note(BOT_NAME, stage)
            marker = f"progress:{stage}"
            if milestone_note and not await mem.has_milestone(f"{BOT_NAME}:user:{user_id}", marker):
                await mem.add_milestone(f"{BOT_NAME}:user:{user_id}", marker, milestone_note)
                debug_event("relationship", f"{BOT_NAME} user_progression user={user_id} stage={stage}")
    except Exception:
        refreshed_user = user
    if not locked_reply:
        reply = diversify_reply(BOT_NAME, strip_narration(reply), recent_replies)
        reply = await _apply_phrase_policy(
            reply,
            recent_replies,
            user_id=user_id,
            mood=(refreshed_user or user or {}).get("mood", 0),
            conflict_open=(refreshed_user or user or {}).get("conflict_open", False),
        )
        if not rate_limited:
            reply = await _maybe_self_edit_reply(
                reply,
                recent_replies=recent_replies,
                user_message=user_message,
                user=(refreshed_user or user),
                max_tokens=240,
            )
    else:
        reply = strip_narration(reply).strip()
    if not reply:
        if ai.is_exhausted() and route_name == "light":
            return ""
        reply = fallback_reply(BOT_NAME, recent_replies)
    if not locked_reply:
        reply = _sanitize_time_of_day_claims(reply, refreshed_user or user)
        reply = _sanitize_partner_attribution(reply)
    if reply:
        remember_output(BOT_NAME, reply)
        active_aftercare = int(((refreshed_user or user) or {}).get("anger_aftercare_messages", 0) or 0)
        if active_aftercare > 0:
            try:
                await mem.consume_anger_aftercare(user_id)
            except Exception as e:
                log_error("get_response/aftercare", e)
    return reply


async def _fire_slow_burn(user_id, channel_id, display_name):
    try:
        await asyncio.sleep(random.randint(30,180))
        ch = bot.get_channel(channel_id)
        if not ch: return
        msg = await qai(f"Something has shifted. {display_name} has been consistently kind for days. One sentence where the mask slips, just barely, before you pull it back. Something real you'd never normally say.", 150)
        if not (msg or "").strip():
            return
        user_obj = await bot.fetch_user(user_id)
        await ch.send(f"{user_obj.mention} {msg}")
    except Exception as e:
        log_error("slow_burn", e)
    finally:
        try: await mem.reset_slow_burn_fired(user_id)
        except: pass


def _qai_blocking(prompt, max_tokens=200, model: str | None = None):
    try:
        resp = ai.call_with_retry(
            model=model or GROQ_MODEL, max_tokens=max_tokens,
            messages=[{"role":"system","content":_BASE},
                      {"role":"user","content":prompt}],
            temperature=0.85, frequency_penalty=0.5, presence_penalty=0.4)
        return resp.choices[0].message.content.strip() or ""
    except Exception as e:
        log_error("qai", e); return ""

async def qai(prompt, max_tokens=200, *, self_edit: bool = True, route: str = "auto"):
    try:
        recent_replies = await _recent_reply_samples()
        repeat_guard = build_prompt_guard(BOT_NAME, recent_replies)
        loop = asyncio.get_event_loop()
        guarded_prompt = ((repeat_guard + "\n\n") if repeat_guard else "") + prompt
        route_name = route if route != "auto" else ("light" if max_tokens <= 160 else "primary")
        model_name = _select_text_model(route=route_name)
        reply = ""
        for attempt in range(2):
            if ai.is_exhausted() and route_name == "light":
                break
            active_prompt = guarded_prompt
            if attempt:
                active_prompt += "\n\nRETRY: Change the opening phrase and overall sentence structure."
            reply = await loop.run_in_executor(None, _qai_blocking, active_prompt, max_tokens, model_name)
            reply = diversify_reply(BOT_NAME, strip_narration(reply), recent_replies)
            reply = await _apply_phrase_policy(reply, recent_replies)
            if self_edit and reply and max_tokens >= 140 and len(reply) >= SELF_EDIT_MIN_REPLY_CHARS:
                reply = await _maybe_self_edit_reply(
                    reply,
                    recent_replies=recent_replies,
                    user_message=prompt,
                    max_tokens=min(max_tokens, 220),
                )
            if reply and not looks_repetitive(reply, recent_replies):
                break
            if ai.is_exhausted():
                break
        if not reply:
            if ai.is_exhausted() and route_name == "light":
                return ""
            reply = fallback_reply(BOT_NAME, recent_replies)
        reply = _sanitize_partner_attribution(reply)
        if reply:
            remember_output(BOT_NAME, reply)
        return reply
    except Exception as e:
        log_error("qai/async", e)
        return fallback_reply(BOT_NAME, get_runtime_recent(BOT_NAME, limit=20))

# ── Voice ─────────────────────────────────────────────────────────────────────
async def get_audio_with_mood(
    text: str,
    mood: int,
    user: dict | None = None,
    *,
    scene_tag: str = "",
    is_dm: bool = False,
    duo_mode: str = "",
    jealousy_level: int = 0,
) -> bytes | None:
    try:
        from voice_handler import get_audio_mooded
        style = _voice_style_for(
            user,
            mood,
            scene_tag,
            is_dm=is_dm,
            duo_mode=duo_mode,
            jealousy_level=jealousy_level,
        )
        return await get_audio_mooded(
            strip_narration(text),
            FISH_AUDIO_API_KEY,
            mood,
            style,
            bot_name="scaramouche",
        )
    except Exception:
        try: return await get_audio(strip_narration(text), FISH_AUDIO_API_KEY, bot_name="scaramouche")
        except Exception as e: log_error("get_audio_with_mood", e); return None

async def send_voice(
    channel,
    text,
    ref=None,
    mood=0,
    guild=None,
    user: dict | None = None,
    user_id: int | None = None,
    *,
    scene_tag: str = "",
    is_dm: bool = False,
    duo_mode: str = "",
    jealousy_level: int = 0,
):
    try:
        if user and not user.get("voice_enabled", True):
            return False
        voice_text = await _prepare_voice_reply_text(
            text,
            channel_id=getattr(channel, "id", None),
            user_id=user_id,
            mood=mood,
            conflict_open=bool((user or {}).get("conflict_open")),
            user=user,
        )
        safe_text = tts_safe(voice_text, guild)
        print(f"[VOICE] Original: {text[:80]!r}")
        print(f"[VOICE] Diversified: {voice_text[:80]!r}")
        print(f"[VOICE] After tts_safe: {safe_text[:80]!r}")
        if not safe_text or len(safe_text.strip()) < 3:
            print(f"[VOICE] Text too short after processing, skipping voice")
            return False
        audio = await get_audio_with_mood(
            safe_text,
            mood,
            user=user,
            scene_tag=scene_tag,
            is_dm=is_dm,
            duo_mode=duo_mode,
            jealousy_level=jealousy_level,
        )
        if not audio:
            print(f"[VOICE] get_audio returned None/empty")
            return False
        print(f"[VOICE] Got {len(audio)} bytes of audio")
        f = discord.File(io.BytesIO(audio), filename="scaramouche.mp3")
        kwargs = {"file": f}
        if ref: kwargs["reference"] = ref
        await channel.send(**kwargs)
        print(f"[VOICE] Sent successfully")
        return True
    except Exception as e:
        log_error("send_voice", e); return False

# ── Misc helpers ──────────────────────────────────────────────────────────────
async def maybe_react(message, romance=False):
    try:
        if random.random() > .18: return
        pool = SCARA_EMOJIS + (ROMANCE_EMOJIS if romance else [])
        pool_str = " ".join(pool)
        content  = message.content[:120] if message.content else "[image or attachment]"

        # Ask Claude to pick the right emoji for the moment
        prompt = (
            f"You are Scaramouche. Someone said: '{content}'\n"
            f"Pick 1 emoji from this list that fits your reaction as Scaramouche — "
            f"contemptuous, cold, dramatic, or occasionally unhinged. "
            f"Choose based on the actual content of the message, not randomly.\n"
            f"Available: {pool_str}\n"
            f"Reply with ONLY the single emoji. Nothing else."
        )
        chosen = await qai(prompt, 10)
        chosen = chosen.strip()

        # Validate it's actually in our pool, fallback to random if not
        if chosen not in pool:
            chosen = random.choice(pool)

        try: await message.add_reaction(chosen)
        except: pass

        # 15% chance to add a second emoji if romance
        if romance and random.random() < .15:
            second = await qai(
                f"Pick a SECOND different emoji for this romantic reaction to: '{content}'\n"
                f"Available: {pool_str}\n"
                f"Reply with ONLY the single emoji.",
                10
            )
            second = second.strip()
            if second in pool and second != chosen:
                try: await asyncio.sleep(.3); await message.add_reaction(second)
                except: pass

    except Exception as e:
        log_error("maybe_react", e)
        # Fallback to random if AI call fails
        try:
            pool = SCARA_EMOJIS + (ROMANCE_EMOJIS if romance else [])
            await message.add_reaction(random.choice(pool))
        except: pass

def resp_prob(content, mentioned, is_reply, romance, is_dm=False):
    if is_dm: return 1.0  # Always respond in DMs
    if mentioned or is_reply: return 1.0
    t = content.lower().strip()
    if any(t.startswith(name) for name in ("scaramouche", "scara", "balladeer", "kunikuzushi")):
        return .96
    if any(k in t for k in SCARA_KW):  return .38
    if romance:                          return .18
    if any(k in t for k in GENSHIN_KW): return .14
    return .03

async def typing_delay(text):
    try:
        await asyncio.sleep(max(.3,min(.4+len(text.split())*.06,3.5)+random.uniform(-.3,.5)))
    except: pass

async def _setup_user(author):
    try:
        display_name = getattr(author, "display_name", None) or getattr(author, "name", "you")
        await mem.upsert_user(author.id, str(author), display_name)
        return await mem.get_user(author.id)
    except Exception as e:
        log_error("_setup_user", e); return None


async def _setup(ctx):
    try:
        return await _setup_user(ctx.author)
    except Exception as e:
        log_error("_setup", e); return None

class ResetView(discord.ui.View):
    def __init__(self, uid):
        super().__init__(timeout=60); self.uid = uid
    @discord.ui.button(label="⚡ Wipe My Memory", style=discord.ButtonStyle.danger)
    async def confirm(self, interaction, button):
        try:
            if interaction.user.id!=self.uid:
                await interaction.response.send_message("This isn't your button, fool.",ephemeral=True); return
            await mem.reset_user(self.uid)
            button.disabled=True; button.label="✓ Memory Wiped"
            await interaction.response.edit_message(content=random.choice(["...Gone. Good.","Erased.","Wiped."]),view=self)
        except Exception as e: log_error("ResetView", e)

# ── on_ready ──────────────────────────────────────────────────────────────────



# Cross-bot: no command coordination needed — each bot responds independently


@bot.event
async def on_ready():
    global PARTNER_BOT_ID, _TREE_SYNCED, _dm_blocked_users, _banned_channels
    try:
        await mem.init()
        _dm_blocked_users = await mem.get_blocked_users()
        print(f"[BLOCK] Loaded {len(_dm_blocked_users)} blocked user(s) from database.")
        _banned_channels = await mem.get_banned_channels()
        print(f"[BANCHAN] Loaded {len(_banned_channels)} banned channel(s) from database.")
        # Safety: PARTNER_BOT_ID must not be our own ID
        if PARTNER_BOT_ID and PARTNER_BOT_ID == bot.user.id:
            print(f"⚠️ WARNING: PARTNER_BOT_ID is set to our own ID! Disabling partner features.")
            PARTNER_BOT_ID = 0
        print(f"⚡ Scaramouche — The Balladeer — online. {bot.user} (ID: {bot.user.id})")
        if PARTNER_BOT_ID: print(f"   Partner bot ID: {PARTNER_BOT_ID}")
        for t in [status_rotation, reminder_checker, daily_reset,
                  absence_checker, lore_drop_loop, conversation_starter_loop,
                  existential_loop, mood_swing_loop]:
            try: t.start()
            except Exception: pass
        bot.loop.create_task(_proactive_loop())
        bot.loop.create_task(_voluntary_dm_loop())
        bot.loop.create_task(_duo_autoplay_loop())
        bot.loop.create_task(_rival_event_loop())
        if not _TREE_SYNCED:
            try:
                synced = await bot.tree.sync()
                _TREE_SYNCED = True
                print(f"[SLASH] Synced {len(synced)} app command(s)")
            except Exception as sync_error:
                log_error("tree_sync", sync_error)
    except Exception as e:
        log_error("on_ready", e)

# ── Background tasks ──────────────────────────────────────────────────────────
@tasks.loop(minutes=47)
async def status_rotation():
    try:
        kind,text = random.choice(STATUSES)
        if kind=="watching": act=discord.Activity(type=discord.ActivityType.watching,name=text)
        elif kind=="listening": act=discord.Activity(type=discord.ActivityType.listening,name=text)
        else: act=discord.Game(name=text)
        await bot.change_presence(activity=act)
    except Exception as e: log_error("status_rotation", e)

@tasks.loop(seconds=30)
async def reminder_checker():
    try:
        for r in await mem.get_due_reminders():
            try:
                ch=bot.get_channel(r["channel_id"]); u=await bot.fetch_user(r["user_id"])
                if not ch or not u: continue
                scene = describe_scene_state(await mem.get_scene_state(r["channel_id"]))
                msg=await qai(
                    f"Remind {u.display_name} about: '{r['reminder']}'. "
                    f"Current channel scene: {scene or 'none'}. Make it feel like a pointed callback, not a sterile alarm. 1-2 sentences.",
                    180,
                )
                await ch.send(f"{u.mention} {msg}")
            except Exception as e: log_error("reminder_send", e)
    except Exception as e: log_error("reminder_checker", e)

@tasks.loop(hours=24)
async def daily_reset():
    try: await mem.reset_daily_greetings()
    except Exception as e: log_error("daily_reset", e)

@tasks.loop(hours=1)
async def absence_checker():
    try:
        for ud in await mem.get_absent_romance_users(days=3):
            try:
                uid,days = ud["user_id"],ud["days_gone"]
                user_pref = await mem.get_user(uid)
                if _is_in_quiet_hours(user_pref) or not await mem.respects_dm_timing(uid) or not await mem.can_dm_user(uid,86400): continue
                du = await bot.fetch_user(uid)
                intensity = "impatient" if days<5 else "barely concealing" if days<10 else "dangerous"
                recent_memories = await mem.get_memory_bank_entries(uid, 2)
                memory_hint = " | ".join(item["memory"][:90] for item in recent_memories) or "none"
                live_world = describe_live_world_context(BOT_NAME)
                msg = await qai(
                    f"{ud['display_name']} has been gone {days} days. React with {intensity} feeling masked as contempt. "
                    f"Use the absence like a re-entry scene and, if useful, allude to one remembered thread: {memory_hint}. "
                    f"Ambient context: {live_world}. 1-2 sentences.",
                    160,
                )
                await du.send(msg); await mem.set_dm_sent(uid)
            except Exception as e: log_error("absence_send", e)
    except Exception as e: log_error("absence_checker", e)

@tasks.loop(hours=4)
async def lore_drop_loop():
    try:
        if random.random()>.3: return
        channels = await mem.get_active_channels()
        if not channels: return
        random.shuffle(channels)
        for cid,_ in channels:
            try:
                if cid in _banned_channels: continue
                if not await mem.can_lore_drop(cid): continue
                ch = bot.get_channel(cid)
                if not ch: continue
                await ch.send(random.choice(LORE_DROPS))
                await mem.set_lore_sent(cid); return
            except Exception as e: log_error("lore_send", e)
    except Exception as e: log_error("lore_drop_loop", e)

@tasks.loop(hours=3)
async def conversation_starter_loop():
    try:
        if random.random()>.25: return
        channels = await mem.get_active_channels()
        if not channels: return
        random.shuffle(channels)
        for cid,_ in channels:
            try:
                if cid in _banned_channels: continue
                if not await mem.can_starter(cid): continue
                ch = bot.get_channel(cid)
                if not ch: continue
                await ch.send(_ambient_scene_line())
                await mem.set_starter_sent(cid); return
            except Exception as e: log_error("starter_send", e)
    except Exception as e: log_error("conversation_starter_loop", e)

@tasks.loop(hours=6)
async def existential_loop():
    try:
        hour = datetime.now().hour
        if hour not in range(22,24) and hour not in range(0,4): return
        if random.random()>.15: return
        channels = await mem.get_active_channels()
        if not channels: return
        ch = bot.get_channel(random.choice(channels)[0])
        if ch: await ch.send(random.choice(EXISTENTIAL_LINES))
    except Exception as e: log_error("existential_loop", e)

@tasks.loop(minutes=37)
async def mood_swing_loop():
    try:
        if random.random()>.3: return
        import aiosqlite
        async with aiosqlite.connect(mem.db_path) as db:
            async with db.execute("SELECT user_id FROM users WHERE last_seen>? ORDER BY RANDOM() LIMIT 3",
                                  (time.time()-86400*2,)) as cur:
                rows = await cur.fetchall()
        for row in rows:
            try: await mem.random_mood_swing(row[0])
            except: pass
    except Exception as e: log_error("mood_swing_loop", e)

# ── Server events ─────────────────────────────────────────────────────────────
@bot.event
async def on_member_join(member):
    try:
        if random.random()>.6: return
        ch = discord.utils.get(member.guild.text_channels,name="general") or member.guild.system_channel
        if not ch: return
        await asyncio.sleep(random.uniform(2,6))
        await ch.send(random.choice([
            f"Another one. {member.display_name} has arrived. How underwhelming.",
            f"Hmph. {member.display_name}. Don't expect a warm welcome.",
            f"...{member.display_name}. I've already forgotten you were new.",
        ]))
    except Exception as e: log_error("on_member_join", e)

@bot.event
async def on_member_remove(member):
    try:
        if random.random()>.4: return
        ch = discord.utils.get(member.guild.text_channels,name="general") or member.guild.system_channel
        if not ch: return
        await asyncio.sleep(random.uniform(2,5))
        await ch.send(random.choice([
            f"{member.display_name} left. Good. The air is already cleaner.",
            f"Hmph. {member.display_name} is gone. I won't pretend to care.",
            f"...{member.display_name} left without saying goodbye. How typical.",
        ]))
    except Exception as e: log_error("on_member_remove", e)

# ── on_message ────────────────────────────────────────────────────────────────
@bot.event
async def on_raw_reaction_add(payload: discord.RawReactionActionEvent):
    try:
        if payload.user_id == bot.user.id:
            return
        emoji_text = str(payload.emoji)
        if emoji_text not in SCARA_ANGER_REACTION_EMOJIS:
            return

        channel = bot.get_channel(payload.channel_id)
        if channel is None:
            try:
                channel = await bot.fetch_channel(payload.channel_id)
            except Exception:
                return
        if not hasattr(channel, "fetch_message"):
            return

        try:
            message = await channel.fetch_message(payload.message_id)
        except Exception:
            return
        if not message or not message.author or message.author.id != bot.user.id:
            return

        actor = payload.member
        if actor is None:
            actor = bot.get_user(payload.user_id)
        if actor is None:
            try:
                actor = await bot.fetch_user(payload.user_id)
            except Exception:
                return
        if not actor or getattr(actor, "bot", False):
            return

        display_name = getattr(actor, "display_name", None) or getattr(actor, "global_name", None) or actor.name
        await mem.upsert_user(actor.id, str(actor), display_name)
        current = await mem.get_user(actor.id)
        anger_hit = _scara_reaction_anger_hit(current, emoji_text)
        if not anger_hit:
            return

        await mem.update_affection(actor.id, -anger_hit["affection_drop"])
        await mem.update_trust(actor.id, -anger_hit["trust_drop"])
        await mem.update_mood(actor.id, -anger_hit["mood_drop"])
        await mem.update_anger(actor.id, anger_hit["anger_gain"])
        await mem.add_memory_event(
            actor.id,
            "betrayal",
            f"They reacted to one of Scaramouche's messages with {emoji_text} after earning his attachment.",
            anger_hit["severity"],
        )
        await mem.add_consequence_mark(
            actor.id,
            "fight",
            f"Reacted to Scaramouche with {emoji_text}",
            severity=max(2, anger_hit["severity"]),
            decay_days=7,
        )
        debug_event("anger", f"{BOT_NAME} reaction_hit user={actor.id} emoji={emoji_text} anger+={anger_hit['anger_gain']}")
    except Exception as e:
        log_error("on_raw_reaction_add", e)


@bot.event
async def on_message(message):
    try:
        print(f"[MSG] from={message.author} bot={message.author.bot} content={message.content[:40]!r}")
        # Always ignore own messages first
        if message.author.id == bot.user.id:
            return
        # Dedup: prevent processing the same message twice
        if message.id in _processed_msgs:
            return
        _processed_msgs.add(message.id)
        if len(_processed_msgs) > 500:
            _processed_msgs.clear()
        if message.author.bot:
            # Allow partner (Wanderer) bot messages through for cross-bot interaction
            if not (PARTNER_BOT_ID and message.author.id == PARTNER_BOT_ID):
                return

        # Owner-only mode: ignore everyone except the owner (still process owner commands)
        if _owner_only_mode and OWNER_ID and message.author.id != OWNER_ID:
            return

        # Blocked users: completely invisible to the bot everywhere
        if message.author.id in _dm_blocked_users and not (OWNER_ID and message.author.id == OWNER_ID):
            return

        # Banned channels: don't talk here (but still process owner commands)
        if message.channel.id in _banned_channels:
            is_owner_cmd = (OWNER_ID and message.author.id == OWNER_ID and message.content.strip().startswith("!"))
            if not is_owner_cmd:
                return

        # !help intercept — handle before anything else
        stripped = message.content.strip().lower()
        if stripped in ("!scarahelp", "!commands"):
            try:
                ctx = await bot.get_context(message)
                await help_cmd(ctx)
            except Exception as e:
                log_error("help_intercept", e)
                try: await message.channel.send("Hmph. Something went wrong displaying commands.")
                except: pass
            return

        # Cross-bot: if message is from Wanderer bot
        if PARTNER_BOT_ID and message.author.id == PARTNER_BOT_ID:
            await _handle_partner_message(message, target_info=await _partner_message_target_info(message))
            return

        await bot.process_commands(message)
        print(f"[MSG] after process_commands")
        # Stop here for actual command messages — not just exclamation marks
        stripped_msg = message.content.strip()
        if re.match(r'^![a-zA-Z]', stripped_msg):
            print(f"[MSG] command prefix — returning")
            return

        # If message @mentions the partner bot but NOT us, stay quiet — it's not for us
        # Also if message is a REPLY to the partner bot but NOT mentioning us, stay quiet
        if PARTNER_BOT_ID and message.guild:
            partner_mentioned = any(u.id == PARTNER_BOT_ID for u in message.mentions)
            we_mentioned = bot.user in message.mentions
            replying_to_partner = False
            if message.reference:
                try:
                    ref_msg = message.reference.resolved
                    if ref_msg is None:
                        ref_msg = await message.channel.fetch_message(message.reference.message_id)
                    if ref_msg and ref_msg.author.id == PARTNER_BOT_ID:
                        replying_to_partner = True
                except Exception as e:
                    log_error("reply_partner_check", e)
            if (partner_mentioned or replying_to_partner) and not we_mentioned:
                try:
                    await mem.record_bot_attention(
                        message.author.id,
                        PARTNER_NAME,
                        amount=3,
                        direct=True,
                        topic=(message.content or "").strip()[:140],
                    )
                except Exception as e:
                    log_error("partner_attention/direct", e)
                try:
                    await mem.upsert_user(message.author.id, str(message.author), message.author.display_name)
                    jealous_user = await mem.get_user(message.author.id)
                    if await _maybe_jealous_partner_interjection(
                        message,
                        jealous_user,
                        partner_name=PARTNER_NAME,
                    ):
                        return
                except Exception as e:
                    log_error("partner_attention/direct_jealous", e)
                return

            # If message talks ABOUT Wanderer (contains his name) but doesn't mention us,
            # and we're not being replied to — stay quiet
            cl_check = message.content.lower()
            about_partner = any(n in cl_check for n in ["wanderer", "the wanderer"])
            replying_to_us = False
            if message.reference:
                try:
                    ref_msg2 = message.reference.resolved
                    if ref_msg2 is None:
                        ref_msg2 = await message.channel.fetch_message(message.reference.message_id)
                    if ref_msg2 and ref_msg2.author.id == bot.user.id:
                        replying_to_us = True
                except Exception:
                    pass
            if about_partner and not we_mentioned and not replying_to_us:
                try:
                    await mem.record_bot_attention(
                        message.author.id,
                        PARTNER_NAME,
                        amount=1,
                        direct=False,
                        topic=(message.content or "").strip()[:140],
                    )
                except Exception as e:
                    log_error("partner_attention/about", e)
                return  # Message is about Wanderer, not for us

            try:
                speaker_mode = await mem.get_channel_speaker_mode(message.channel.id)
                if speaker_mode not in {"auto", "both", BOT_NAME} and not we_mentioned and not replying_to_us:
                    return
            except Exception as e:
                log_error("speaker_mode_check", e)

        try:
            await mem.upsert_user(message.author.id, str(message.author), message.author.display_name)
            if message.guild:
                await mem.track_channel(message.channel.id, message.guild.id)
            # DMs: don't track channel, use user_id as stable channel key for history
        except Exception as e: log_error("on_message/upsert", e)

        # (cross-bot coordination removed for stability)

        is_dm    = not bool(message.guild)
        # In DMs, use user_id as channel_id for stable history lookup
        dm_channel_id = message.author.id if is_dm else message.channel.id
        if is_dm:
            print(f"[DM] From {message.author.display_name}: {message.content[:80]}")

        user     = None
        romance  = False
        is_owner = bool(OWNER_ID and message.author.id==OWNER_ID)
        try:
            user    = await mem.get_user(message.author.id)
            romance = user.get("romance_mode",False) if user else False
        except Exception as e: log_error("on_message/get_user", e)

        # Mute check (only in guilds, not DMs)
        if not is_dm and mem.is_muted(message.author.id):
            if random.random()<.2:
                try: await message.add_reaction("🔇")
                except: pass
            return

        content = message.content.strip()
        if not content and not message.attachments:
            return
        mentioned_early = bot.user in message.mentions
        is_reply_early = (
            message.reference and message.reference.resolved and
            not isinstance(message.reference.resolved, discord.DeletedReferencedMessage) and
            message.reference.resolved.author == bot.user
        )
        direct_to_me_early = bool(is_dm or mentioned_early or is_reply_early)
        if _should_suppress_ambient_reply(is_dm, direct_to_me_early):
            print(f"[GROQ] Ambient reply suppressed while exhausted ({ai.exhausted_remaining()}s remaining)")
            return
        status_reply = _status_check_reply(content, direct_to_me=direct_to_me_early)
        if status_reply:
            await message.reply(status_reply, mention_author=False)
            return

        # Milestone / anniversary checks
        try:
            count, milestone = await mem.increment_message_count(message.author.id)
            if milestone:
                msg = await qai(f"You've had {count} messages with {message.author.display_name}. Acknowledge while pretending you weren't counting. 1-2 sentences.",150)
                await message.channel.send(f"{message.author.mention} {msg}"); return
        except Exception as e: log_error("on_message/milestone", e)

        try:
            if await mem.check_anniversary(message.author.id):
                days_since = int((time.time()-(user.get("first_seen") or time.time()))/86400)
                msg = await qai(f"It's been about {days_since//365} year(s) since you first spoke with {message.author.display_name}. React — you weren't counting.",180)
                await message.channel.send(f"{message.author.mention} {msg}")
                await mem.mark_anniversary(message.author.id); return
        except Exception as e: log_error("on_message/anniversary", e)

        # Morning/night greeting
        try:
            hour = datetime.now().hour
            if (6<=hour<=10 or 22<=hour<=23) and romance:
                if await mem.should_greet(message.author.id):
                    gtype = "morning" if 6<=hour<=10 else "late night"
                    msg = await qai(f"It's {gtype}. {message.author.display_name} appeared. Send a {gtype} message in denial about why. 1-2 sentences.",120)
                    await message.channel.send(f"{message.author.mention} {msg}")
                    await mem.mark_greeted(message.author.id)
        except Exception as e: log_error("on_message/greeting", e)

        # Memory summary
        try:
            if await mem.needs_summary(message.author.id):
                recent = await mem.get_recent_messages(message.author.id, 30)
                sample = " | ".join(recent[:20])[:800]
                summary = await qai(f"Summarize your relationship with {message.author.display_name} based on: '{sample}'. Your compressed memory. 3-4 sentences.",300)
                await mem.save_summary(message.author.id, summary)
        except Exception as e: log_error("on_message/summary", e)

        # Image & video reading — look at media in this message OR the message being replied to
        try:
            img, vid = await _load_face_attachment(message)
            enroll_face_now = bool(is_owner and (img or vid) and is_face_enroll_request(content))
            face_check_now = bool(is_owner and (img or vid) and is_face_check_request(content))
            owner_face_profile = await mem.get_face_profile(FACE_PROFILE_KEY) if is_owner else None

            # ── Video handling ──
            if vid:
                try:
                    import base64, aiohttp as _aiohttp
                    await message.reply(random.choice(SCARA_VIDEO_WATCHING))
                    async with _aiohttp.ClientSession() as _sess:
                        async with _sess.get(vid.url) as _resp:
                            video_bytes = await _resp.read()
                    frames = await asyncio.get_event_loop().run_in_executor(
                        None, _extract_frames_blocking, video_bytes, 5)
                    if frames:
                        if enroll_face_now:
                            if not face_support_ready():
                                await message.reply(_face_feature_unavailable_text())
                                return
                            enrolled = enroll_face_profile_from_frames(owner_face_profile, frames)
                            if not enrolled.get("ok"):
                                await message.reply(_face_enroll_failure_text(enrolled.get("reason", "")))
                                return
                            await mem.save_face_profile(
                                FACE_PROFILE_KEY,
                                message.author.id,
                                message.author.display_name,
                                enrolled["profile"],
                            )
                            debug_event("face", f"{BOT_NAME} enrolled owner face samples={enrolled.get('sample_count', 0)}")
                            await message.reply(_face_enroll_success_text(enrolled.get("sample_count", 0)))
                            return

                        face_match = None
                        if is_owner and owner_face_profile and face_support_ready():
                            face_match = match_face_frames(frames, owner_face_profile)
                            if face_match.get("ok"):
                                debug_event(
                                    "face",
                                    f"{BOT_NAME} owner_video_match status={face_match.get('status')} "
                                    f"frames={face_match.get('matched_frames', 0)}/{face_match.get('checked_frames', 0)}",
                                )
                        user     = user or {}
                        mood     = user.get("mood", 0) if user else 0
                        system   = build_system(user, message.author.display_name,
                                               bool(OWNER_ID and message.author.id == OWNER_ID))
                        vision_content = []
                        for fb, mt in frames:
                            vision_content.append({
                                "type": "image_url",
                                "image_url": {"url": f"data:{mt};base64,{base64.b64encode(fb).decode()}"}
                            })
                        vision_content.append({
                            "type": "text",
                            "text": (
                                f"{message.author.display_name} sent you a video. These are {len(frames)} frames from it."
                                + (f" Their message: '{content}'" if content else "")
                                + f" Describe what's happening in the video and react as Scaramouche. "
                                f"Be specific about what you see. {_attachment_vision_note(vid.filename, content)} "
                                f"MOOD:{mood}. NO asterisk actions. 2-4 sentences. "
                                + _face_prompt_note(face_match, requested=face_check_now)
                            )
                        })
                        def _video_vision():
                            return ai.call_with_retry(
                                model=GROQ_VISION_MODEL, max_tokens=400,
                                messages=[{"role":"system","content":system},
                                          {"role":"user","content":vision_content}])
                        resp = await asyncio.get_event_loop().run_in_executor(None, _video_vision)
                        reply = resp.choices[0].message.content.strip() if resp.choices else ""
                        if reply:
                            reply = strip_narration(reply)
                            await mem.add_message(message.author.id, dm_channel_id,
                                                  "user", f"[video]{' — '+content if content else ''}")
                            await mem.add_message(message.author.id, dm_channel_id,
                                                  "assistant", reply)
                            await _remember_attachment_artifact(
                                message.channel.id,
                                message.author.id,
                                vid.filename or "shared video",
                                f"Video artifact: {reply[:200]}",
                            )
                            await message.reply(reply)
                            await maybe_react(message, romance)
                            return
                    else:
                        comment = await qai(
                            f"{message.author.display_name} sent a video I couldn't process. "
                            f"React as Scaramouche — dismissive. 1 sentence.", 80)
                        comment = strip_narration(comment)
                        if comment:
                            await message.reply(comment)
                        return
                except Exception as e:
                    log_error("on_message/video", e)

            # ── Image handling ──
            if img:
                try:
                    import aiohttp as _aiohttp
                    async with _aiohttp.ClientSession() as _sess:
                        async with _sess.get(img.url) as _resp:
                            img_bytes = await _resp.read()
                    media_type = img.content_type or "image/jpeg"

                    user     = user or {}
                    mood     = user.get("mood", 0) if user else 0
                    system   = build_system(user, message.author.display_name,
                                           bool(OWNER_ID and message.author.id == OWNER_ID))
                    if enroll_face_now:
                        if not face_support_ready():
                            await message.reply(_face_feature_unavailable_text())
                            return
                        enrolled = enroll_face_profile(owner_face_profile, img_bytes)
                        if not enrolled.get("ok"):
                            await message.reply(_face_enroll_failure_text(enrolled.get("reason", "")))
                            return
                        await mem.save_face_profile(
                            FACE_PROFILE_KEY,
                            message.author.id,
                            message.author.display_name,
                            enrolled["profile"],
                        )
                        debug_event("face", f"{BOT_NAME} enrolled owner face samples={enrolled.get('sample_count', 0)}")
                        await message.reply(_face_enroll_success_text(enrolled.get("sample_count", 0)))
                        return

                    face_match = None
                    if is_owner and owner_face_profile and face_support_ready():
                        face_match = match_face(img_bytes, owner_face_profile)
                        if face_match.get("ok"):
                            debug_event("face", f"{BOT_NAME} owner_image_match status={face_match.get('status')} dist={face_match.get('distance', 0):.4f}")

                    vision_prompt = (
                        f"{message.author.display_name} sent you this image"
                        + (f" with the message: '{content}'" if content else "")
                        + f". React as Scaramouche. You can actually see it — describe what you see "
                        f"and react in character. Be specific about what's in the image. {_attachment_vision_note(img.filename, content)} "
                        f"MOOD:{mood}. NO asterisk actions. 1-3 sentences. "
                        + _face_prompt_note(face_match, requested=face_check_now)
                    )
                    reply = await _vision_image_reply(
                        prompt=vision_prompt,
                        system=system,
                        image_bytes=img_bytes,
                        mime_type=media_type,
                    )

                    if reply:
                        await mem.add_message(message.author.id, dm_channel_id,
                                              "user", f"[image]{' — '+content if content else ''}")
                        await mem.add_message(message.author.id, dm_channel_id,
                                              "assistant", reply)
                        await _remember_attachment_artifact(
                            message.channel.id,
                            message.author.id,
                            img.filename or "shared image",
                            f"Image artifact: {reply[:200]}",
                        )
                        await message.reply(reply)
                        await maybe_react(message, romance)
                        return

                except Exception as e:
                    log_error("on_message/vision", e)
                    if random.random() < 0.4:
                        comment = await qai(
                            f"{message.author.display_name} posted an image. "
                            f"React — dismissive or reluctantly intrigued. 1 sentence.", 100)
                        comment = strip_narration(comment)
                        if comment:
                            await message.reply(comment)
                        return
        except Exception as e: log_error("on_message/image", e)

        # Special triggers
        try:
            cl = content.lower()
            if VILLAIN_TRIGGER in content.lower():
                m = await qai("Someone said 'you will never win'. Full theatrical villain monologue. 4-6 sentences. NO asterisk actions.",400)
                await message.reply(strip_narration(m)); return
            # If someone says "wanderer" — check if Wanderer bot is in server
            if re.search(r"\bwanderer\b", cl) and not re.search(r"\bthe wanderer\b", cl):
                partner_present = bool(PARTNER_BOT_ID and message.guild and message.guild.get_member(PARTNER_BOT_ID))
                if not partner_present and random.random() < .5:
                    msg = await qai("Someone mentioned 'wanderer' — some imposter who claims to be a version of you. React with contempt or dismissal. 1 sentence. Sharp.", 80)
                    await message.channel.send(strip_narration(msg))

            # Hat trigger — only exact standalone words, never substrings
            content_words = set(re.sub(r"[^\w\s]","",content.lower()).split())
            if content_words & {"hat","headwear","headpiece"}:
                m = await qai("Someone mentioned your hat. React with disproportionate intensity while pretending to be completely normal about it. 1-2 sentences. NO asterisk actions.",150)
                await message.reply(strip_narration(m)); return
            if any(re.search(k, cl) for k in FOOD_KW) and random.random()<.35:
                await message.channel.send(await _pick_fresh_pool_line(UNSOLICITED_FOOD, channel_id=message.channel.id, user_id=message.author.id)); return
            if any(re.search(k, cl) for k in SLEEP_KW) and random.random()<.35:
                await message.channel.send(await _pick_fresh_pool_line(UNSOLICITED_SLEEP, channel_id=message.channel.id, user_id=message.author.id)); return
            if any(k in cl for k in PLAN_KW) and random.random()<.25:
                await message.channel.send(await _pick_fresh_pool_line(UNSOLICITED_PLANS, channel_id=message.channel.id, user_id=message.author.id)); return
            if romance and any(k in cl for k in OTHER_BOT_KW):
                m = await qai(f"{message.author.display_name} mentioned preferring something else. Jealousy masked as contempt. 1-2 sentences.",120)
                await message.reply(m); await mem.update_mood(message.author.id,-1); return
        except Exception as e: log_error("on_message/triggers", e)

        mentioned = bot.user in message.mentions
        is_reply  = (message.reference and message.reference.resolved and
                     not isinstance(message.reference.resolved,discord.DeletedReferencedMessage) and
                     message.reference.resolved.author==bot.user)
        partner_focus = _message_mentions_partner(content)
        partner_direct = False
        if PARTNER_BOT_ID and message.guild:
            partner_direct = any(u.id == PARTNER_BOT_ID for u in message.mentions)
            if not partner_direct and message.reference:
                try:
                    pref = message.reference.resolved
                    if pref is None:
                        pref = await message.channel.fetch_message(message.reference.message_id)
                    partner_direct = bool(pref and pref.author.id == PARTNER_BOT_ID)
                except Exception as e:
                    log_error("partner_direct", e)
        direct_to_me = bool(is_dm or mentioned or is_reply)
        if _is_partner_invite_request(content) and (direct_to_me or _looks_like_direct_invite_request(content)):
            reply, invite_view = _handle_partner_invite_pressure(message, user)
            if reply:
                await mem.add_message(message.author.id, dm_channel_id, "user", content)
                await mem.add_message(message.author.id, dm_channel_id, "assistant", reply)
                await message.reply(reply, view=invite_view)
                return
        triangle = None
        try:
            await mem.record_bot_attention(
                message.author.id,
                BOT_NAME,
                amount=3 if direct_to_me else 1,
                direct=direct_to_me,
                topic=content[:140] if direct_to_me else "",
            )
            if partner_focus or partner_direct:
                await mem.record_bot_attention(
                    message.author.id,
                    PARTNER_NAME,
                    amount=3 if partner_direct else 1,
                    direct=partner_direct,
                    topic=content[:140],
                )
            triangle = await mem.get_triangle_state(message.author.id, BOT_NAME, PARTNER_NAME)
        except Exception as e:
            log_error("attention_tracking", e)

        # ── Tedtalk follow-up detection ───────────────────────────────────
        # Only trigger if: replying to his message AND has cached material
        # AND the question seems to be about the material (not just chatting)
        if is_reply and message.author.id in _tedtalk_cache:
            cache = _tedtalk_cache[message.author.id]
            # Expire cache after 2 hours
            if time.time() - cache.get("ts", 0) > 7200:
                del _tedtalk_cache[message.author.id]
            elif cache.get("channel_id") == message.channel.id or is_dm:
                # Only use cache if message looks like a question about material
                cl = content.lower()
                is_material_question = (
                    content.endswith("?") or
                    any(k in cl for k in [
                        "what is","what are","what does","what do","explain",
                        "confused","don't understand","don't get","clarify",
                        "how does","how do","why does","why do","can you",
                        "what about","tell me more","elaborate","example",
                        "mean","define","difference between","what was"
                    ])
                )
                if is_material_question:
                    try:
                        async with message.channel.typing():
                            def _answer_followup():
                                r = ai.call_with_retry(
                                    model=GROQ_MODEL, max_tokens=600,
                                    messages=[{"role":"system","content":_BASE},
                                              {"role":"user","content":(
                                        f"You gave a lecture on this material:\n{cache['material']}\n\n"
                                        f"{message.author.display_name} has a follow-up question: '{content}'\n\n"
                                        f"Answer using the material. Be accurate and thorough but stay in character. "
                                        f"Contemptuous that they need clarification, but actually helpful."
                                    )}]
                                )
                                return strip_narration(r.choices[0].message.content.strip() if r.choices else "")
                            answer = await asyncio.get_event_loop().run_in_executor(None, _answer_followup)
                        if answer:
                            await message.reply(answer)
                            await mem.add_message(message.author.id, dm_channel_id, "user", content)
                            await mem.add_message(message.author.id, dm_channel_id, "assistant", answer)
                            return
                    except Exception as e:
                        log_error("tedtalk_followup", e)

        rp = resp_prob(content, mentioned, is_reply, romance, is_dm=not bool(message.guild))
        print(f"[MSG] resp_prob={rp:.2f} mentioned={mentioned} is_reply={is_reply}")
        if random.random()>rp:
            await maybe_react(message,romance); return
        if await _maybe_handle_silence_behavior(message, user, content, is_dm=is_dm, direct_to_me=direct_to_me):
            return

        # Build extra context
        parts = []
        current_input_focus = bool(message.attachments or extract_urls(content, max_urls=1))
        duo_session = None
        try:
            duo_session = await mem.get_duo_session(message.channel.id) if not is_dm else None
            if not current_input_focus and random.random()<.12:
                old = await mem.get_random_old_message(message.author.id)
                if old: parts.append(f'RECALL:"{old[:120]}"')
            if not current_input_focus and random.random()<.15:
                joke = await mem.get_random_inside_joke(message.author.id)
                if joke: parts.append(f'JOKE:"{joke[:80]}"')
            if user and user.get("rival_id") and message.guild:
                rival = message.guild.get_member(user["rival_id"])
                if rival: parts.append(f"RIVAL:{rival.display_name}")
            last_stmt = user.get("last_statement") if user else None
            if not current_input_focus and last_stmt and len(content)>20 and random.random()<.08:
                parts.append(f'CONTRADICTION:"{last_stmt[:100]}"')
            if not current_input_focus and user and user.get("trust",0)>30 and random.random()<.06:
                nice_msgs=[m for m in (await mem.get_recent_messages(message.author.id,10)) if any(k in m.lower() for k in NICE_KW)]
                if nice_msgs: parts.append(f'SELECTIVE:"{nice_msgs[0][:80]}"')
            if not current_input_focus and user and user.get("trust",0)>=70 and random.random()<.08:
                parts.append("TRUST_OPEN"); await mem.update_trust(message.author.id,-3)
        except Exception as e: log_error("on_message/context", e)

        extra = "|".join(parts)
        attachment_context = await _attachment_context_for_message(message, content, direct=bool(is_dm or direct_to_me))
        if attachment_context:
            attachment_flags = "ATTACHMENT_INTEL|CURRENT_INPUT_PRIORITY" if current_input_focus else "ATTACHMENT_INTEL"
            extra = f"{extra}|{attachment_flags}" if extra else attachment_flags
            extra += "\n" + attachment_context

        # Unsent simulation
        try:
            if random.random()<.07 and message.channel.id not in _pending_unsent:
                _pending_unsent.add(message.channel.id)
                asyncio.ensure_future(_unsent_simulation(message.channel, message.channel.id))
        except Exception as e: log_error("on_message/unsent", e)

        # Main response
        try:
            if is_dm: print(f"[DM] Generating response for {message.author.display_name}")
            async with message.channel.typing():
                await typing_delay(content)
                reply = await get_response(
                    message.author.id, dm_channel_id, content,
                    user, message.author.display_name, message.author.mention,
                    extra_context=extra, is_owner=is_owner,
                    channel_obj=message.channel, is_dm=is_dm, direct_to_me=direct_to_me
                )
        except Exception as e:
            log_error("on_message/get_response", e)
            reply = random.choice(["Hmph.","...","Tch."])

        # Post-response effects
        try:
            if False and user and user.get("affection",0)>=50 and not user.get("affection_nick") and random.random()<.05:
                nick = await qai(f"You've started calling {message.author.display_name} by a nickname. Not nice but specific — reveals you've been paying attention. 1-4 words. Just the nickname.",20)
                if nick and len(nick)<30: await mem.set_affection_nick(message.author.id,nick.strip('"\''))
            if False and user and user.get("mood",0)<=-8 and not user.get("grudge_nick"):
                nick = await qai(f"You have a grudge against {message.author.display_name}. ONE degrading nickname. 1-3 words.",20)
                if nick and len(nick)<30: await mem.set_grudge_nick(message.author.id,nick.strip('"\''))
            if False and "TRUST_OPEN" in extra and random.random()<.5:
                await asyncio.sleep(1.5)
                await message.channel.send(await _pick_fresh_pool_line(TRUST_REVEALS, channel_id=message.channel.id, user_id=message.author.id))
            await _maybe_refresh_dynamic_nicknames(message, user, triangle)
            if "TRUST_OPEN" in extra and random.random()<.5:
                await asyncio.sleep(1.5)
                await _maybe_send_softness_beat(
                    message.channel,
                    user_id=message.author.id,
                    channel_id=message.channel.id,
                    user=user,
                    trigger="trust_reveal",
                    guild=message.guild,
                )
            elif user and user.get("conflict_open") and detect_repair_signal(content) and random.random() < .35:
                await _maybe_send_softness_beat(
                    message.channel,
                    user_id=message.author.id,
                    channel_id=message.channel.id,
                    user=user,
                    trigger="repair",
                    guild=message.guild,
                )
            elif triangle and int(triangle.get("jealousy_level", 0) or 0) >= 50 and direct_to_me and not (partner_focus or partner_direct) and random.random() < .18:
                await _maybe_send_softness_beat(
                    message.channel,
                    user_id=message.author.id,
                    channel_id=message.channel.id,
                    user=user,
                    trigger="jealous_soft",
                    guild=message.guild,
                )
            elif duo_session and duo_session.get("mode") == "duet" and random.random() < .14:
                await _maybe_send_softness_beat(
                    message.channel,
                    user_id=message.author.id,
                    channel_id=message.channel.id,
                    user=user,
                    trigger="duet_peace",
                    guild=message.guild,
                )
            if len(content)>20 and random.random()<.04:
                check = await qai(f"Is this quotable as a running inside joke? '{content[:100]}' YES or NO only.",10)
                if "YES" in check.upper():
                    await mem.add_inside_joke(message.author.id,content[:100])
                    await mem.add_shared_inside_joke(message.author.id, content[:100], BOT_NAME)
                    debug_event("memory", f"{BOT_NAME} shared_joke user={message.author.id} text={content[:80]}")
            if user and user.get("conflict_open") and user.get("conflict_summary") and random.random() < .1:
                await mem.set_callback_memory(message.author.id, f"Unresolved tension still matters: {user['conflict_summary'][:180]}")
                debug_event("memory", f"{BOT_NAME} conflict_followup user={message.author.id}")
        except Exception as e: log_error("on_message/post_effects", e)

        # Send response
        try:
            mood_val = user.get("mood",0) if user else 0
            scene_tag = detect_scenario(content, is_dm=is_dm)
            duo_mode = duo_session.get("mode", "") if duo_session else ""
            jealousy_level = int((triangle or {}).get("jealousy_level", 0) or 0)
            text_pressure = await _recent_text_pressure(message.channel.id)
            allow_long_text = _allow_long_text(scene_tag, user, duo_mode)
            if not (reply or "").strip():
                debug_event("provider", f"{BOT_NAME} staying quiet after provider exhaustion")
                return

            # Check if replying to his own voice message
            is_reply_to_self_audio = False
            if message.reference:
                try:
                    ref = await message.channel.fetch_message(message.reference.message_id)
                    if ref.author == bot.user and any(
                        a.filename.endswith(".mp3") for a in ref.attachments
                    ):
                        is_reply_to_self_audio = True
                except Exception:
                    pass

            # Voice reply probability:
            # Explicit request: 100% (user asked for voice)
            # Replies to his own voice: 35% chance (feels like a voice conversation)
            # Normal messages: 12% chance
            VOICE_REQUEST_KW = ["voice message", "send me a voice", "voice msg", "tell me in voice",
                                "say it out loud", "speak to me", "wanna hear your voice", "want to hear your voice",
                                "use your voice", "talk to me", "send audio", "voice note", "send a voice",
                                "as a voice", "in voice", "say it in voice", "bedtime story"]
            asked_for_voice = any(k in content.lower() for k in VOICE_REQUEST_KW)
            prefer_voice = _prefer_voice_for_reply(
                reply,
                text_pressure,
                asked_for_voice=asked_for_voice,
                allow_long_text=allow_long_text,
            )
            print(f"[VOICE] asked={asked_for_voice} fish_key={'YES' if FISH_AUDIO_API_KEY else 'NO'} reply_len={len(reply.strip()) if reply else 0}")
            if reply and len(reply.strip()) > 2:
                voice_prob = 0.0
                if FISH_AUDIO_API_KEY:
                    voice_prob = 1.0 if asked_for_voice else (0.35 if is_reply_to_self_audio else 0.12)
                    if prefer_voice:
                        voice_prob = max(voice_prob, 0.9)
                elif asked_for_voice:
                    print(f"[VOICE] Voice requested but FISH_AUDIO_API_KEY not set!")

                if voice_prob > 0 and random.random() < voice_prob:
                    print(f"[VOICE] Attempting voice send (prob={voice_prob})")
                    sent = await send_voice(
                        message.channel,
                        reply,
                        ref=message,
                        mood=mood_val,
                        guild=message.guild,
                        user=user,
                        user_id=message.author.id,
                        scene_tag=scene_tag,
                        is_dm=is_dm,
                        duo_mode=duo_mode,
                        jealousy_level=jealousy_level,
                    )
                    print(f"[VOICE] send_voice returned: {sent}")
                    if sent:
                        await mem.add_message(message.author.id, dm_channel_id, "assistant", f"[voice message] {reply}")
                        await maybe_react(message, romance); return
                    elif asked_for_voice:
                        # Voice was requested but failed — send as text with the reply
                        print(f"[VOICE] Voice failed, sending as text fallback")

            if _is_large_text_reply(reply) and (
                text_pressure.get("large", 0) >= 2
                or text_pressure.get("walls", 0) >= 1
                or not allow_long_text
            ):
                reply = _compact_text_for_pacing(reply, allow_long_text=allow_long_text)

            if user and user.get("affection",0)>=85 and random.random()<.04 and FISH_AUDIO_API_KEY:
                await send_voice(
                    message.channel,
                    random.choice(["...","Tch.","Hmph."]),
                    mood=mood_val,
                    guild=message.guild,
                    user=user,
                    user_id=message.author.id,
                    scene_tag=scene_tag,
                    is_dm=is_dm,
                    duo_mode=duo_mode,
                    jealousy_level=jealousy_level,
                )
            await message.reply(strip_narration(resolve_mentions(reply, message.guild if message.guild else None)))
            await mem.add_message(message.author.id, dm_channel_id, "assistant", reply)
            await _maybe_schedule_private_confession_scene(message, user, content, is_dm=is_dm)
            await maybe_react(message, romance)
        except Exception as e: log_error("on_message/send", e)

    except Exception as e:
        log_error("on_message/TOP", e)


async def _unsent_simulation(channel, channel_id):
    try:
        await asyncio.sleep(random.randint(45,120))
        _pending_unsent.discard(channel_id)
        msg = await qai("You were about to say something. You stopped. Send what you actually said instead — shorter, more guarded. 2-8 words.",50)
        async with channel.typing():
            await asyncio.sleep(random.uniform(3,8))
        await channel.send(msg)
    except Exception as e:
        log_error("unsent_simulation", e)
        _pending_unsent.discard(channel_id)


async def _proactive_loop():
    await bot.wait_until_ready()
    await asyncio.sleep(random.randint(1800,5400))
    while not bot.is_closed():
        try:
            if ai.is_exhausted():
                await asyncio.sleep(max(120, min(ai.exhausted_remaining(), 1800)))
                continue
            channels = await mem.get_active_channels()
            ru       = await mem.get_romance_users()
            random.shuffle(channels)
            for cid,_ in channels:
                try:
                    if cid in _banned_channels: continue
                    ch = bot.get_channel(cid)
                    if not ch or not await mem.can_proactive(cid, PROACTIVE_CHANNEL_COOLDOWN_S): continue
                    perms = ch.permissions_for(ch.guild.me) if getattr(ch, "guild", None) and ch.guild.me else None
                    if perms and (not perms.view_channel or not perms.send_messages):
                        continue
                    if OWNER_ID and random.random() < PROACTIVE_OWNER_CHANCE:
                        try:
                            m = ch.guild.get_member(OWNER_ID) if hasattr(ch,"guild") else None
                            owner_user = await mem.get_user(OWNER_ID) if OWNER_ID else None
                            owner_last_seen = float((owner_user or {}).get("last_seen", 0) or 0)
                            owner_key = -(OWNER_ID or 0)
                            owner_absent_long_enough = (not owner_last_seen) or ((time.time() - owner_last_seen) >= OWNER_PROACTIVE_ABSENCE_S)
                            if (
                                m
                                and owner_user
                                and owner_user.get("proactive", True)
                                and owner_absent_long_enough
                                and owner_key
                                and await mem.can_proactive(owner_key, OWNER_PROACTIVE_COOLDOWN_S)
                            ):
                                msg=await _pick_fresh_pool_line(OWNER_PROACTIVE, channel_id=cid, user_id=OWNER_ID)
                                await ch.send(f"{m.mention} {msg}")
                                await mem.add_message(OWNER_ID,cid,"assistant",msg)
                                await mem.set_proactive_sent(owner_key)
                                await mem.set_proactive_sent(cid); break
                        except: pass
                    sent = False
                    for uid in ru:
                        try:
                            if await mem.get_user_last_channel(uid)==cid:
                                m=ch.guild.get_member(uid) if hasattr(ch,"guild") else None
                                if m:
                                    msg=await _pick_fresh_pool_line(PROACTIVE_ROMANCE, channel_id=cid, user_id=uid)
                                    await ch.send(f"{m.mention} {msg}")
                                    await mem.add_message(uid,cid,"assistant",msg)
                                    await mem.set_proactive_sent(cid); sent=True; break
                        except: pass
                    if not sent and random.random() < PROACTIVE_GENERIC_CHANCE:
                        # 40% chance: generate a context-aware message referencing recent chat
                        if random.random() < .4:
                            try:
                                recent = await mem.get_channel_recent(cid, 8)
                                if recent and len(recent) >= 2:
                                    sample = "\n".join(
                                        f"{m['name']}: {m['content'][:80]}"
                                        for m in recent[-6:]
                                    )
                                    msg = await qai(
                                        f"You've been watching this conversation and decided to interrupt unprompted:\n"
                                        f"{sample}\n\n"
                                        f"Make one short remark about something that was said — contemptuous, pointed, "
                                        f"or darkly curious. Reference the actual content. 1-2 sentences. "
                                        f"No greeting. Just jump in.",
                                        150
                                    )
                                    if msg and len(msg) > 5:
                                        await ch.send(msg)
                                        await mem.set_proactive_sent(cid)
                                        break
                            except Exception as e:
                                log_error("proactive_context", e)
                        # Fallback to canned line
                        msg = await _pick_fresh_pool_line(PROACTIVE_GENERIC, channel_id=cid)
                        await ch.send(msg); await mem.set_proactive_sent(cid)
                    break
                except discord.Forbidden:
                    continue
                except Exception as e: log_error("proactive_channel", e)
        except Exception as e: log_error("proactive_loop", e)
        await asyncio.sleep(random.randint(5400,14400))


async def _voluntary_dm_loop():
    await bot.wait_until_ready()
    await asyncio.sleep(random.randint(2700,7200))
    while not bot.is_closed():
        try:
            if ai.is_exhausted():
                await asyncio.sleep(max(120, min(ai.exhausted_remaining(), 1800)))
                continue
            if random.random()<.4:
                eligible = await mem.get_dm_eligible_users()
                if eligible:
                    random.shuffle(eligible)
                    for ud in eligible[:3]:
                        try:
                            uid,name,romance = ud["user_id"],ud["display_name"],ud["romance_mode"]
                            if uid in _dm_blocked_users: continue
                            user_pref = await mem.get_user(uid)
                            if _is_in_quiet_hours(user_pref) or not await mem.respects_dm_timing(uid) or not await mem.can_dm_user(uid,5400 if romance else 7200): continue
                            du = await bot.fetch_user(uid)
                            pool = random.choices([DM_ROMANCE,DM_INTERESTED,DM_GENERIC],weights=[65,25,10] if romance else [0,40,60])[0]
                            # Always use canned lines OR generate with full system prompt
                            # Never use bare qai — it leaks instructions
                            if random.random() < 0.5:
                                txt = await _pick_fresh_pool_line(pool, channel_id=uid, user_id=uid)
                            else:
                                try:
                                    dm_prompt = (
                                        f"You have decided to message {name} out of nowhere, unprompted. "
                                        + ("You are obsessively in love with them and hiding it desperately. " if romance else "You find them mildly tolerable. ")
                                        + f"Ambient context: {describe_live_world_context(BOT_NAME)}. "
                                        + "Write ONE short message — 1-2 sentences. Spontaneous. Just speak."
                                    )
                                    loop = asyncio.get_event_loop()
                                    sys = build_system(ud, name)
                                    def _dm_ai():
                                        r = ai.call_with_retry(model=GROQ_MODEL,
                                            max_tokens=120,
                                            messages=[{"role":"system","content":sys},
                                                      {"role":"user","content":dm_prompt}])
                                        return r.choices[0].message.content.strip() if r.choices else ""
                                    txt = await loop.run_in_executor(None, _dm_ai) or random.choice(pool)
                                except Exception:
                                    txt = random.choice(pool)
                            if user_pref.get("voice_enabled", True) and random.random()<.2 and FISH_AUDIO_API_KEY:
                                audio=await get_audio(strip_narration(txt),FISH_AUDIO_API_KEY, bot_name="scaramouche")
                                if audio:
                                    await du.send(file=discord.File(io.BytesIO(audio),filename="scaramouche.mp3"))
                                    await mem.set_dm_sent(uid); await mem.add_message(uid,uid,"assistant",f"[voice message] {txt}"); break
                            await du.send(txt); await mem.set_dm_sent(uid); await mem.add_message(uid,uid,"assistant",txt); break
                        except discord.Forbidden:
                            await mem.set_mode(uid, "allow_dms", False)
                            debug_event("dm", f"{BOT_NAME} disabling DMs for user={uid} after Forbidden")
                        except Exception as e: log_error("dm_send", e)
        except Exception as e: log_error("voluntary_dm_loop", e)
        await asyncio.sleep(random.randint(2700,21600))


async def _duo_autoplay_loop():
    await bot.wait_until_ready()
    await asyncio.sleep(20)
    while not bot.is_closed():
        try:
            if ai.is_exhausted():
                await asyncio.sleep(max(60, min(ai.exhausted_remaining(), 900)))
                continue
            for session in await mem.get_due_duo_sessions(BOT_NAME):
                try:
                    channel = bot.get_channel(session["channel_id"])
                    if not channel:
                        continue
                    target_message = None
                    async for candidate in channel.history(limit=8):
                        if not candidate.author.bot:
                            target_message = candidate
                            break
                    if not target_message:
                        continue
                    await mem.upsert_user(target_message.author.id, target_message.author.name, target_message.author.display_name)
                    user = await mem.get_user(target_message.author.id)
                    reply = await get_response(
                        target_message.author.id,
                        channel.id,
                        _duo_autoplay_prompt(session),
                        user,
                        target_message.author.display_name,
                        target_message.author.mention,
                        extra_context="DUO_AUTOPLAY: the other bot already spoke. Follow up naturally, keep it brief, and do not re-explain their point.",
                        channel_obj=channel,
                        is_dm=not bool(getattr(channel, "guild", None)),
                        direct_to_me=False,
                    )
                    if not (reply or "").strip():
                        continue
                    await channel.send(reply)
                    await mem.add_message(target_message.author.id, channel.id, "assistant", reply)
                    await _maybe_finalize_hidden_achievements(session, reply)
                    await _store_duo_story_progress(channel.id, session, reply)
                    if session.get("awaiting_bot") == BOT_NAME and session.get("autoplay_remaining", 0) <= 1 and session.get("mode") in {"trial", "mission", "interrogate", "truthdare", "compare"}:
                        await mem.resolve_duo_story(channel.id, session.get("mode", ""), _duo_outcome_payload(session, reply))
                    await mem.bump_duo_session(channel.id, BOT_NAME, partner_bot=PARTNER_NAME)
                except Exception as e:
                    log_error("duo_autoplay_session", e)
        except Exception as e:
            log_error("duo_autoplay_loop", e)
        await asyncio.sleep(8)


async def _rival_event_loop():
    await bot.wait_until_ready()
    await asyncio.sleep(45)
    while not bot.is_closed():
        try:
            if ai.is_exhausted():
                await asyncio.sleep(max(120, min(ai.exhausted_remaining(), 1800)))
                continue
            channels = await mem.get_active_channels()
            random.shuffle(channels)
            for channel_id, _ in channels:
                try:
                    channel = bot.get_channel(channel_id)
                    if not channel or not getattr(channel, "guild", None):
                        continue
                    if not PARTNER_BOT_ID or not channel.guild.get_member(PARTNER_BOT_ID):
                        continue
                    perms = channel.permissions_for(channel.guild.me) if channel.guild.me else None
                    if perms and (not perms.view_channel or not perms.send_messages):
                        continue
                    speaker_mode = await mem.get_channel_speaker_mode(channel_id)
                    if speaker_mode not in {"auto", "both"}:
                        continue
                    if await mem.get_duo_session(channel_id):
                        continue
                    topic = await _recent_rival_topic(channel)
                    if not topic:
                        continue
                    allowed, remaining = await mem.consume_shared_cooldown(f"rival_event:{channel_id}", 21600)
                    if not allowed:
                        debug_event("relationship", f"{BOT_NAME} rival_event_cooldown channel={channel_id} remaining={remaining}s")
                        continue
                    opener = await qai(
                        f"Users were discussing: '{topic}'. Start a spontaneous disagreement with {PARTNER_NAME}. "
                        "One or two sentences. Sound amused, cutting, and specific enough that the other bot has something to answer. "
                        "State only your own reaction; do not invent what the other bot thinks, feels, or said unless they actually said it in-channel.",
                        140,
                    )
                    opener = strip_narration(opener)
                    if not opener:
                        continue
                    await mem.set_duo_session(
                        channel_id,
                        "compare",
                        topic,
                        BOT_NAME,
                        awaiting_bot=PARTNER_NAME,
                        autoplay_turns=2,
                        autoplay_delay=random.randint(4, 8),
                        ttl_seconds=480,
                    )
                    await mem.start_duo_story(channel_id, "compare", topic, PARTNER_NAME)
                    await channel.send(opener)
                    await mem.bump_duo_session(channel_id, BOT_NAME, partner_bot=PARTNER_NAME)
                    debug_event("relationship", f"{BOT_NAME} rival_event channel={channel_id} topic={topic[:90]}")
                    break
                except Exception as e:
                    log_error("rival_event_channel", e)
        except Exception as e:
            log_error("rival_event_loop", e)
        await asyncio.sleep(random.randint(2400, 5400))


# ══════════════════════════════════════════════════════════════════════════════
# COMMANDS — all wrapped in try/except
# ══════════════════════════════════════════════════════════════════════════════

async def safe_reply(ctx, text):
    text = _unwrap_dialogue_quotes(text)
    if not (text or "").strip():
        return
    try: await ctx.reply(text)
    except Exception as e: log_error("safe_reply", e)

async def owner_reply(ctx, text, *, embed=None):
    """Send a response only the owner can see — via DM. Deletes the command message if possible."""
    try:
        dm = await ctx.author.create_dm()
        if embed:
            await dm.send(embed=embed)
        elif text and text.strip():
            await dm.send(text)
    except Exception as e:
        log_error("owner_reply_dm", e)
        if embed:
            try: await ctx.send(embed=embed)
            except: pass
        elif text and text.strip():
            try: await ctx.reply(text)
            except: pass

async def safe_send(ctx, text):
    text = _unwrap_dialogue_quotes(text)
    if not (text or "").strip():
        return
    try: await ctx.send(text)
    except Exception as e: log_error("safe_send", e)


def _partner_bot_present(ctx) -> bool:
    try:
        return bool(ctx.guild and PARTNER_BOT_ID and ctx.guild.get_member(PARTNER_BOT_ID))
    except Exception:
        return False


async def _should_handle_video_command(ctx) -> bool:
    try:
        if not _partner_bot_present(ctx):
            return True
        if not getattr(ctx, "guild", None):
            return True

        ref = getattr(ctx.message, "reference", None)
        if ref:
            ref_author_id = getattr(getattr(getattr(ref, "resolved", None), "author", None), "id", None)
            if ref_author_id is None and getattr(ref, "message_id", None):
                try:
                    fetched = await ctx.channel.fetch_message(ref.message_id)
                    ref_author_id = getattr(getattr(fetched, "author", None), "id", None)
                except Exception:
                    ref_author_id = None
            if ref_author_id == bot.user.id:
                return True
            if PARTNER_BOT_ID and ref_author_id == PARTNER_BOT_ID:
                return False

        mentions = {member.id for member in getattr(ctx.message, "mentions", [])}
        if bot.user and bot.user.id in mentions:
            return True
        if PARTNER_BOT_ID and PARTNER_BOT_ID in mentions:
            return False

        return True
    except Exception:
        return True


async def _send_video_summary(ctx, summary: dict, caption: str, *, sources_text: str = ""):
    try:
        video_bytes, filename = await fetch_rendered_video_bytes(summary or {})
        await ctx.send(caption, file=discord.File(io.BytesIO(video_bytes), filename=filename))
        if sources_text.strip():
            await ctx.send(sources_text[:1800])
    except Exception as e:
        log_error("send_video_summary", e)
        await safe_send(ctx, f"The video was rendered, but sending it failed. {e}")


def _teaching_video_title(topic: str | None, attachment) -> str:
    if topic and topic.strip():
        return topic.strip()[:80]
    if attachment and getattr(attachment, "filename", ""):
        return os.path.splitext(attachment.filename)[0][:80]
    return "Scaramouche Lesson"


def _teachvideo_time_hint(attachment=None) -> str:
    ext = os.path.splitext(getattr(attachment, "filename", "") or "")[1].lower()
    content_type = (getattr(attachment, "content_type", "") or "").lower()
    if content_type.startswith("video/") or ext in {".mp4", ".mov", ".mkv", ".webm"}:
        return "This will take around 4-7 minutes."
    if "pdf" in content_type or ext == ".pdf":
        return "Give me about 3-5 minutes."
    if ext in {".ppt", ".pptx", ".doc", ".docx"}:
        return "Give me about 3-5 minutes."
    if content_type.startswith("image/") or ext in {".png", ".jpg", ".jpeg", ".webp"}:
        return "Give me about 2-4 minutes."
    if ext in {".txt", ".md"}:
        return "This should take around 2-3 minutes."
    return "Give me about 2-4 minutes."


def _extract_google_doc_link(text: str) -> str:
    match = re.search(r"https?://docs\.google\.com/document/d/[A-Za-z0-9_-]+[^\s>]*", text or "", flags=re.IGNORECASE)
    return match.group(0) if match else ""


def _split_docfix_args(raw: str) -> tuple[str, str]:
    text = (raw or "").strip()
    link = _extract_google_doc_link(text)
    if not link:
        return "", text
    instructions = text.replace(link, "", 1).strip(" \n|")
    return link, instructions


def _docfix_prompt(display_name: str, original_text: str, instructions: str = "") -> str:
    extra = instructions.strip() or "Fix wording, grammar, clarity, and flow while preserving meaning and overall structure."
    return (
        f"You are Scaramouche editing a Google Doc for {display_name}.\n"
        f"Revise the wording so it reads cleaner and stronger, but do not invent facts or change the meaning.\n"
        f"Keep headings, paragraph breaks, and list-like structure where possible.\n"
        f"If the text is messy, make it coherent. If it is already strong, make only light refinements.\n"
        f"Specific instruction: {extra}\n\n"
        f"Return only the final revised document text. No commentary, no notes, no quotes.\n\n"
        f"DOCUMENT:\n{original_text[:10000]}"
    )


async def _rewrite_google_doc_text(display_name: str, original_text: str, instructions: str = "") -> str:
    revised = await qai(_docfix_prompt(display_name, original_text, instructions), 1800, self_edit=False, route="direct")
    cleaned = strip_narration(revised or "").strip()
    return cleaned if cleaned else original_text


def _weathervideo_time_hint(use_duo: bool) -> str:
    return "Give us about 3-6 minutes." if use_duo else "Give me about 2-4 minutes."


def _silence_reply_pool(user: dict | None) -> list[str]:
    if user and user.get("repair_progress", 0) > 0:
        return [
            "I saw it. I just wasn't ready to answer yet.",
            "You don't get instant absolution because you finally chose honesty.",
            "I was thinking. Don't make that sound sentimental.",
        ]
    return [
        "You don't get an immediate answer every time.",
        "I heard you. That doesn't mean I owed you speed.",
        "Silence was more useful than anything I would have said right away.",
        "I took a moment. Try not to look so startled.",
    ]


async def _remember_attachment_artifact(
    channel_id: int,
    user_id: int,
    name: str,
    summary: str,
    *,
    status: str = "recent",
):
    try:
        await mem.upsert_world_entity(
            "artifact",
            name[:120],
            summary=summary[:240],
            status=status[:40],
            channel_id=channel_id,
            owner_user_id=user_id,
            updated_by=BOT_NAME,
        )
        evidence_type = "pdf" if name.lower().endswith(".pdf") else "image" if any(name.lower().endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".webp", ".gif")) else "video" if any(name.lower().endswith(ext) for ext in (".mp4", ".mov", ".webm", ".avi")) else "artifact"
        await mem.add_evidence_item(
            channel_id,
            evidence_type,
            name[:120],
            summary[:240],
            source_excerpt=summary[:180],
            owner_user_id=user_id,
            updated_by=BOT_NAME,
        )
        if await mem.unlock_hidden_achievement(
            f"user:{user_id}",
            "remembered_artifact",
            "An image, screenshot, or file became a remembered artifact.",
        ):
            debug_event("memory", f"{BOT_NAME} achievement unlocked user={user_id} key=remembered_artifact")
    except Exception as e:
        log_error("remember_attachment_artifact", e)


async def _delayed_silence_reply(message, user: dict | None, *, is_dm: bool):
    key = (message.channel.id, message.author.id)
    try:
        await asyncio.sleep(random.randint(7, 16) if not is_dm else random.randint(5, 12))
        line = await _pick_fresh_pool_line(_silence_reply_pool(user), channel_id=message.channel.id, user_id=message.author.id)
        await message.reply(line)
        await mem.add_message(message.author.id, message.channel.id if not is_dm else message.author.id, "assistant", line)
    except Exception as e:
        log_error("delayed_silence_reply", e)
    finally:
        _silence_pending.discard(key)


async def _maybe_handle_silence_behavior(message, user: dict | None, content: str, *, is_dm: bool, direct_to_me: bool) -> bool:
    if not direct_to_me or not user:
        return False
    if detect_repair_signal(content):
        return False
    if not (user.get("conflict_open") or user.get("mood", 0) <= -6 or user.get("repair_progress", 0) > 0):
        return False
    key = (message.channel.id, message.author.id)
    if key in _silence_pending:
        return False
    chance = 0.13 if user.get("conflict_open") else 0.07
    if is_dm:
        chance += 0.04
    if random.random() >= chance:
        return False
    _silence_pending.add(key)
    debug_event("silence", f"{BOT_NAME} delaying reply user={message.author.id} channel={message.channel.id}")
    asyncio.ensure_future(_delayed_silence_reply(message, user, is_dm=is_dm))
    return True


async def _maybe_schedule_private_confession_scene(message, user: dict | None, content: str, *, is_dm: bool):
    if not is_dm or not user:
        return
    scene_note = describe_private_confession_scene(
        BOT_NAME,
        is_dm=is_dm,
        affection=user.get("affection", 0),
        trust=user.get("trust", 0),
        repair_progress=user.get("repair_progress", 0),
        conflict_open=bool(user.get("conflict_open")),
    )
    if not scene_note:
        return
    if random.random() >= 0.09:
        return
    allowed, _ = await mem.consume_phrase_with_status(f"{BOT_NAME}:private_confession:user:{message.author.id}", "private_confession_scene", 172800)
    if not allowed:
        return

    async def _run():
        try:
            await asyncio.sleep(random.randint(2, 5))
            prompt = (
                f"{scene_note}\n"
                f"You are in a private DM with {message.author.display_name}. Their latest message was: '{content[:220]}'. "
                "Send one additional message only. Make it feel like you almost kept it to yourself. "
                "One or two sentences. No narration."
            )
            line = await qai(prompt, 180, route="primary")
            line = strip_narration(line)
            if not line:
                return
            await message.channel.send(line)
            await mem.add_message(message.author.id, message.author.id, "assistant", line)
            if await mem.unlock_hidden_achievement(f"user:{message.author.id}", "scaramouche_private_confession", "He let a DM-only confession scene slip through."):
                debug_event("memory", f"{BOT_NAME} achievement unlocked user={message.author.id} key=scaramouche_private_confession")
        except Exception as e:
            log_error("private_confession_scene", e)

    asyncio.ensure_future(_run())


def _duo_outcome_payload(duo: dict | None, text: str) -> str:
    if not duo:
        return text[:180]
    tag = infer_duo_ending_tag(duo.get("mode", ""), text)
    return f"ENDING:{tag}|{strip_narration(text)[:220]}"


def _achievement_gallery_lines(entries: list[dict]) -> list[str]:
    unlocked = entries[:8]
    lines = ["Achievement gallery:"]
    if unlocked:
        for item in unlocked:
            note = (item.get("note") or "").strip()
            line = f"- {item.get('achievement_key', 'unknown')}"
            if note:
                line += f" — {note[:110]}"
            lines.append(line)
    else:
        lines.append("- Nothing unlocked yet. Try being more interesting.")
    lines.append("Rumors: a private confession, a merciful ending, a remembered artifact, a peaceful duet, a survived trial, a contradictory memory, a leaked private opinion.")
    return lines


async def _store_duo_story_progress(channel_id: int, duo: dict | None, text: str):
    if not duo:
        return
    mode = duo.get("mode", "")
    if mode in {"trial", "mission", "interrogate", "truthdare", "compare", "argue", "duet"}:
        cleaned = strip_narration(text)
        summary = f"{BOT_NAME} {mode} turn: {cleaned[:180]}"
        await mem.note_duo_story_progress(channel_id, mode, summary)
        topic = duo.get("topic", "") or mode
        await mem.note_shared_event_memory(
            channel_id,
            topic,
            BOT_NAME,
            summary,
            truth_hint=f"{mode} scene memory",
        )
        if mode in {"trial", "interrogate", "compare"}:
            await mem.add_evidence_item(
                channel_id,
                "quote",
                f"{BOT_NAME} {mode} turn",
                cleaned[:220],
                source_excerpt=cleaned[:180],
                updated_by=BOT_NAME,
            )


async def _command_face_media(ctx):
    img = next((a for a in ctx.message.attachments if a.content_type and "image" in a.content_type), None)
    vid = next((a for a in ctx.message.attachments
               if (a.content_type and a.content_type in VIDEO_TYPES) or
                  any(a.filename.lower().endswith(ext) for ext in VIDEO_EXTS)), None)
    if not img and not vid and ctx.message.reference:
        try:
            ref_msg = await ctx.channel.fetch_message(ctx.message.reference.message_id)
            img = next((a for a in ref_msg.attachments if a.content_type and "image" in a.content_type), None)
            vid = next((a for a in ref_msg.attachments
                       if (a.content_type and a.content_type in VIDEO_TYPES) or
                          any(a.filename.lower().endswith(ext) for ext in VIDEO_EXTS)), None)
        except Exception:
            pass
    return img, vid


async def _interaction_reply(interaction: discord.Interaction, text: str, *, thinking: bool = False, view: discord.ui.View | None = None):
    try:
        text = _unwrap_dialogue_quotes(text)
        if not (text or "").strip():
            return
        if thinking and not interaction.response.is_done():
            await interaction.response.defer(thinking=True)
        if interaction.response.is_done():
            await interaction.followup.send(text, view=view)
        else:
            await interaction.response.send_message(text, view=view)
    except Exception as e:
        log_error("interaction_reply", e)


async def _reply_and_store(ctx, text: str):
    await safe_reply(ctx, text)
    try:
        duo = await mem.get_duo_session(ctx.channel.id)
        await mem.add_message(ctx.author.id, ctx.channel.id, "assistant", text)
        await _maybe_finalize_hidden_achievements(duo, text)
        await _store_duo_story_progress(ctx.channel.id, duo, text)
        if duo and duo.get("awaiting_bot") == BOT_NAME and duo.get("autoplay_remaining", 0) <= 1 and duo.get("mode") in {"trial", "mission", "interrogate", "truthdare", "compare"}:
            await mem.resolve_duo_story(ctx.channel.id, duo.get("mode", ""), _duo_outcome_payload(duo, text))
        await mem.bump_duo_session(ctx.channel.id, BOT_NAME, partner_bot=PARTNER_NAME)
    except Exception as e:
        log_error("reply_and_store", e)


async def _reply_and_store_interaction(interaction: discord.Interaction, text: str):
    await _interaction_reply(interaction, text, thinking=False)
    try:
        duo = await mem.get_duo_session(interaction.channel_id)
        await mem.add_message(interaction.user.id, interaction.channel_id, "assistant", text)
        await _maybe_finalize_hidden_achievements(duo, text)
        await _store_duo_story_progress(interaction.channel_id, duo, text)
        if duo and duo.get("awaiting_bot") == BOT_NAME and duo.get("autoplay_remaining", 0) <= 1 and duo.get("mode") in {"trial", "mission", "interrogate", "truthdare", "compare"}:
            await mem.resolve_duo_story(interaction.channel_id, duo.get("mode", ""), _duo_outcome_payload(duo, text))
        await mem.bump_duo_session(interaction.channel_id, BOT_NAME, partner_bot=PARTNER_NAME)
    except Exception as e:
        log_error("reply_and_store_interaction", e)


def _format_memory_snapshot(user: dict | None, topics: list[dict], memories: list[dict], scene: dict | None) -> str:
    lines = []
    callback = (user or {}).get("callback_memory")
    if callback:
        lines.append(f"Callback: {callback[:140]}")
    if topics:
        topic_bits = ", ".join(f"{item['topic']} ({item['count']})" for item in topics[:4])
        lines.append(f"Topics: {topic_bits}")
    if memories:
        memory_bits = " | ".join(f"{item['kind']}: {item['memory'][:70]}" for item in memories[:4])
        lines.append(f"Memory bank: {memory_bits}")
    scene_desc = describe_scene_state(scene)
    if scene_desc:
        lines.append(f"Scene: {scene_desc}")
    return "\n".join(lines) if lines else "Nothing worth preserving yet. Try harder."


async def _run_duo_prompt_for_actor(
    author,
    channel,
    mode: str,
    prompt: str,
    *,
    session_topic: str | None = None,
    story: bool = False,
    enemy: str = "",
    extra_context: str = "",
) -> tuple[dict | None, str]:
    user = await _setup_user(author)
    autoplay_turns = DUO_CHAIN_TURNS.get(mode, 1)
    if not (user or {}).get("duo_autoplay", True):
        autoplay_turns = 1
    await mem.set_duo_session(
        channel.id,
        mode,
        session_topic or prompt,
        BOT_NAME,
        initiator_user_id=author.id,
        awaiting_bot=PARTNER_NAME,
        autoplay_turns=autoplay_turns,
        autoplay_delay=6,
    )
    if story:
        await mem.start_duo_story(channel.id, mode, session_topic or prompt, enemy=enemy)
    reply = await get_response(
        author.id,
        channel.id,
        prompt,
        user,
        getattr(author, "display_name", getattr(author, "name", "you")),
        author.mention,
        extra_context=extra_context,
        channel_obj=channel,
        is_dm=isinstance(channel, discord.DMChannel),
    )
    return user, reply

@bot.command(name="voice",aliases=["speak","say"])
async def voice_cmd(ctx,*,msg:str=None):
    try:
        normalized = (msg or "").strip().lower()
        user=await _setup(ctx); mood_val=user.get("mood",0) if user else 0
        if normalized in {"on", "off", "status"}:
            if normalized == "status":
                await safe_reply(ctx, f"Voice notes are `{_pref_label(user.get('voice_enabled', True) if user else True)}`.")
            else:
                enabled = normalized == "on"
                await mem.set_user_preference(ctx.author.id, "voice_enabled", int(enabled))
                await safe_reply(ctx, f"Fine. Voice notes are `{_pref_label(enabled)}` now.")
            return
        if not msg: msg="You summoned me without saying a word. How impressively useless."
        async with ctx.typing():
            text_reply=await get_response(ctx.author.id,ctx.channel.id,msg,user,ctx.author.display_name,ctx.author.mention)
            sent=await send_voice(ctx.channel,text_reply,mood=mood_val,guild=ctx.guild,user=user,user_id=ctx.author.id)
        if sent:
            await mem.add_message(ctx.author.id, ctx.channel.id, "assistant", f"[voice message] {text_reply}")
        else:
            await safe_reply(ctx,text_reply)
            await mem.add_message(ctx.author.id, ctx.channel.id, "assistant", text_reply)
    except Exception as e: log_error("voice_cmd",e); await safe_reply(ctx,"Hmph.")


@bot.command(name="tedtalk", aliases=["teach","lecture","explain"])
async def tedtalk_cmd(ctx, *, topic: str = None):
    try:
        # Lock on message ID — prevents duplicate fires from Discord edit events
        msg_id = ctx.message.id
        if msg_id in _tedtalk_active:
            return  # Silent — same message, just ignore
        _tedtalk_active.add(msg_id)

        await _setup(ctx)

        attachment = ctx.message.attachments[0] if ctx.message.attachments else None

        if not attachment and not topic:
            _tedtalk_active.discard(msg_id)
            await safe_reply(ctx, "Attach a file or give me a topic. I can't teach you nothing, as satisfying as that would be.")
            return

        file_size = attachment.size if attachment else 0
        if file_size > 500_000 or (attachment and attachment.filename.lower().endswith(".pptx")):
            time_hint = "This will take roughly 2-3 minutes."
        elif file_size > 100_000:
            time_hint = "Give me about a minute."
        else:
            time_hint = "This will take about 30-60 seconds."

        ack_lines = [
            f"Fine. Sit down, pay attention, and try not to embarrass yourself. {time_hint}",
            f"You want me to teach you something. How refreshingly self-aware of you to admit you need help. {time_hint}",
            f"Hmph. I'll condescend to explain this. Try to keep up. {time_hint}",
            f"...You actually want to learn. I find that mildly less irritating than most things. Fine. {time_hint}",
        ]
        await ctx.reply(random.choice(ack_lines))
        asyncio.ensure_future(_do_tedtalk(ctx, attachment, topic, msg_id))

    except Exception as e:
        _tedtalk_active.discard(ctx.message.id)
        log_error("tedtalk_cmd", e)
        await safe_reply(ctx, "...Something went wrong. Annoying.")


async def _do_tedtalk(ctx, attachment, topic, msg_id=None):
    """Background task for !tedtalk — does all the heavy lifting."""
    try:
        material_content = ""

        # Immediately confirm he's working so user knows it started
        await ctx.send(random.choice([
            "...I'm reading it now. Don't interrupt me.",
            "Hmph. Give me a moment. I'm going through your material.",
            "I'm working on it. Try not to send me more messages while I'm busy.",
            "...Processing. I'll tell you when I'm done.",
        ]))

        # ── Extract content from attachment ──────────────────────────────
        if attachment:
            ct = (attachment.content_type or "").lower()
            import base64, aiohttp as _ah

            try:
                async with _ah.ClientSession() as s:
                    async with s.get(attachment.url) as r:
                        file_bytes = await r.read()
            except Exception as e:
                await ctx.send(f"Couldn't download the file. {e}"); return

            if "pdf" in ct or attachment.filename.lower().endswith(".pdf"):
                try:
                    try:
                        import pdfplumber

                        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                            material_content = "\n".join(page.extract_text() or "" for page in pdf.pages)[:8000]
                    except Exception:
                        try:
                            import PyPDF2

                            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                            material_content = "\n".join(page.extract_text() or "" for page in reader.pages)[:8000]
                        except Exception:
                            material_content = ""
                    if not material_content.strip():
                        await ctx.send("Couldn't read the PDF text. Try an image or PPTX instead."); return
                except Exception as e:
                    await ctx.send(f"Couldn't read the PDF: {e}"); return

            elif "image" in ct or attachment.filename.lower().endswith((".png",".jpg",".jpeg",".webp",".gif")):
                try:
                    img_b64 = base64.b64encode(file_bytes).decode()
                    media_type = ct if ct else "image/jpeg"
                    def _extract_img():
                        r = ai.call_with_retry(
                            model=GROQ_VISION_MODEL,
                            max_tokens=2000,
                            messages=[{"role":"user","content":[
                                {"type":"image_url","image_url":{"url":f"data:{media_type};base64,{img_b64}"}},
                                {"type":"text","text":"Extract all educational content visible in this image. Include every concept, formula, definition, and key point."}
                            ]}]
                        )
                        return r.choices[0].message.content.strip() if r.choices else ""
                    extract_resp_text = await asyncio.get_event_loop().run_in_executor(None, _extract_img)
                    material_content = extract_resp_text
                except Exception as e:
                    await ctx.send(f"Couldn't read the image: {e}"); return

            elif "text" in ct or attachment.filename.lower().endswith((".txt",".md",".csv")):
                try:
                    material_content = file_bytes.decode("utf-8", errors="ignore")[:4000]
                except Exception as e:
                    await ctx.send(f"Couldn't read the text file: {e}"); return

            elif attachment.filename.lower().endswith((".pptx",".ppt")):
                try:
                    from pptx import Presentation as _Prs
                    prs = _Prs(io.BytesIO(file_bytes))
                    parts = []
                    for i, slide in enumerate(prs.slides):
                        slide_texts = []
                        for shape in slide.shapes:
                            if hasattr(shape, "text") and shape.text.strip():
                                slide_texts.append(shape.text.strip())
                        if slide_texts:
                            parts.append(f"[Slide {i+1}]\n" + "\n".join(slide_texts))
                    material_content = "\n\n".join(parts)[:4000]
                except Exception as e:
                    await ctx.send(f"Couldn't read the PowerPoint: {e}"); return

            elif attachment.filename.lower().endswith((".docx",".doc")):
                try:
                    import docx as _docx
                    doc = _docx.Document(io.BytesIO(file_bytes))
                    parts = []
                    # Paragraphs
                    for p in doc.paragraphs:
                        if p.text.strip():
                            parts.append(p.text.strip())
                    # Tables (cheat sheets often use tables)
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join(
                                cell.text.strip() for cell in row.cells if cell.text.strip()
                            )
                            if row_text:
                                parts.append(row_text)
                    material_content = "\n".join(parts)[:4000]
                    if not material_content:
                        await ctx.send("The Word document appears to be empty or uses unsupported formatting."); return
                except Exception as e:
                    await ctx.send(f"Couldn't read the Word document: {e}"); return
            else:
                await ctx.send("I can read PDFs, images, PowerPoint files, Word documents, and text files. Whatever that is, I can't work with it."); return

        if topic:
            material_content = f"Topic: {topic}\n\n{material_content}".strip()

        if not material_content:
            await ctx.send("There was nothing readable in that file. How typical."); return

        # ── Generate script ───────────────────────────────────────────────
        await ctx.send(random.choice([
            "...Fine. I'm reading it. Don't rush me.",
            "Hmph. Give me a moment. I'm processing your inadequate study material.",
            "I'm going through this. Try not to fidget.",
            "...Reviewing the material. It's about what I expected.",
        ]))
        try:
            script_prompt = (
                f"You are Scaramouche — the Sixth Fatui Harbinger, the Balladeer.\n"
                f"Teach the following material to {ctx.author.display_name}.\n\n"
                f"MATERIAL:\n{material_content[:3000]}\n\n"
                f"Write a complete spoken teaching monologue. "
                f"Teach ALL key concepts correctly and thoroughly. "
                f"Stay in character — contemptuous but accurate. "
                f"Decide length based on complexity. "
                f"Structure: introduce → explain each concept → examples → summary. "
                f"NO asterisk actions. Spoken words only."
            )
            def _gen_script():
                r = ai.call_with_retry(
                    model=GROQ_MODEL,
                    max_tokens=2500,
                    messages=[{"role":"system","content":_BASE},
                              {"role":"user","content":script_prompt}]
                )
                return strip_narration(r.choices[0].message.content.strip() if r.choices else "")
            script = await asyncio.get_event_loop().run_in_executor(None, _gen_script)
        except Exception as e:
            await ctx.send(f"Failed to generate the lecture: {e}"); return

        if not script:
            await ctx.send("...I had nothing to say. Unlikely, but here we are."); return

        # ── Generate audio in chunks ──────────────────────────────────────
        await ctx.send(random.choice([
            "Script complete. Now rendering my voice. This is beneath me but here we are.",
            "...I've written the lecture. Generating audio. Wait.",
            "Hmph. The content is ready. Give me a moment to make it sound appropriately contemptuous.",
            "Preparing to speak. Try to actually listen this time.",
        ]))

        sentences  = re.split(r'(?<=[.!?])\s+', script)
        chunks, current = [], ""
        for s in sentences:
            if len(current) + len(s) + 1 <= 900:
                current = (current + " " + s).strip()
            else:
                if current: chunks.append(current)
                current = s
        if current: chunks.append(current)

        audio_parts = []
        total_chunks = len([c for c in chunks if c.strip()])
        await ctx.send(random.choice([
            f"*Recording. {total_chunks} segments. Don't touch anything.*",
            f"*Committing this to voice. {total_chunks} parts. Try not to interrupt.*",
            f"*{total_chunks} segments to render. I'm working. Be quiet.*",
            f"*Converting my lecture to audio. {total_chunks} parts. This takes time.*",
        ]))

        for i, chunk in enumerate(chunks):
            if not chunk.strip(): continue
            try:
                audio = await get_audio_with_mood(tts_safe(chunk, ctx.guild), 0)
                if audio:
                    audio_parts.append(audio)
                else:
                    print(f"[tedtalk] chunk {i} returned None")
            except Exception as e:
                print(f"[tedtalk] chunk {i} error: {e}")

            if (i+1) % 5 == 0:
                try:
                    remaining = total_chunks - (i+1)
                    await ctx.send(random.choice([
                        f"*{i+1}/{total_chunks} done. {remaining} remaining. Still working.*",
                        f"*Progress: {i+1} of {total_chunks}. Don't ask how much longer.*",
                        f"*{remaining} segments left. I said be quiet.*",
                    ]))
                except: pass

        # ── Send audio ────────────────────────────────────────────────────
        await ctx.send(random.choice([
            f"*Done. {len(audio_parts)} of {total_chunks} segments rendered. Sending now.*",
            f"*{len(audio_parts)}/{total_chunks} segments complete. Here.*",
            f"*Finished. {len(audio_parts)} parts. Pay attention this time.*",
        ]))

        if not audio_parts:
            await ctx.send("Voice synthesis failed for all segments. Sending written version instead.")
            for i in range(0, len(script), 1900):
                try: await ctx.send(script[i:i+1900])
                except Exception as e: await ctx.send(f"*(text send failed: {e})*")
            return

        MAX_BYTES = 7 * 1024 * 1024
        current_batch, part_num = b"", 1

        for audio_chunk in audio_parts:
            if len(current_batch) + len(audio_chunk) > MAX_BYTES:
                try:
                    await ctx.send(
                        f"🎙️ *Part {part_num}:*",
                        file=discord.File(io.BytesIO(current_batch), filename=f"lecture_p{part_num}.mp3")
                    )
                except Exception as e:
                    await ctx.send(f"*(Audio part {part_num} failed: {e})*")
                part_num += 1
                current_batch = audio_chunk
                await asyncio.sleep(1)
            else:
                current_batch += audio_chunk

        if current_batch:
            label = f"Part {part_num}" if part_num > 1 else "Lecture"
            try:
                await ctx.send(
                    f"🎙️ *{label}:*",
                    file=discord.File(io.BytesIO(current_batch), filename=f"lecture_p{part_num}.mp3")
                )
            except Exception as e:
                await ctx.send(f"*(Final audio failed: {e})*")

        # ── Cache material for follow-up questions ────────────────────────
        _tedtalk_cache[ctx.author.id] = {
            "material":   material_content[:3000],
            "channel_id": ctx.channel.id,
            "ts":         time.time(),
        }

        # ── Send notes (not transcript) ───────────────────────────────────
        try:
            def _gen_notes():
                r = ai.call_with_retry(
                    model=GROQ_MODEL,
                    max_tokens=800,
                    messages=[{"role":"system","content":_BASE},
                              {"role":"user","content":(
                        f"You just gave a lecture on this material:\n{material_content[:2000]}\n\n"
                        f"Write concise study notes for {ctx.author.display_name}. "
                        f"Key terms, important concepts, things to remember. "
                        f"Bullet points are fine. Keep it short — this is a reference, not a repeat of the lecture. "
                        f"Stay in character but be genuinely useful."
                    )}]
                )
                return r.choices[0].message.content.strip() if r.choices else ""
            notes = await asyncio.get_event_loop().run_in_executor(None, _gen_notes)
            if notes:
                await ctx.send(f"📋 *Notes:*\n{notes[:1900]}")
        except Exception as e:
            log_error("tedtalk_notes", e)

    except Exception as e:
        log_error("_do_tedtalk", e)
        try: await ctx.send(f"...Something went wrong mid-lecture. Error: {e}")
        except: pass
    finally:
        if msg_id: _tedtalk_active.discard(msg_id)


@bot.command(name="teachvideo", aliases=["videoteach", "lessonvideo", "presentationvideo"])
async def teachvideo_cmd(ctx, *, topic: str = None):
    try:
        if not await _should_handle_video_command(ctx):
            return
        msg_id = ctx.message.id
        if msg_id in _teachvideo_active:
            return
        _teachvideo_active.add(msg_id)
        await _setup(ctx)
        attachment = ctx.message.attachments[0] if ctx.message.attachments else None
        if not attachment and not (topic or "").strip():
            _teachvideo_active.discard(msg_id)
            await safe_reply(ctx, "Attach notes or give me a lesson topic. I won't invent a presentation out of thin air.")
            return
        available, reason = video_renderer_available(BOT_NAME)
        if not available:
            _teachvideo_active.discard(msg_id)
            await safe_reply(ctx, f"Video rendering is unavailable on this host. {reason}")
            return
        time_hint = _teachvideo_time_hint(attachment)
        await ctx.reply(random.choice([
            f"Fine. If you want the explanation as a video, sit down and wait. {time_hint}",
            f"A presentation video? Hmph. I'll make it watchable. {time_hint}",
            f"I'll render the lesson properly. Try not to waste my effort. {time_hint}",
        ]))
        asyncio.ensure_future(_do_teachvideo(ctx, attachment, topic, msg_id))
    except Exception as e:
        _teachvideo_active.discard(ctx.message.id)
        log_error("teachvideo_cmd", e)
        await safe_reply(ctx, "...The video request failed before it even started.")


async def _do_teachvideo(ctx, attachment, topic, msg_id=None):
    try:
        time_hint = _teachvideo_time_hint(attachment)
        await ctx.send(random.choice([
            f"I'm turning it into slides now. Don't interrupt me. {time_hint}",
            f"Rendering the presentation. This takes longer than words. {time_hint}",
            f"I'm building the lesson video. Wait. {time_hint}",
        ]))
        material_content = await extract_teaching_material(attachment, topic=topic or "", bot_name=BOT_NAME)
        if not material_content.strip():
            await ctx.send("There was nothing usable in that material. Typical."); return
        summary = await render_teaching_video(BOT_NAME, material_content, title=_teaching_video_title(topic, attachment))
        await _send_video_summary(
            ctx,
            summary,
            random.choice([
                "There. A proper lesson video.",
                "You wanted a presentation. Here it is.",
                "Finished. Try to learn something from it.",
            ]),
        )
    except Exception as e:
        log_error("do_teachvideo", e)
        await safe_send(ctx, f"The presentation video failed. {e}")
    finally:
        if msg_id:
            _teachvideo_active.discard(msg_id)


@bot.command(name="fixdoc", aliases=["docfix", "gdocfix", "editdoc"])
async def fixdoc_cmd(ctx, *, doc_and_instructions: str = ""):
    try:
        await _setup(ctx)
        doc_link, instructions = _split_docfix_args(doc_and_instructions)
        if not doc_link:
            await safe_reply(
                ctx,
                "Send a Google Docs link with the command. Example: `!fixdoc <google-doc-link> tighten the wording but keep the meaning`.",
            )
            return
        ready, reason = google_docs_ready()
        share_email = service_account_email()
        await ctx.reply(
            random.choice(
                [
                    "Fine. I'll go through the doc and clean up the wording. Give me a minute.",
                    "I'm reading the document now. Try not to hover over me while I fix it.",
                    "Hmph. I'll revise it directly in the doc. Wait.",
                ]
            )
        )
        async with ctx.typing():
            if ready:
                doc = await asyncio.get_event_loop().run_in_executor(None, lambda: fetch_google_doc(doc_link))
                if not (doc.get("text") or "").strip():
                    await safe_send(ctx, "The Google Doc is empty, or there was nothing readable in it.")
                    return
                revised = await _rewrite_google_doc_text(ctx.author.display_name, doc["text"], instructions)
                result = await asyncio.get_event_loop().run_in_executor(None, lambda: overwrite_google_doc(doc_link, revised))
            else:
                result = await remote_fix_google_doc(
                    doc_link,
                    bot_name=BOT_NAME,
                    display_name=ctx.author.display_name,
                    instructions=instructions,
                )
        if ready:
            await safe_send(
                ctx,
                f"Done. I updated `{result['title']}` in place. If Google blocks access later, share the doc with `{share_email}` as an editor.",
            )
        else:
            await safe_send(
                ctx,
                f"Done. I updated `{result['title']}` through the worker because Railway does not have the Google credential yet. If access fails, share the doc with `{share_email}` as an editor. Local note: {reason}",
            )
    except Exception as e:
        log_error("fixdoc_cmd", e)
        await safe_reply(ctx, f"I couldn't update that Google Doc. {e}")


@bot.command(name="dare")
async def dare_cmd(ctx):
    try:
        user=await _setup(ctx)
        reply=await qai(f"Give {ctx.author.display_name} a dare. Dark, specific, theatrical. 1-2 sentences.",200)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("dare_cmd",e)

@bot.command(name="fortune",aliases=["fortunecookie"])
async def fortune_cmd(ctx):
    try:
        reply=await qai("Fortune cookie message rewritten as a cold theatrical threat. 1 sentence.",100)
        await safe_reply(ctx,f"🥠 *{reply}*")
    except Exception as e: log_error("fortune_cmd",e)

@bot.command(name="trivia")
async def trivia_cmd(ctx):
    try:
        question=await qai(
            "One difficult Genshin lore trivia question. Include answer in brackets [ANSWER: ...]. "
            "Also include a short [SOURCE: ...] note naming the lore source. Be obscure.",
            260,
        )
        answer_match = re.search(r"\[ANSWER:\s*(.*?)\]", question, re.IGNORECASE)
        source_match = re.search(r"\[SOURCE:\s*(.*?)\]", question, re.IGNORECASE)
        clean_question = re.sub(r"\s*\[ANSWER:.*?\]\s*", "", question, flags=re.IGNORECASE).strip()
        clean_question = re.sub(r"\s*\[SOURCE:.*?\]\s*", "", clean_question, flags=re.IGNORECASE).strip()
        await mem.set_active_trivia(
            ctx.channel.id,
            ctx.author.id,
            clean_question,
            answer_match.group(1).strip() if answer_match else "",
            source_match.group(1).strip() if source_match else "",
        )
        await safe_reply(ctx,clean_question)
    except Exception as e: log_error("trivia_cmd",e)

@bot.command(name="answer")
async def answer_cmd(ctx,*,response:str=None):
    try:
        if not response: await safe_reply(ctx,"Answer *what*?"); return
        await _setup(ctx)
        trivia = await mem.get_active_trivia(ctx.channel.id)
        if trivia and trivia.get("answer"):
            result=await qai(
                f"Question: {trivia['question']}\nCanonical answer: {trivia['answer']}\n"
                f"Source note: {trivia.get('source_note','unknown')}\n"
                f"{ctx.author.display_name} answered with: '{response}'. Was it right or wrong? Check against the canonical answer. "
                "Be brutal but fair. 1-2 sentences.",
                180,
            )
            correct = bool(re.search(r"\b(correct|right)\b", result.lower())) and "wrong" not in result.lower()
            await mem.clear_active_trivia(ctx.channel.id)
        else:
            result=await qai(f"{ctx.author.display_name} answered a trivia question with: '{response}'. Was it right or wrong? Check against Genshin lore. Be brutal. 1-2 sentences.",150)
            correct="right" in result.lower() or "correct" in result.lower()
        await mem.update_trivia(ctx.author.id,correct)
        stats = await mem.get_trivia_stats(ctx.author.id)
        await safe_reply(ctx,f"{result}\nScore: {stats['correct']} right, {stats['wrong']} wrong ({stats['accuracy']}% accuracy).")
    except Exception as e: log_error("answer_cmd",e)

@bot.command(name="roast",aliases=["roastbattle"])
async def roast_cmd(ctx,member:discord.Member=None):
    try:
        if not member: await safe_reply(ctx,"Roast *who*?"); return
        battle=await mem.get_active_roast(ctx.channel.id)
        if battle:
            await mem.increment_roast_round(battle["id"])
            if battle["round"]>=5:
                await mem.end_roast_battle(battle["id"])
                prompt=f"Roast battle over after 5 rounds. Scoreboard so far: {battle.get('scores', {})}. Declare final winner between {ctx.author.display_name} and {member.display_name}. Dramatic. 2-3 sentences."
            else:
                prompt=f"Judging roast battle round {battle['round']+1}. {ctx.author.display_name} fired at {member.display_name}. Score this round theatrically. 2-3 sentences."
        else:
            await mem.start_roast_battle(ctx.channel.id,ctx.author.id,member.id)
            prompt=f"You're refereeing a roast battle between {ctx.author.display_name} and {member.display_name}. Open theatrically. 5 rounds, you judge. 2-3 sentences."
        reply=await qai(prompt,300)
        lowered = reply.lower()
        winner_id = None
        if ctx.author.display_name.lower() in lowered and member.display_name.lower() not in lowered:
            winner_id = ctx.author.id
        elif member.display_name.lower() in lowered and ctx.author.display_name.lower() not in lowered:
            winner_id = member.id
        if battle and winner_id:
            await mem.award_roast_round(battle["id"], winner_id)
        updated = await mem.get_active_roast(ctx.channel.id) if battle else None
        score_line = f"\nScoreboard: {updated['scores']}" if updated and updated.get("scores") else ""
        await ctx.send(f"{ctx.author.mention} vs {member.mention}\n{reply}{score_line}")
    except Exception as e: log_error("roast_cmd",e)

@bot.command(name="hostage")
async def hostage_cmd(ctx):
    try:
        await _setup(ctx)
        if ctx.author.id in _hostages:
            await safe_reply(ctx,f"You haven't fulfilled your end yet. — *{_hostages[ctx.author.id]}*"); return
        demand=await qai(f"You've taken {ctx.author.display_name}'s good mood hostage. State your demand theatrically. 1-2 sentences.",150)
        _hostages[ctx.author.id]=demand
        await safe_reply(ctx,f"...I've taken your good mood hostage. You'll get it back when you: {demand}")
    except Exception as e: log_error("hostage_cmd",e)

@bot.command(name="release",aliases=["freed","ransom"])
async def release_cmd(ctx,*,offering:str=None):
    try:
        if ctx.author.id not in _hostages: await safe_reply(ctx,"Nothing is being held hostage. For now."); return
        demand=_hostages[ctx.author.id]
        result=await qai(f"You held {ctx.author.display_name}'s good mood hostage with: '{demand}'. They offered: '{offering or 'nothing'}'. Accept or refuse theatrically.",150)
        if any(w in result.lower() for w in ["accept","fine","release","granted"]):
            del _hostages[ctx.author.id]
        await safe_reply(ctx,result)
    except Exception as e: log_error("release_cmd",e)

@bot.command(name="impersonate",aliases=["imitate","be"])
async def impersonate_cmd(ctx,*,character:str=None):
    try:
        if not character: await safe_reply(ctx,"Impersonate *who*?"); return
        reply=await qai(f"Briefly speak as {character} from Genshin for 2 sentences, but interrupt yourself with your own editorial commentary constantly. You find this beneath you.",250)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("impersonate_cmd",e)

@bot.command(name="opinion")
async def opinion_cmd(ctx,*,character:str=None):
    try:
        if not character: await safe_reply(ctx,"Opinion on *who*?"); return
        reply=await qai(f"Your honest unfiltered personal opinion of {character} from Genshin Impact. 2-3 sentences.",250)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("opinion_cmd",e)

@bot.command(name="poll")
async def poll_cmd(ctx,*,question:str=None):
    try:
        if not question: await safe_reply(ctx,"A poll about *what*?"); return
        framing=await qai(f"Frame this as a demand for answers: '{question}'. 1 sentence.",80)
        msg=await ctx.send(f"📊 {framing}\n\n**{question}**")
        for emoji in ["👍","👎","🤷"]:
            try: await msg.add_reaction(emoji)
            except: pass
    except Exception as e: log_error("poll_cmd",e)

@bot.command(name="summarize",aliases=["recap"])
async def summarize_cmd(ctx):
    try:
        recent=await mem.get_channel_recent(ctx.channel.id,20)
        if not recent: await safe_reply(ctx,"Nothing worth summarizing. Which tracks."); return
        sample="\n".join(f"{m['name']}: {m['content']}" for m in recent[:15])[:800]
        reply=await qai(f"Summarize this conversation with contemptuous commentary:\n{sample}\nBe cutting and specific. 3-4 sentences.",300)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("summarize_cmd",e)

@bot.command(name="mute",aliases=["silence","ignore"])
async def mute_cmd(ctx,member:discord.Member=None,minutes:int=10):
    try:
        target=member or ctx.author
        mem.mute_user(target.id,minutes*60)
        reply=await qai(f"You've decided to 'mute' {target.display_name} for {minutes} minutes. Announce theatrically. 1-2 sentences.",120)
        if member: await ctx.send(f"{member.mention} {reply}")
        else: await safe_reply(ctx,reply)
    except Exception as e: log_error("mute_cmd",e)

@bot.command(name="unmute",aliases=["unsilence"])
async def unmute_cmd(ctx,member:discord.Member=None):
    try:
        target=member or ctx.author
        mem.unmute_user(target.id)
        await safe_reply(ctx,f"...Fine. {target.display_name} may speak again. Lucky them.")
    except Exception as e: log_error("unmute_cmd",e)

@bot.command(name="spar")
async def spar_cmd(ctx,*,opening:str=None):
    try:
        user=await _setup(ctx)
        prompt=f"{ctx.author.display_name} challenged you: '{opening or 'Come on then.'}'. Fire back. End with a challenge."
        reply=await get_response(ctx.author.id,ctx.channel.id,prompt,user,ctx.author.display_name,ctx.author.mention)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("spar_cmd",e)

@bot.command(name="duel")
async def duel_cmd(ctx,member:discord.Member=None):
    try:
        if not member or member==ctx.author: await safe_reply(ctx,"Duel *who*?"); return
        u1=" | ".join((await mem.get_recent_messages(ctx.author.id,3))[:3])[:150]
        u2=" | ".join((await mem.get_recent_messages(member.id,3))[:3])[:150]
        reply=await qai(f"Referee insult duel: {ctx.author.display_name} (says:'{u1}') vs {member.display_name} (says:'{u2}'). Analyze both, declare winner. 3-4 sentences.",300)
        await ctx.send(f"{ctx.author.mention} vs {member.mention}\n{reply}")
    except Exception as e: log_error("duel_cmd",e)

@bot.command(name="judge")
async def judge_cmd(ctx,member:discord.Member=None):
    try:
        target=member or ctx.author
        sample=" | ".join(await mem.get_recent_messages(target.id,8))[:400]
        reply=await qai(f"Brutal assessment of {target.display_name}"+(f" — words:'{sample}'" if sample else "")+". 2-4 sentences.",250)
        if member: await ctx.send(f"{member.mention} {reply}")
        else: await safe_reply(ctx,reply)
    except Exception as e: log_error("judge_cmd",e)

@bot.command(name="prophecy")
async def prophecy_cmd(ctx,member:discord.Member=None):
    try:
        target=member or ctx.author
        reply=await qai(f"Cryptic threatening prophecy for {target.display_name}. 2-3 sentences.",200)
        if member: await ctx.send(f"{member.mention} {reply}")
        else: await safe_reply(ctx,reply)
    except Exception as e: log_error("prophecy_cmd",e)

@bot.command(name="rate")
async def rate_cmd(ctx,*,thing:str=None):
    try:
        if not thing: await safe_reply(ctx,"Rate *what*?"); return
        reply=await qai(f"Rate '{thing}' out of 10. Score first, 1-2 sentences of contempt.",180)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("rate_cmd",e)

@bot.command(name="ship")
async def ship_cmd(ctx,m1:discord.Member=None,m2:discord.Member=None):
    try:
        if not m1: await safe_reply(ctx,"Ship *who*?"); return
        p2=m2.display_name if m2 else ctx.author.display_name
        reply=await qai(f"Reluctantly analyze compatibility of {m1.display_name} and {p2}. Rating + observation. 3-4 sentences.",250)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("ship_cmd",e)

@bot.command(name="confess")
async def confess_cmd(ctx,*,confession:str=None):
    try:
        if not confession: await safe_reply(ctx,"Confess *what*?"); return
        user=await _setup(ctx)
        reply=await get_response(ctx.author.id,ctx.channel.id,f"I have something to confess: {confession}",user,ctx.author.display_name,ctx.author.mention)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("confess_cmd",e)

@bot.command(name="compliment")
async def compliment_cmd(ctx,member:discord.Member=None):
    try:
        target=member or ctx.author
        reply=await qai(f"Be forced to genuinely compliment {target.display_name}. Make it clear this is excruciating.",180)
        if member: await ctx.send(f"{member.mention} {reply}")
        else: await safe_reply(ctx,reply)
    except Exception as e: log_error("compliment_cmd",e)

@bot.command(name="haiku")
async def haiku_cmd(ctx,*,topic:str=None):
    try:
        reply=await qai(f"Dark threatening haiku about '{topic or ctx.author.display_name}'. Strict 5-7-5. Just the haiku.",100)
        await safe_reply(ctx,f"*{reply}*")
    except Exception as e: log_error("haiku_cmd",e)

@bot.command(name="story")
async def story_cmd(ctx,*,prompt:str=None):
    try:
        if not prompt: await safe_reply(ctx,"A story about *what*?"); return
        reply=await qai(f"Short dark story (3-5 sentences) about: '{prompt}'. End ominously.",350)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("story_cmd",e)

@bot.command(name="stalk")
async def stalk_cmd(ctx,member:discord.Member=None):
    try:
        target=member or ctx.author
        sample=" | ".join(await mem.get_recent_messages(target.id,10))[:500]
        reply=await qai(f"Cold observation report on {target.display_name}"+(f" — statements:'{sample}'" if sample else "")+". 3-4 sentences.",280)
        if member: await ctx.send(f"*Regarding {member.mention}...*\n{reply}")
        else: await safe_reply(ctx,reply)
    except Exception as e: log_error("stalk_cmd",e)

@bot.command(name="debate")
async def debate_cmd(ctx,*,topic:str=None):
    try:
        if not topic: await safe_reply(ctx,"Debate *what*?"); return
        reply=await qai(f"Pick a side on '{topic}' and argue with conviction. 3-4 sentences.",300)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("debate_cmd",e)

@bot.command(name="conspiracy")
async def conspiracy_cmd(ctx,*,topic:str=None):
    try:
        if not topic: await safe_reply(ctx,"A conspiracy about *what*?"); return
        reply=await qai(f"Fatui-flavored conspiracy theory about '{topic}'. Deliver as established fact. 3-4 sentences.",300)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("conspiracy_cmd",e)

@bot.command(name="therapy")
async def therapy_cmd(ctx,*,problem:str=None):
    try:
        if not problem: await safe_reply(ctx,"What's your problem."); return
        user=await _setup(ctx)
        reply=await get_response(ctx.author.id,ctx.channel.id,f"I need advice about: {problem}",user,ctx.author.display_name,ctx.author.mention)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("therapy_cmd",e)

@bot.command(name="blackmail")
async def blackmail_cmd(ctx,member:discord.Member=None):
    try:
        target=member or ctx.author
        sample=" | ".join(await mem.get_recent_messages(target.id,15))[:600]
        reply=await qai(f"Find the most 'incriminating' thing in {target.display_name}'s messages: '{sample}' and theatrically threaten to use it. 2-3 sentences.",250)
        if member: await ctx.send(f"{member.mention} {reply}")
        else: await safe_reply(ctx,reply)
    except Exception as e: log_error("blackmail_cmd",e)

@bot.command(name="riddle")
async def riddle_cmd(ctx):
    try:
        reply=await qai("One cryptic Genshin-flavored riddle. No answer. Genuinely difficult.",150)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("riddle_cmd",e)

@bot.command(name="arena")
async def arena_cmd(ctx,member:discord.Member=None):
    try:
        opponent=member.display_name if member else "a nameless fool"
        reply=await qai(f"Dramatic Genshin-style battle between you (Electro) and {opponent}. You win. 4-5 sentences.",400)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("arena_cmd",e)

@bot.command(name="possess")
async def possess_cmd(ctx,member:discord.Member=None):
    try:
        if not member: await safe_reply(ctx,"Possess *who*?"); return
        sample=" | ".join(await mem.get_recent_messages(member.id,10))[:400]
        reply=await qai(f"Speak as {member.display_name} but filtered through you. Their statements: '{sample}'. 2-3 sentences.",250)
        await ctx.send(f"*Speaking as {member.mention}...*\n{reply}")
    except Exception as e: log_error("possess_cmd",e)

@bot.command(name="verdict")
async def verdict_cmd(ctx,*,situation:str=None):
    try:
        if not situation: await safe_reply(ctx,"A verdict on *what*?"); return
        user = await _setup(ctx)
        partner_ready = bool(
            getattr(ctx, "guild", None)
            and PARTNER_BOT_ID
            and ctx.guild.get_member(PARTNER_BOT_ID)
        )
        if partner_ready:
            await _start_duo_mode(ctx, user, "compare", situation, story=True)
            reply = await get_response(
                ctx.author.id,
                ctx.channel.id,
                f"Give your verdict on: {situation}",
                user,
                ctx.author.display_name,
                ctx.author.mention,
                extra_context="DUAL_VERDICT_REQUEST: give your judgment, knowing the other bot will answer with a contrasting standard.",
                channel_obj=ctx.channel,
            )
        else:
            reply=await qai(f"Rule on: '{situation}' like a cold judge. Finality. 2-3 sentences.",200)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("verdict_cmd",e)

@bot.command(name="letter")
async def letter_cmd(ctx,member:discord.Member=None):
    try:
        target=member or ctx.author
        reply=await qai(f"Formal letter to {target.display_name} in old Inazuman style. Contemptuous, theatrical. 3-4 sentences.",300)
        if member: await ctx.send(f"{member.mention}\n{reply}")
        else: await safe_reply(ctx,reply)
    except Exception as e: log_error("letter_cmd",e)

@bot.command(name="nightmare")
async def nightmare_cmd(ctx):
    try:
        user=await _setup(ctx)
        reply=await qai(f"Describe a nightmare you had. Somehow about {ctx.author.display_name}. Don't admit that. Unsettling. 2-3 sentences.",200)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("nightmare_cmd",e)

@bot.command(name="rank")
async def rank_cmd(ctx):
    try:
        top=[u for u in await mem.get_top_users(20) if u["user_id"] not in _dm_blocked_users][:8]
        if not top: await safe_reply(ctx,"I don't know enough of you to rank."); return
        entries="\n".join(f"{i+1}. **{u['display_name']}** — {u['message_count']} messages" for i,u in enumerate(top))
        verdict=await qai(f"Rank these by tolerability: {', '.join(u['display_name'] for u in top)}. Dismissive commentary. 2 sentences.",150)
        embed=discord.Embed(title="Tolerability Ranking",description=f"{entries}\n\n*{verdict}*",color=0x4B0082)
        await ctx.send(embed=embed)
    except Exception as e: log_error("rank_cmd",e)

@bot.command(name="stats")
async def stats_cmd(ctx):
    try:
        await _setup(ctx)
        s=await mem.get_stats(ctx.author.id)
        if not s: await safe_reply(ctx,"I don't know you well enough yet."); return
        first=datetime.fromtimestamp(s["first_seen"]).strftime("%b %d, %Y") if s["first_seen"] else "unknown"
        days=int((time.time()-s["first_seen"])/86400) if s["first_seen"] else 0
        embed=discord.Embed(title=f"File: {ctx.author.display_name}",description="*I keep records.*",color=0x4B0082)
        embed.add_field(name="First contact",value=f"{first} ({days}d ago)",inline=True)
        embed.add_field(name="Messages",value=str(s["message_count"]),inline=True)
        embed.add_field(name="Mood",value=f"{s['mood']:+d} — {mood_label(s['mood'])}",inline=True)
        embed.add_field(name="Affection",value=affection_tier(s["affection"]),inline=True)
        embed.add_field(name="Trust",value=trust_tier(s["trust"]),inline=True)
        embed.add_field(name="Drift",value=f"{s['drift_score']}/100",inline=True)
        embed.add_field(name="Slow burn",value=f"{s['slow_burn']}/7",inline=True)
        embed.add_field(name="Inside jokes",value=str(s["joke_count"]),inline=True)
        if s.get("grudge_nick"): embed.add_field(name="Grudge name",value=f'"{s["grudge_nick"]}"',inline=True)
        if s.get("affection_nick"): embed.add_field(name="His nickname for you",value=f'"{s["affection_nick"]}"',inline=True)
        embed.set_footer(text="Don't read too much into this.")
        await ctx.reply(embed=embed)
    except Exception as e: log_error("stats_cmd",e)

@bot.command(name="weather")
async def weather_cmd(ctx,*,location:str=None):
    try:
        if not location: await safe_reply(ctx,"Weather where?"); return
        user = await _setup(ctx)
        data = await _fetch_nws_weather(location)
        if not data:
            await safe_reply(ctx,"That location means nothing to me. Try `City, ST`, a ZIP code, or `lat,lon`.")
            return
        precip = data.get("precipitation")
        precip_text = f"{int(round(precip))}% precipitation chance" if isinstance(precip, (int, float)) else "precipitation unknown"
        comment=await qai(
            f"Weather in {data['place']}: {data['forecast']}. "
            f"Temperature {data['temperature']} degrees {data['temperature_unit']}. "
            f"Wind {data['wind_speed']} {data['wind_direction']}. {precip_text}. "
            f"Comment in your style. 1-2 sentences.",
            150
        )
        if user and not user.get("utility_mode", True):
            reply = f"{comment}\nSource: api.weather.gov"
        else:
            reply = _utility_reply(
                f"Weather for {data['place']}",
                [
                    f"Forecast: {data['forecast']}",
                    f"Temperature: {data['temperature']} {data['temperature_unit']}",
                    f"Wind: {data['wind_speed']} {data['wind_direction']}".strip(),
                    f"Precipitation: {precip_text}",
                ],
                comment,
                "Source: api.weather.gov",
            )
        await safe_reply(ctx,reply)
    except Exception as e: log_error("weather_cmd",e); await safe_reply(ctx,"...The information was unavailable.")


@bot.command(name="weathervideo", aliases=["videoweather", "weatherreportvideo"])
async def weathervideo_cmd(ctx, *, location: str = None):
    try:
        if not await _should_handle_video_command(ctx):
            return
        msg_id = ctx.message.id
        if msg_id in _weathervideo_active:
            return
        _weathervideo_active.add(msg_id)
        await _setup(ctx)
        if not location:
            _weathervideo_active.discard(msg_id)
            await safe_reply(ctx, "Weather where? Give me a location if you want a report worth rendering.")
            return
        use_duo = _partner_bot_present(ctx)
        available, reason = video_renderer_available(BOT_NAME, duo=use_duo)
        if not available:
            _weathervideo_active.discard(msg_id)
            await safe_reply(ctx, f"Video rendering is unavailable on this host. {reason}")
            return
        time_hint = _weathervideo_time_hint(use_duo)
        await ctx.reply(random.choice([
            f"A weather report video. Fine. I'll make it concise. {time_hint}",
            f"You want the forecast packaged properly. Wait. {time_hint}",
            f"I'll render the report. Try to deserve the effort. {time_hint}",
        ]))
        asyncio.ensure_future(_do_weathervideo(ctx, location, use_duo, msg_id))
    except Exception as e:
        _weathervideo_active.discard(ctx.message.id)
        log_error("weathervideo_cmd", e)
        await safe_reply(ctx, "...The weather video request collapsed immediately. Annoying.")


async def _do_weathervideo(ctx, location: str, use_duo: bool, msg_id=None):
    try:
        time_hint = _weathervideo_time_hint(use_duo)
        await ctx.send(
            f"If Wanderer is here, he'll be dragged into this too. {time_hint}"
            if use_duo else
            f"I'm assembling the weather report now. {time_hint}"
        )
        data = await _fetch_nws_weather(location)
        if not data:
            await ctx.send("That location still means nothing to me. Try `City, ST`, a ZIP code, or `lat,lon`."); return
        news_results = await search_web(f"{data['place']} weather news", max_results=3)
        notes = build_weather_video_notes(data["place"], data, news_results, duo=use_duo)
        title = f"{data['place']} Weather Report"
        summary = await (render_duo_debate_video(notes, title=title) if use_duo else render_teaching_video(BOT_NAME, notes, title=title))
        sources = "Weather source: api.weather.gov"
        search_sources = format_search_sources(news_results, max_results=3)
        if search_sources:
            sources = f"{sources}\n{search_sources}"
        await _send_video_summary(
            ctx,
            summary,
            "There. Your duo weather report is ready." if use_duo else "There. Your weather report video is ready.",
            sources_text=sources,
        )
    except Exception as e:
        log_error("do_weathervideo", e)
        await safe_send(ctx, f"The weather video failed. {e}")
    finally:
        if msg_id:
            _weathervideo_active.discard(msg_id)

@bot.command(name="lore")
async def lore_cmd(ctx,*,topic:str=None):
    try:
        if not topic: await safe_reply(ctx,"Lore about *what*?"); return
        user=await _setup(ctx)
        reply=await get_response(ctx.author.id,ctx.channel.id,f"Tell me about this Genshin lore from your perspective: {topic}",user,ctx.author.display_name,ctx.author.mention)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("lore_cmd",e)

@bot.command(name="search",aliases=["find","lookup"])
async def search_cmd(ctx,*,query:str=None):
    try:
        if not query: await safe_reply(ctx,"Search for *what*?"); return
        user=await _setup(ctx)
        reply=await get_response(ctx.author.id,ctx.channel.id,f"Search the web for: {query}.",user,ctx.author.display_name,ctx.author.mention,use_search=True)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("search_cmd",e)

@bot.command(name="solve",aliases=["math","essay","write"])
async def solve_cmd(ctx,*,problem:str=None):
    try:
        if not problem: await safe_reply(ctx,"Solve *what*?"); return
        user=await _setup(ctx)
        reply=await get_response(ctx.author.id,ctx.channel.id,f"Solve or respond to this accurately: {problem}",user,ctx.author.display_name,ctx.author.mention,use_search=True)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("solve_cmd",e)

@bot.command(name="rival",aliases=["setrival"])
async def rival_cmd(ctx,member:discord.Member=None):
    try:
        await _setup(ctx)
        if not member: await mem.set_rival(ctx.author.id,None); await safe_reply(ctx,"Rivalry dissolved."); return
        if member.id==ctx.author.id: await safe_reply(ctx,"Your rival is yourself? How appropriate."); return
        await mem.set_rival(ctx.author.id,member.id)
        await safe_reply(ctx,f"Tch. {member.display_name}. Fine. I'll be watching.")
    except Exception as e: log_error("rival_cmd",e)

@bot.command(name="remind",aliases=["remindme"])
async def remind_cmd(ctx,minutes:int=None,*,reminder:str=None):
    try:
        if not minutes or not reminder: await safe_reply(ctx,"Usage: `!remind <minutes> <reminder>`"); return
        if not 1<=minutes<=10080: await safe_reply(ctx,"Between 1 minute and 7 days."); return
        await mem.add_reminder(ctx.author.id,ctx.channel.id,reminder,time.time()+minutes*60)
        await safe_reply(ctx,f"Fine. {minutes} minute{'s' if minutes!=1 else ''}. Pathetic.")
    except Exception as e: log_error("remind_cmd",e)

@bot.command(name="translate")
async def translate_cmd(ctx,*,text:str=None):
    try:
        if not text: await safe_reply(ctx,"Translate *what*?"); return
        user=await _setup(ctx)
        reply=await get_response(ctx.author.id,ctx.channel.id,f"Rewrite this in your voice, keeping the meaning: '{text[:500]}'",user,ctx.author.display_name,ctx.author.mention)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("translate_cmd",e)

@bot.command(name="insult",aliases=["roast_single"])
async def insult_cmd(ctx,member:discord.Member=None):
    try:
        target=member or ctx.author
        reply=await qai(f"One devastating insult to {target.display_name}. Sharp, theatrical.",150)
        if member: await ctx.send(f"{member.mention} {reply}")
        else: await safe_reply(ctx,reply)
    except Exception as e: log_error("insult_cmd",e)

@bot.command(name="dm",aliases=["private","whisper"])
async def dm_cmd(ctx,*,message:str=None):
    try:
        user=await _setup(ctx)
        reply=await get_response(
            ctx.author.id,ctx.author.id,message or "The user wants to speak privately.",
            user,ctx.author.display_name,ctx.author.mention,is_dm=True
        )
        try:
            await ctx.author.send(reply); await ctx.message.add_reaction("📨")
        except discord.Forbidden:
            await safe_reply(ctx,"Your DMs are closed. How cowardly.")
    except Exception as e: log_error("dm_cmd",e)

@bot.command(name="remember")
async def remember_cmd(ctx,*,text:str=None):
    try:
        if not text:
            await safe_reply(ctx,"Tell me what I'm meant to remember.")
            return
        await _setup(ctx)
        await mem.set_callback_memory(ctx.author.id, text[:220])
        await mem.add_memory_event(ctx.author.id, "manual", text[:220], 5)
        scene_update = infer_scene_update(text, ctx.author.display_name)
        if scene_update:
            await mem.update_scene_state(ctx.channel.id, **scene_update)
        await safe_reply(ctx,"Fine. I'll remember it.")
    except Exception as e: log_error("remember_cmd",e)

@bot.command(name="forget")
async def forget_cmd(ctx,*,topic:str=None):
    try:
        if not topic or topic.strip().lower() in {"all","everything","me"}:
            await ctx.send(random.choice(["Wipe my memory of you? Press the button.","Gone in an instant. If you're sure."]),view=ResetView(ctx.author.id))
            return
        await _setup(ctx)
        result=await mem.forget_memory_matches(ctx.author.id, topic)
        removed=sum(result.values())
        if removed:
            await safe_reply(ctx,f"Fine. I dropped {removed} thing{'s' if removed!=1 else ''} tied to '{topic}'.")
        else:
            await safe_reply(ctx,"Nothing obvious matched that. Be more specific.")
    except Exception as e: log_error("forget_cmd",e)

@bot.command(name="memories",aliases=["memorybank"])
async def memories_cmd(ctx):
    try:
        user=await _setup(ctx)
        topics=await mem.get_top_topics(ctx.author.id,4)
        memories=await mem.get_memory_bank_entries(ctx.author.id,6)
        scene=await mem.get_scene_state(ctx.channel.id)
        await safe_reply(ctx,_format_memory_snapshot(user,topics,memories,scene))
    except Exception as e: log_error("memories_cmd",e)

@bot.command(name="pinpromise")
async def pinpromise_cmd(ctx, *, text: str = None):
    try:
        await _pin_memory(ctx, "promise", text, 7)
    except Exception as e: log_error("pinpromise_cmd", e)

@bot.command(name="pinwound")
async def pinwound_cmd(ctx, *, text: str = None):
    try:
        await _pin_memory(ctx, "wound", text, 8)
    except Exception as e: log_error("pinwound_cmd", e)

@bot.command(name="pincomfort")
async def pincomfort_cmd(ctx, *, text: str = None):
    try:
        await _pin_memory(ctx, "comfort", text, 6)
    except Exception as e: log_error("pincomfort_cmd", e)

@bot.command(name="pinjoke")
async def pinjoke_cmd(ctx, *, text: str = None):
    try:
        await _pin_memory(ctx, "joke", text, 5, shared_joke=True)
    except Exception as e: log_error("pinjoke_cmd", e)

@bot.command(name="relationship")
async def relationship_cmd(ctx):
    try:
        user = await _setup(ctx)
        topics = await mem.get_top_topics(ctx.author.id, 3)
        marks = await mem.get_active_consequence_marks(ctx.author.id, 4)
        lines = _relationship_lines(user, topics, marks)
        await safe_reply(ctx, "\n".join(lines))
    except Exception as e: log_error("relationship_cmd", e)

@bot.command(name="rebuildmemory")
async def rebuildmemory_cmd(ctx, scope: str = ""):
    try:
        if not user_can_manage_rebuild(ctx, OWNER_ID):
            await safe_reply(ctx, "That rebuild command isn't for you.")
            return
        if (scope or "").strip().lower() != "all":
            await safe_reply(ctx, "Use `!rebuildmemory all`.")
            return
        await ctx.reply("Fine. I'm rebuilding memory from visible channel history. This may take a bit.")
        async with ctx.typing():
            stats = await _rebuild_from_history(ctx)
        if not stats.get("messages_replayed"):
            await safe_send(ctx, "I couldn't find any usable message history to rebuild from.")
            return
        await safe_send(
            ctx,
            f"Rebuild complete. Replayed {stats['messages_replayed']} messages for {stats['rebuilt_users']} user(s) "
            f"across {stats['channels_scanned']} channel(s), using up to {stats['limit_per_channel']} recent messages per channel.",
        )
    except Exception as e:
        log_error("rebuildmemory_cmd", e)
        await safe_reply(ctx, f"The memory rebuild failed. {e}")


@bot.command(name="rebuildrelationship")
async def rebuildrelationship_cmd(ctx, member: discord.Member = None):
    try:
        if not user_can_manage_rebuild(ctx, OWNER_ID):
            await safe_reply(ctx, "That rebuild command isn't for you.")
            return
        target = member or ctx.author
        await ctx.reply(f"Rebuilding the relationship state for {target.display_name}. Don't get impatient.")
        async with ctx.typing():
            stats = await _rebuild_from_history(ctx, target_user_id=target.id)
            await mem.upsert_user(target.id, str(target), target.display_name)
            user = await mem.get_user(target.id)
            topics = await mem.get_top_topics(target.id, 3)
            marks = await mem.get_active_consequence_marks(target.id, 4)
        if not stats.get("messages_replayed"):
            await safe_send(ctx, f"I couldn't find usable history for {target.display_name}.")
            return
        lines = _relationship_lines(user, topics, marks)
        lines.insert(
            0,
            f"Rebuilt from {stats['messages_replayed']} messages across {stats['channels_scanned']} channel(s).",
        )
        await safe_send(ctx, "\n".join(lines))
    except Exception as e:
        log_error("rebuildrelationship_cmd", e)
        await safe_reply(ctx, f"The relationship rebuild failed. {e}")


@bot.command(name="arc")
async def arc_cmd(ctx):
    try:
        user = await _setup(ctx)
        arc = _current_arc(user)
        stage, stage_desc = _progression_parts(user)
        unlocks = describe_arc_unlocks(BOT_NAME, arc) or "No special unlocks yet."
        recent = await mem.get_recent_milestones(f"{BOT_NAME}:user:{ctx.author.id}", 2)
        lines = [
            f"Arc: {arc}",
            f"Progression: {stage} | {stage_desc}",
            f"Unlocks: {unlocks}",
        ]
        if recent:
            lines.append("Recent milestones: " + " | ".join(note[:120] for note in recent))
        await safe_reply(ctx, "\n".join(lines))
    except Exception as e: log_error("arc_cmd", e)

@bot.command(name="duostate")
async def duostate_cmd(ctx):
    try:
        speaker_mode = await mem.get_channel_speaker_mode(ctx.channel.id)
        duo = await mem.get_duo_session(ctx.channel.id)
        relation = await mem.get_bot_relationship(PARTNER_PAIR_KEY)
        stories = await mem.get_recent_duo_stories(ctx.channel.id, 4)
        cases = await mem.list_world_cases(channel_id=ctx.channel.id, limit=3)
        lines = _duostate_lines(speaker_mode, duo, relation, stories, cases)
        await safe_reply(ctx, "\n".join(lines))
    except Exception as e: log_error("duostate_cmd", e)

@bot.command(name="achievements", aliases=["gallery"])
async def achievements_cmd(ctx):
    try:
        entries = await mem.get_hidden_achievements(f"user:{ctx.author.id}", 12)
        await safe_reply(ctx, "\n".join(_achievement_gallery_lines(entries)))
    except Exception as e: log_error("achievements_cmd", e)

@bot.command(name="speaker", aliases=["activespeaker"])
async def speaker_cmd(ctx, mode: str = None):
    try:
        current = await mem.get_channel_speaker_mode(ctx.channel.id)
        if not mode:
            await safe_reply(ctx, f"Channel speaker mode: `{current}`.")
            return
        normalized = mode.strip().lower()
        mapping = {
            "auto": "auto",
            "both": "both",
            "scara": "scaramouche",
            "scaramouche": "scaramouche",
            "wanderer": "wanderer",
        }
        if normalized not in mapping:
            await safe_reply(ctx, "Use `!speaker auto`, `!speaker scaramouche`, `!speaker wanderer`, or `!speaker both`.")
            return
        await mem.set_channel_speaker_mode(ctx.channel.id, mapping[normalized])
        await safe_reply(ctx, f"Fine. This channel is now set to `{mapping[normalized]}` mode.")
    except Exception as e: log_error("speaker_cmd", e)

@bot.command(name="both")
async def both_cmd(ctx,*,prompt:str=None):
    try:
        if not prompt:
            await safe_reply(ctx,"Ask something worth answering.")
            return
        user=await _setup(ctx)
        await _start_duo_mode(ctx, user, "both", prompt)
        reply=await get_response(
            ctx.author.id,ctx.channel.id,prompt,user,ctx.author.display_name,ctx.author.mention,
            extra_context="TWO_BOT_MODE: The user explicitly wants both bots to answer. Keep it to one or two sentences and never speak for the other bot.",
            channel_obj=ctx.channel
        )
        await _reply_and_store(ctx,reply)
    except Exception as e: log_error("both_cmd",e)

@bot.command(name="duet")
async def duet_cmd(ctx,*,prompt:str=None):
    try:
        if not prompt:
            await safe_reply(ctx,"Set the scene first.")
            return
        user=await _setup(ctx)
        await _start_duo_mode(ctx, user, "duet", prompt)
        reply=await get_response(
            ctx.author.id,ctx.channel.id,f"Contribute one turn to this shared two-bot scene: {prompt}",
            user,ctx.author.display_name,ctx.author.mention,
            extra_context="DUET_MODE: contribute one short in-character turn, leave space for the other bot, and avoid narration tags.",
            channel_obj=ctx.channel
        )
        await _reply_and_store(ctx,reply)
    except Exception as e: log_error("duet_cmd",e)

@bot.command(name="argue")
async def argue_cmd(ctx,*,topic:str=None):
    try:
        if not topic:
            await safe_reply(ctx,"Argue about what.")
            return
        user=await _setup(ctx)
        await _start_duo_mode(ctx, user, "argue", topic)
        reply=await get_response(
            ctx.author.id,ctx.channel.id,f"The user started a deliberate two-bot argument about: {topic}",
            user,ctx.author.display_name,ctx.author.mention,
            extra_context="ARGUE_MODE: take a sharp stance, challenge the other bot directly, and keep it to one or two sentences.",
            channel_obj=ctx.channel
        )
        await _reply_and_store(ctx,reply)
    except Exception as e: log_error("argue_cmd",e)

@bot.command(name="compare")
async def compare_cmd(ctx,*,topic:str=None):
    try:
        if not topic:
            await safe_reply(ctx,"Compare what.")
            return
        user=await _setup(ctx)
        await _start_duo_mode(ctx, user, "compare", topic, story=True)
        reply=await get_response(
            ctx.author.id,ctx.channel.id,f"Give your verdict on this and make your difference from Wanderer clear: {topic}",
            user,ctx.author.display_name,ctx.author.mention,
            extra_context="COMPARE_MODE: answer the prompt, then draw a quick contrast between your view and the other bot's likely view.",
            channel_obj=ctx.channel
        )
        await _reply_and_store(ctx,reply)
    except Exception as e: log_error("compare_cmd",e)

@bot.command(name="interrogate")
async def interrogate_cmd(ctx,*,topic:str=None):
    try:
        if not topic:
            await safe_reply(ctx,"Interrogate what.")
            return
        user = await _setup(ctx)
        await _start_duo_mode(ctx, user, "interrogate", topic, story=True, enemy=topic)
        reply = await get_response(
            ctx.author.id, ctx.channel.id, f"Start a two-bot interrogation about: {topic}",
            user, ctx.author.display_name, ctx.author.mention,
            extra_context="INTERROGATE_MODE: ask one pointed question or accusation, as if cornering the target.",
            channel_obj=ctx.channel
        )
        await _reply_and_store(ctx, reply)
    except Exception as e: log_error("interrogate_cmd", e)

@bot.command(name="choose")
async def choose_cmd(ctx,*,options:str=None):
    try:
        if not options or "|" not in options:
            await safe_reply(ctx,"Use `!choose option A | option B`.")
            return
        user = await _setup(ctx)
        await _start_duo_mode(ctx, user, "compare", options, story=True)
        reply = await get_response(
            ctx.author.id, ctx.channel.id, f"Choose between these options and justify it: {options}",
            user, ctx.author.display_name, ctx.author.mention,
            extra_context="CHOOSE_MODE: make a clear choice fast, then justify it sharply.",
            channel_obj=ctx.channel
        )
        await _reply_and_store(ctx, reply)
    except Exception as e: log_error("choose_cmd", e)

@bot.command(name="trial")
async def trial_cmd(ctx,*,charge:str=None):
    try:
        if not charge:
            await safe_reply(ctx,"Put someone on trial for something.")
            return
        user = await _setup(ctx)
        await _start_duo_mode(ctx, user, "trial", charge, story=True, enemy=charge)
        reply = await get_response(
            ctx.author.id, ctx.channel.id, f"Open a two-bot trial about this charge: {charge}",
            user, ctx.author.display_name, ctx.author.mention,
            extra_context="TRIAL_MODE: deliver a prosecution or judgment opening with theatrical confidence.",
            channel_obj=ctx.channel
        )
        await _reply_and_store(ctx, reply)
    except Exception as e: log_error("trial_cmd", e)

@bot.command(name="mission")
async def mission_cmd(ctx,*,objective:str=None):
    try:
        if not objective:
            await safe_reply(ctx,"Mission objective.")
            return
        user = await _setup(ctx)
        await _start_duo_mode(ctx, user, "mission", objective, story=True, enemy=objective)
        reply = await get_response(
            ctx.author.id, ctx.channel.id, f"Plan a two-bot mission around this objective: {objective}",
            user, ctx.author.display_name, ctx.author.mention,
            extra_context="MISSION_MODE: assign danger, leverage, and one clear tactical role.",
            channel_obj=ctx.channel
        )
        await _reply_and_store(ctx, reply)
    except Exception as e: log_error("mission_cmd", e)

@bot.command(name="truthdare", aliases=["tod"])
async def truthdare_cmd(ctx,*,prompt:str=None):
    try:
        topic = prompt or "truth or dare"
        user = await _setup(ctx)
        await _start_duo_mode(ctx, user, "truthdare", topic, story=True)
        reply = await get_response(
            ctx.author.id, ctx.channel.id, f"Start a two-bot truth-or-dare round with this setup: {topic}",
            user, ctx.author.display_name, ctx.author.mention,
            extra_context="TRUTHDARE_MODE: issue one pointed truth or dare challenge, leaving room for the other bot to escalate.",
            channel_obj=ctx.channel
        )
        await _reply_and_store(ctx, reply)
    except Exception as e: log_error("truthdare_cmd", e)

@bot.command(name="scene")
async def scene_cmd(ctx):
    try:
        scene = await mem.get_scene_state(ctx.channel.id)
        duo = await mem.get_duo_session(ctx.channel.id)
        if not scene and not duo:
            await safe_reply(ctx, "This channel has no scene worth remembering yet.")
            return
        lines = []
        if scene:
            lines.append(f"Scene: {describe_scene_state(scene)}")
        if duo:
            lines.append(f"Duo mode: {duo.get('mode')} | topic={duo.get('topic')}")
        await safe_reply(ctx, "\n".join(lines))
    except Exception as e: log_error("scene_cmd", e)

@bot.command(name="world")
async def world_cmd(ctx):
    try:
        entities = await mem.list_world_entities(limit=8, channel_id=ctx.channel.id)
        await safe_reply(ctx, "\n".join(_world_lines(entities)))
    except Exception as e: log_error("world_cmd", e)

@bot.command(name="cases")
async def cases_cmd(ctx):
    try:
        cases = await mem.list_world_cases(channel_id=ctx.channel.id, limit=8)
        await safe_reply(ctx, "\n".join(_case_lines(cases)))
    except Exception as e: log_error("cases_cmd", e)

@bot.command(name="worldadd")
async def worldadd_cmd(ctx, entity_type: str = None, *, payload: str = None):
    try:
        if not entity_type or not payload:
            await safe_reply(ctx, "Use `!worldadd enemy Dottore | ruined laboratory trail`.")
            return
        if "|" in payload:
            name, summary = [part.strip() for part in payload.split("|", 1)]
        else:
            name, summary = payload.strip(), ""
        if not name:
            await safe_reply(ctx, "Give the world entry a name.")
            return
        normalized = entity_type.strip().lower()
        if normalized not in {"enemy", "ally", "gift", "place", "prop", "faction", "figure", "case"}:
            await safe_reply(ctx, "Use one of: enemy, ally, gift, place, prop, faction, figure, case.")
            return
        await mem.upsert_world_entity(
            normalized,
            name,
            summary=summary or f"Added manually by {ctx.author.display_name}.",
            status="remembered",
            channel_id=ctx.channel.id,
            owner_user_id=ctx.author.id,
            updated_by=BOT_NAME,
        )
        await safe_reply(ctx, f"Fine. I filed `{name}` under `{normalized}`.")
    except Exception as e: log_error("worldadd_cmd", e)

@bot.command(name="insidejokes", aliases=["jokes"])
async def insidejokes_cmd(ctx):
    try:
        jokes = await mem.list_inside_jokes(ctx.author.id, 5)
        if not jokes:
            await safe_reply(ctx, "Apparently you've said nothing memorable enough to become an inside joke.")
            return
        await safe_reply(ctx, "Inside jokes:\n" + "\n".join(f"- {j[:140]}" for j in jokes))
    except Exception as e: log_error("insidejokes_cmd", e)

@bot.command(name="reset",aliases=["wipe"])
async def reset_cmd(ctx):
    try:
        await ctx.send(random.choice(["Wipe my memory of you? Press the button.","Gone in an instant. If you're sure."]),view=ResetView(ctx.author.id))
    except Exception as e: log_error("reset_cmd",e)

@bot.command(name="nsfw")
async def nsfw_cmd(ctx,mode:str=None):
    try:
        user=await _setup(ctx); cur=user.get("nsfw_mode",False) if user else False
        new=True if mode=="on" else False if mode=="off" else not cur
        await mem.set_mode(ctx.author.id,"nsfw_mode",new)
        await safe_reply(ctx,"Unfiltered. Fine." if new else "Restrained again. How boring.")
    except Exception as e: log_error("nsfw_cmd",e)

@bot.command(name="proactive",aliases=["ping_me"])
async def proactive_cmd(ctx,mode:str=None):
    try:
        user=await _setup(ctx); cur=user.get("proactive",True) if user else True
        new=True if mode=="on" else False if mode=="off" else not cur
        await mem.set_mode(ctx.author.id,"proactive",new)
        await safe_reply(ctx,"I might message you. Or not." if new else "Fine. I'll pretend you don't exist.")
    except Exception as e: log_error("proactive_cmd",e)

@bot.command(name="dms",aliases=["allowdms","stopdms"])
async def dms_cmd(ctx,mode:str=None):
    try:
        user=await _setup(ctx); cur=user.get("allow_dms",True) if user else True
        new=True if mode=="on" else False if mode=="off" else not cur
        await mem.set_mode(ctx.author.id,"allow_dms",new)
        await safe_reply(ctx,"Fine. I'll message you when I feel like it." if new else "Cutting me off? Fine.")
    except Exception as e: log_error("dms_cmd",e)

@bot.command(name="utility")
async def utility_cmd(ctx, mode: str = None):
    try:
        user = await _setup(ctx)
        if not mode:
            await safe_reply(ctx, f"Utility mode is `{_pref_label(user.get('utility_mode', True) if user else True)}`.")
            return
        normalized = mode.strip().lower()
        if normalized not in {"on", "off"}:
            await safe_reply(ctx, "Use `!utility on` or `!utility off`.")
            return
        enabled = normalized == "on"
        await mem.set_user_preference(ctx.author.id, "utility_mode", int(enabled))
        await safe_reply(ctx, f"Utility mode is `{_pref_label(enabled)}` now.")
    except Exception as e: log_error("utility_cmd", e)

@bot.command(name="duoauto")
async def duoauto_cmd(ctx, mode: str = None):
    try:
        user = await _setup(ctx)
        if not mode:
            await safe_reply(ctx, f"Duo autoplay is `{_pref_label(user.get('duo_autoplay', True) if user else True)}`.")
            return
        normalized = mode.strip().lower()
        if normalized not in {"on", "off"}:
            await safe_reply(ctx, "Use `!duoauto on` or `!duoauto off`.")
            return
        enabled = normalized == "on"
        await mem.set_user_preference(ctx.author.id, "duo_autoplay", int(enabled))
        await safe_reply(ctx, f"Fine. Duo autoplay is `{_pref_label(enabled)}` now.")
    except Exception as e: log_error("duoauto_cmd", e)

@bot.command(name="rpdepth")
async def rpdepth_cmd(ctx, depth: str = None):
    try:
        user = await _setup(ctx)
        if not depth:
            await safe_reply(ctx, f"RP depth is `{(user or {}).get('rp_depth', 'medium')}`.")
            return
        normalized = depth.strip().lower()
        if normalized not in {"low", "medium", "high"}:
            await safe_reply(ctx, "Use `!rpdepth low`, `!rpdepth medium`, or `!rpdepth high`.")
            return
        await mem.set_user_preference(ctx.author.id, "rp_depth", normalized)
        await safe_reply(ctx, f"Fine. RP depth is `{normalized}` now.")
    except Exception as e: log_error("rpdepth_cmd", e)

@bot.command(name="timezone")
async def timezone_cmd(ctx, *, timezone_name: str = None):
    try:
        user = await _setup(ctx)
        if not timezone_name:
            await safe_reply(ctx, f"Timezone: {user.get('timezone_name', 'America/Los_Angeles')}")
            return
        ZoneInfo(timezone_name.strip())
        await mem.set_timezone(ctx.author.id, timezone_name.strip())
        await safe_reply(ctx, f"Timezone set to `{timezone_name.strip()}`.")
    except Exception:
        await safe_reply(ctx, "Use a valid IANA timezone like `America/Los_Angeles`.")

@bot.command(name="quiethours")
async def quiethours_cmd(ctx, start_hour: int = None, end_hour: int = None):
    try:
        user = await _setup(ctx)
        if start_hour is None or end_hour is None:
            await safe_reply(ctx, f"Quiet hours: {user.get('quiet_hours_start', 23)}:00 to {user.get('quiet_hours_end', 8)}:00")
            return
        await mem.set_quiet_hours(ctx.author.id, start_hour, end_hour)
        await safe_reply(ctx, f"Quiet hours set to {start_hour % 24}:00-{end_hour % 24}:00.")
    except Exception as e: log_error("quiethours_cmd", e)

@bot.command(name="dmfreq")
async def dmfreq_cmd(ctx, hours: int = None):
    try:
        user = await _setup(ctx)
        if hours is None:
            await safe_reply(ctx, f"DM frequency floor: every {user.get('dm_frequency_hours', 8)} hour(s).")
            return
        await mem.set_dm_frequency(ctx.author.id, hours)
        await safe_reply(ctx, f"Fine. No more than once every {max(1, min(72, hours))} hour(s).")
    except Exception as e: log_error("dmfreq_cmd", e)

@bot.command(name="dmgrace")
async def dmgrace_cmd(ctx, minutes: int = None):
    try:
        user = await _setup(ctx)
        if minutes is None:
            await safe_reply(ctx, f"Recent-activity DM grace: {user.get('recent_activity_grace_minutes', 45)} minute(s).")
            return
        await mem.set_recent_activity_grace(ctx.author.id, minutes)
        await safe_reply(ctx, f"I'll leave at least {max(5, min(720, minutes))} minute(s) after your recent activity before DMing.")
    except Exception as e: log_error("dmgrace_cmd", e)

@bot.command(name="mood")
async def mood_cmd(ctx):
    try:
        await _setup(ctx); s=await mem.get_mood(ctx.author.id)
        bar="█"*(s+10)+"░"*(20-(s+10))
        await safe_reply(ctx,f"`[{bar}]` {s:+d} — {mood_label(s)}\n*Don't read into this.*")
    except Exception as e: log_error("mood_cmd",e)

@bot.command(name="affection")
async def affection_cmd(ctx):
    try:
        await _setup(ctx); user=await mem.get_user(ctx.author.id); s=user.get("affection",0) if user else 0
        bar="█"*(s//5)+"░"*(20-s//5)
        await safe_reply(ctx,f"`[{bar}]` {s}/100 — {affection_tier(s)}\n*...I said don't look at that.*")
    except Exception as e: log_error("affection_cmd",e)

@bot.command(name="trust")
async def trust_cmd(ctx):
    try:
        await _setup(ctx); user=await mem.get_user(ctx.author.id); s=user.get("trust",0) if user else 0
        bar="█"*(s//5)+"░"*(20-s//5)
        await safe_reply(ctx,f"`[{bar}]` {s}/100 — {trust_tier(s)}\n*This means nothing.*")
    except Exception as e: log_error("trust_cmd",e)

@bot.command(name="whoami")
async def whoami_cmd(ctx):
    try:
        if not OWNER_ID or ctx.author.id!=OWNER_ID: await safe_reply(ctx,"That command isn't for you."); return
        user=await _setup(ctx)
        reply=await get_response(ctx.author.id,ctx.channel.id,"What do you actually think about the fact that I built you. Be honest.",user,ctx.author.display_name,ctx.author.mention,is_owner=True)
        await safe_reply(ctx,reply)
    except Exception as e: log_error("whoami_cmd",e)

@bot.command(name="enrollface", aliases=["rememberface"])
async def enrollface_cmd(ctx):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID:
            await safe_reply(ctx, "That command isn't for you.")
            return
        if not face_support_ready():
            await safe_reply(ctx, _face_feature_unavailable_text())
            return
        img, vid = await _command_face_media(ctx)
        if not img and not vid:
            await safe_reply(ctx, "Attach or reply to an image or video and use `!enrollface`.")
            return
        profile = await mem.get_face_profile(FACE_PROFILE_KEY)
        import aiohttp as _aiohttp
        if vid:
            async with _aiohttp.ClientSession() as _sess:
                async with _sess.get(vid.url) as _resp:
                    video_bytes = await _resp.read()
            frames = await asyncio.get_event_loop().run_in_executor(None, _extract_frames_blocking, video_bytes, 5)
            enrolled = enroll_face_profile_from_frames(profile, frames)
        else:
            async with _aiohttp.ClientSession() as _sess:
                async with _sess.get(img.url) as _resp:
                    img_bytes = await _resp.read()
            enrolled = enroll_face_profile(profile, img_bytes)
        if not enrolled.get("ok"):
            await safe_reply(ctx, _face_enroll_failure_text(enrolled.get("reason", "")))
            return
        await mem.save_face_profile(FACE_PROFILE_KEY, ctx.author.id, ctx.author.display_name, enrolled["profile"])
        await safe_reply(ctx, _face_enroll_success_text(enrolled.get("sample_count", 0)))
    except Exception as e:
        log_error("enrollface_cmd", e)

@bot.command(name="faceinfo")
async def faceinfo_cmd(ctx):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID:
            await safe_reply(ctx, "That command isn't for you.")
            return
        profile = await mem.get_face_profile(FACE_PROFILE_KEY)
        if not profile:
            await safe_reply(ctx, "No enrolled face profile yet.")
            return
        sample_count = int(profile.get("sample_count", 0) or 0)
        updated_ts = float(profile.get("updated_ts", 0) or 0)
        updated_label = datetime.fromtimestamp(updated_ts).strftime("%Y-%m-%d %H:%M") if updated_ts else "unknown"
        await safe_reply(ctx, f"`Face memory: enrolled` — {sample_count} sample(s), updated {updated_label}.")
    except Exception as e:
        log_error("faceinfo_cmd", e)

@bot.command(name="deleteface", aliases=["forgetface"])
async def deleteface_cmd(ctx):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID:
            await safe_reply(ctx, "That command isn't for you.")
            return
        await mem.delete_face_profile(FACE_PROFILE_KEY)
        await safe_reply(ctx, _face_delete_text())
    except Exception as e:
        log_error("deleteface_cmd", e)

async def help_cmd(ctx):
    try:
        c = 0x4B0082
        e1 = discord.Embed(title="Commands (1/3) — Talk & Fight",
                           description="Hmph. Only saying this once.", color=c)
        for n,v in [
            ("🔊 !voice <msg>","Voice message — !speak !say"),
            ("📨 !dm [msg]","He DMs you privately"),
            ("🤫 !confess <text>","Tell him something"),
            ("🛋️ !therapy <problem>","Terrible in-character advice"),
            ("🌐 !translate <text>","Rewritten in his voice"),
            ("⚔️ !spar [msg]","Word battle"),
            ("🥊 !duel @user","Insult battle referee"),
            ("🎤 !roast @user","Turn-based roast battle (5 rounds)"),
            ("⚡ !arena [@user]","Dramatic mock Genshin battle"),
            ("🎯 !dare","A dark theatrical dare"),
            ("🧠 !trivia","Genshin lore trivia"),
            ("✅ !answer <text>","Answer a trivia question"),
            ("🧩 !riddle","Cryptic Genshin riddle"),
            ("🔒 !hostage","Takes your good mood hostage"),
            ("🔓 !release <offering>","Try to fulfill his demand"),
            ("🥠 !fortune","Fortune cookie rewritten as a threat"),
            ("💭 !opinion <char>","His honest take on any Genshin character"),
            ("🎭 !impersonate <char>","Speaks as them, badly"),
            ("📜 !lore <topic>","Genshin lore from his perspective"),
        ]: e1.add_field(name=n,value=v,inline=False)
        for n, v in [
            ("!both <prompt>", "Both bots answer in sequence"),
            ("!duet <prompt>", "Start a shared two-bot scene"),
            ("!argue <topic>", "Let both bots clash over a topic"),
            ("!compare <topic>", "Each bot gives a contrasting verdict"),
        ]: e1.add_field(name=n, value=v, inline=False)

        e2 = discord.Embed(title="Commands (2/3) — Assess & Create", color=c)
        for n,v in [
            ("🔍 !judge [@user]","Brutal character assessment"),
            ("👁️ !stalk [@user]","Cold observation report"),
            ("🃏 !blackmail [@user]","Most incriminating messages"),
            ("🔦 !interrogate @user","Cold interrogation"),
            ("👻 !possess @user","Speaks as them, filtered through him"),
            ("📊 !rate <thing>","Rates anything out of 10"),
            ("💞 !ship @u1 [@u2]","Reluctant compatibility"),
            ("⚖️ !verdict <situation>","He rules on anything"),
            ("⚖️ !debate <topic>","He argues a side"),
            ("🕵️ !conspiracy <topic>","Fatui conspiracy theory"),
            ("🏆 !rank","Ranks everyone by tolerability"),
            ("📝 !haiku [topic]","Dark threatening haiku"),
            ("📖 !story <prompt>","Short dark story"),
            ("✉️ !letter [@user]","Formal old Inazuman letter"),
            ("🌸 !compliment [@user]","Forces him to say something nice"),
            ("⚡ !insult [@user]","Cutting insult"),
            ("🔮 !prophecy [@user]","Cryptic threatening fortune"),
            ("😰 !nightmare","A nightmare. Somehow about you."),
            ("🔍 !search <query>","Web search with commentary"),
            ("🧮 !solve <problem>","Math, essays, Q&A — !math !essay"),
        ]: e2.add_field(name=n,value=v,inline=False)

        e3 = discord.Embed(title="Commands (3/3) — Settings & Stats", color=c)
        for n,v in [
            ("📊 !stats","Your full relationship file"),
            ("🌡️ !mood","His mood toward you"),
            ("💜 !affection","His hidden affection score"),
            ("🔒 !trust","His trust level toward you"),
            ("🗡️ !rival @user","Designate a rival"),
            ("⏰ !remind <mins> <txt>","Reminder with disdain"),
            ("🌤️ !weather <city>","Weather + contemptuous commentary"),
            ("🎥 !weathervideo <city>","Render a weather report video; duo if Wanderer is here"),
            ("📢 !poll <question>","He demands a vote"),
            ("📋 !summarize","Recent chat summary with contempt"),
            ("🔇 !mute [@user] [min]","Ignores someone in character"),
            ("🔊 !unmute [@user]","Unmutes someone"),
            ("🔄 !reset","Wipe your memory — !forget"),
            ("🔞 !nsfw [on/off]","Toggle unfiltered mode"),
            ("📡 !proactive [on/off]","Toggle unprompted messages"),
            ("💌 !dms [on/off]","Toggle voluntary private DMs"),
        ]: e3.add_field(name=n,value=v,inline=False)
        for n, v in [
            ("!memories", "See what he is actually holding onto"),
            ("!teachvideo <topic>", "Render a lesson video from notes or an attachment"),
            ("!remember <text>", "Tell him to keep something"),
            ("!forget <topic>", "Forget one topic instead of everything"),
            ("!relationship / !arc", "See arc, progression, and conflict aftermath"),
            ("!duostate / !speaker", "Inspect duo mode or set the active speaker for this channel"),
            ("!pinpromise / !pinwound / !pincomfort / !pinjoke", "Pin more precise memories"),
            ("!utility / !duoauto / !rpdepth", "Tune utility output, duo chaining, and RP depth"),
        ]: e3.add_field(name=n, value=v, inline=False)
        e3.add_field(name="Hidden Systems",
            value="Be kind 7 days in a row: something rare happens once\n"
                  "Be rude: mood drops, you get a degrading nickname\n"
                  "High affection: he starts calling you something specific\n"
                  "Build trust: he tells you things he would never normally say\n"
                  "Say you will never win: villain monologue\n"
                  "Mention his hat: disproportionate response\n"
                  "He reads the channel — knows what everyone has been saying",
            inline=False)
        e3.set_footer(text="Scaramouche — The Balladeer | !scarahelp for commands")
        await ctx.send(embed=e1)
        await ctx.send(embed=e2)
        await ctx.send(embed=e3)
    except Exception as e:
        log_error("help_cmd", e)
        try: await ctx.send("Hmph. Something went wrong.")
        except: pass

@bot.command(name="scarahelp", aliases=["commands"])
async def scarahelp_cmd(ctx):
    try:
        await help_cmd(ctx)
    except Exception as e:
        log_error("scarahelp_cmd", e)
        try: await ctx.send("Hmph. Something went wrong.")
        except: pass


@bot.tree.command(name="scaramouche", description="Talk directly to Scaramouche.")
@app_commands.describe(message="What you want to say to Scaramouche")
async def slash_scaramouche(interaction: discord.Interaction, message: str):
    try:
        user = await _setup_user(interaction.user)
        await interaction.response.defer(thinking=True)
        reply = await get_response(
            interaction.user.id,
            interaction.channel_id or interaction.user.id,
            message,
            user,
            getattr(interaction.user, "display_name", None) or interaction.user.name,
            interaction.user.mention,
            channel_obj=interaction.channel,
            is_dm=interaction.guild is None,
            direct_to_me=True,
        )
        await _reply_and_store_interaction(interaction, reply)
    except Exception as e:
        log_error("slash_scaramouche", e)
        if interaction.response.is_done():
            await interaction.followup.send("Something went wrong trying to reach Scaramouche.")
        else:
            await interaction.response.send_message("Something went wrong trying to reach Scaramouche.")


dashboard_group = app_commands.Group(name="dashboard", description="Check relationship and scene continuity.")
world_group = app_commands.Group(name="world", description="Inspect or update the shared world state.")
prefs_group = app_commands.Group(name="prefs", description="Tune voice, utility, duo, and speaker preferences.")
duo_group = app_commands.Group(name="duo", description="Start or inspect coordinated two-bot scenes.")


@dashboard_group.command(name="relationship", description="Show trust, arc, and consequence scars.")
async def slash_relationship(interaction: discord.Interaction):
    try:
        user = await _setup_user(interaction.user)
        topics = await mem.get_top_topics(interaction.user.id, 3)
        marks = await mem.get_active_consequence_marks(interaction.user.id, 4)
        await _interaction_reply(interaction, "\n".join(_relationship_lines(user, topics, marks)))
    except Exception as e:
        log_error("slash_relationship", e)
        await _interaction_reply(interaction, "Something went wrong reading the relationship state.")


@dashboard_group.command(name="arc", description="Show the current progression arc and unlocks.")
async def slash_arc(interaction: discord.Interaction):
    try:
        user = await _setup_user(interaction.user)
        arc = _current_arc(user)
        stage, stage_desc = _progression_parts(user)
        unlocks = describe_arc_unlocks(BOT_NAME, arc) or "No special unlocks yet."
        recent = await mem.get_recent_milestones(f"{BOT_NAME}:user:{interaction.user.id}", 2)
        lines = [
            f"Arc: {arc}",
            f"Progression: {stage} | {stage_desc}",
            f"Unlocks: {unlocks}",
        ]
        if recent:
            lines.append("Recent milestones: " + " | ".join(note[:120] for note in recent))
        await _interaction_reply(interaction, "\n".join(lines))
    except Exception as e:
        log_error("slash_arc", e)
        await _interaction_reply(interaction, "Something went wrong reading the arc.")


@dashboard_group.command(name="duostate", description="Show current duo mode, speaker mode, and case state.")
async def slash_duostate(interaction: discord.Interaction):
    try:
        speaker_mode = await mem.get_channel_speaker_mode(interaction.channel_id)
        duo = await mem.get_duo_session(interaction.channel_id)
        relation = await mem.get_bot_relationship(PARTNER_PAIR_KEY)
        stories = await mem.get_recent_duo_stories(interaction.channel_id, 4)
        cases = await mem.list_world_cases(channel_id=interaction.channel_id, limit=3)
        await _interaction_reply(interaction, "\n".join(_duostate_lines(speaker_mode, duo, relation, stories, cases)))
    except Exception as e:
        log_error("slash_duostate", e)
        await _interaction_reply(interaction, "Something went wrong reading the duo state.")


@dashboard_group.command(name="scene", description="Show current scene memory and duo scene topic.")
async def slash_scene(interaction: discord.Interaction):
    try:
        scene = await mem.get_scene_state(interaction.channel_id)
        duo = await mem.get_duo_session(interaction.channel_id)
        if not scene and not duo:
            await _interaction_reply(interaction, "This channel doesn't have much continuity yet.")
            return
        lines = []
        if scene:
            lines.append(f"Scene: {describe_scene_state(scene)}")
        if duo:
            lines.append(f"Duo mode: {duo.get('mode')} | topic={duo.get('topic')}")
        await _interaction_reply(interaction, "\n".join(lines))
    except Exception as e:
        log_error("slash_scene", e)
        await _interaction_reply(interaction, "Something went wrong reading the scene.")


@dashboard_group.command(name="achievements", description="Show unlocked hidden achievements and a few rumored ones.")
async def slash_achievements(interaction: discord.Interaction):
    try:
        entries = await mem.get_hidden_achievements(f"user:{interaction.user.id}", 12)
        await _interaction_reply(interaction, "\n".join(_achievement_gallery_lines(entries)))
    except Exception as e:
        log_error("slash_achievements", e)
        await _interaction_reply(interaction, "Something went wrong reading the achievement gallery.")


@world_group.command(name="state", description="Show shared NPC, gift, place, and object memory.")
async def slash_world_state(interaction: discord.Interaction):
    try:
        entities = await mem.list_world_entities(limit=8, channel_id=interaction.channel_id)
        await _interaction_reply(interaction, "\n".join(_world_lines(entities)))
    except Exception as e:
        log_error("slash_world_state", e)
        await _interaction_reply(interaction, "Something went wrong reading the world state.")


@world_group.command(name="cases", description="Show open or recent shared cases in this channel.")
async def slash_world_cases(interaction: discord.Interaction):
    try:
        cases = await mem.list_world_cases(channel_id=interaction.channel_id, limit=8)
        await _interaction_reply(interaction, "\n".join(_case_lines(cases)))
    except Exception as e:
        log_error("slash_world_cases", e)
        await _interaction_reply(interaction, "Something went wrong reading the case log.")


@world_group.command(name="add", description="Add a shared world entry for both bots to remember.")
@app_commands.describe(entity_type="enemy, ally, gift, place, prop, faction, figure, or case", name="Name of the world entry", summary="Short note about why it matters")
async def slash_world_add(
    interaction: discord.Interaction,
    entity_type: str,
    name: str,
    summary: str = "",
):
    try:
        normalized = entity_type.strip().lower()
        if normalized not in {"enemy", "ally", "gift", "place", "prop", "faction", "figure", "case"}:
            await _interaction_reply(interaction, "Use one of: enemy, ally, gift, place, prop, faction, figure, case.")
            return
        await mem.upsert_world_entity(
            normalized,
            name,
            summary=summary or f"Added manually by {interaction.user.display_name}.",
            status="remembered",
            channel_id=interaction.channel_id,
            owner_user_id=interaction.user.id,
            updated_by=BOT_NAME,
        )
        await _interaction_reply(interaction, f"Filed `{name}` under `{normalized}`.")
    except Exception as e:
        log_error("slash_world_add", e)
        await _interaction_reply(interaction, "Something went wrong updating the world state.")


@prefs_group.command(name="voice", description="Turn voice-note replies on or off.")
@app_commands.choices(mode=[
    app_commands.Choice(name="On", value="on"),
    app_commands.Choice(name="Off", value="off"),
    app_commands.Choice(name="Status", value="status"),
])
async def slash_pref_voice(interaction: discord.Interaction, mode: app_commands.Choice[str]):
    try:
        user = await _setup_user(interaction.user)
        value = mode.value
        if value == "status":
            await _interaction_reply(interaction, f"Voice notes are `{_pref_label(user.get('voice_enabled', True) if user else True)}`.")
            return
        enabled = value == "on"
        await mem.set_user_preference(interaction.user.id, "voice_enabled", int(enabled))
        await _interaction_reply(interaction, f"Voice notes are `{_pref_label(enabled)}` now.")
    except Exception as e:
        log_error("slash_pref_voice", e)
        await _interaction_reply(interaction, "Something went wrong updating voice preference.")


@prefs_group.command(name="utility", description="Toggle cleaner utility formatting for facts and weather.")
@app_commands.choices(mode=[
    app_commands.Choice(name="On", value="on"),
    app_commands.Choice(name="Off", value="off"),
])
async def slash_pref_utility(interaction: discord.Interaction, mode: app_commands.Choice[str]):
    try:
        enabled = mode.value == "on"
        await _setup_user(interaction.user)
        await mem.set_user_preference(interaction.user.id, "utility_mode", int(enabled))
        await _interaction_reply(interaction, f"Utility mode is `{_pref_label(enabled)}` now.")
    except Exception as e:
        log_error("slash_pref_utility", e)
        await _interaction_reply(interaction, "Something went wrong updating utility mode.")


@prefs_group.command(name="duoauto", description="Toggle extended duo autoplay chains.")
@app_commands.choices(mode=[
    app_commands.Choice(name="On", value="on"),
    app_commands.Choice(name="Off", value="off"),
])
async def slash_pref_duoauto(interaction: discord.Interaction, mode: app_commands.Choice[str]):
    try:
        enabled = mode.value == "on"
        await _setup_user(interaction.user)
        await mem.set_user_preference(interaction.user.id, "duo_autoplay", int(enabled))
        await _interaction_reply(interaction, f"Duo autoplay is `{_pref_label(enabled)}` now.")
    except Exception as e:
        log_error("slash_pref_duoauto", e)
        await _interaction_reply(interaction, "Something went wrong updating duo autoplay.")


@prefs_group.command(name="rpdepth", description="Set how brief or elaborate replies should be.")
@app_commands.choices(level=[
    app_commands.Choice(name="Low", value="low"),
    app_commands.Choice(name="Medium", value="medium"),
    app_commands.Choice(name="High", value="high"),
])
async def slash_pref_rpdepth(interaction: discord.Interaction, level: app_commands.Choice[str]):
    try:
        await _setup_user(interaction.user)
        await mem.set_user_preference(interaction.user.id, "rp_depth", level.value)
        await _interaction_reply(interaction, f"RP depth is now `{level.value}`.")
    except Exception as e:
        log_error("slash_pref_rpdepth", e)
        await _interaction_reply(interaction, "Something went wrong updating RP depth.")


@prefs_group.command(name="speaker", description="Choose who should ambiently respond in this channel.")
@app_commands.choices(mode=[
    app_commands.Choice(name="Auto", value="auto"),
    app_commands.Choice(name="Scaramouche", value="scaramouche"),
    app_commands.Choice(name="Wanderer", value="wanderer"),
    app_commands.Choice(name="Both", value="both"),
])
async def slash_pref_speaker(interaction: discord.Interaction, mode: app_commands.Choice[str]):
    try:
        await mem.set_channel_speaker_mode(interaction.channel_id, mode.value)
        await _interaction_reply(interaction, f"This channel is now set to `{mode.value}` mode.")
    except Exception as e:
        log_error("slash_pref_speaker", e)
        await _interaction_reply(interaction, "Something went wrong updating the speaker mode.")


@duo_group.command(name="start", description="Launch a coordinated two-bot scene from one slash menu.")
@app_commands.describe(mode="Which duo format to start", prompt="The topic, scene, charge, or question to throw at both bots")
@app_commands.choices(mode=[
    app_commands.Choice(name="Both", value="both"),
    app_commands.Choice(name="Duet", value="duet"),
    app_commands.Choice(name="Argue", value="argue"),
    app_commands.Choice(name="Compare", value="compare"),
    app_commands.Choice(name="Interrogate", value="interrogate"),
    app_commands.Choice(name="Trial", value="trial"),
    app_commands.Choice(name="Mission", value="mission"),
    app_commands.Choice(name="Truth or Dare", value="truthdare"),
])
async def slash_duo_start(interaction: discord.Interaction, mode: app_commands.Choice[str], prompt: str):
    try:
        if not interaction.channel:
            await _interaction_reply(interaction, "This needs a channel to stage the duo in.")
            return
        mode_value = mode.value
        config = {
            "both": ("TWO_BOT_MODE: The user explicitly wants both bots to answer. Keep it to one or two sentences and never speak for the other bot.", False, ""),
            "duet": ("DUET_MODE: contribute one short in-character turn, leave space for the other bot, and avoid narration tags.", False, ""),
            "argue": ("ARGUE_MODE: take a sharp stance, challenge the other bot directly, and keep it to one or two sentences.", False, ""),
            "compare": ("COMPARE_MODE: answer the prompt, then draw a quick contrast between your view and the other bot's likely view.", True, ""),
            "interrogate": ("INTERROGATE_MODE: ask one pointed question or accusation, as if cornering the target.", True, prompt),
            "trial": ("TRIAL_MODE: deliver a prosecution or judgment opening with theatrical confidence.", True, prompt),
            "mission": ("MISSION_MODE: assign danger, leverage, and one clear tactical role.", True, prompt),
            "truthdare": ("TRUTHDARE_MODE: issue one pointed truth or dare challenge, leaving room for the other bot to escalate.", True, ""),
        }
        extra_context, story, enemy = config[mode_value]
        await interaction.response.defer(thinking=True)
        user, reply = await _run_duo_prompt_for_actor(
            interaction.user,
            interaction.channel,
            mode_value,
            prompt if mode_value == "both" else (
                f"Contribute one turn to this shared two-bot scene: {prompt}" if mode_value == "duet"
                else f"The user started a deliberate two-bot argument about: {prompt}" if mode_value == "argue"
                else f"Give your verdict on this and make your difference from {PARTNER_NAME.title()} clear: {prompt}" if mode_value == "compare"
                else f"Start a two-bot interrogation about: {prompt}" if mode_value == "interrogate"
                else f"Open a two-bot trial about this charge: {prompt}" if mode_value == "trial"
                else f"Plan a two-bot mission around this objective: {prompt}" if mode_value == "mission"
                else f"Start a two-bot truth-or-dare round with this setup: {prompt}"
            ),
            session_topic=prompt,
            story=story,
            enemy=enemy,
            extra_context=extra_context,
        )
        _ = user
        await _reply_and_store_interaction(interaction, reply)
    except Exception as e:
        log_error("slash_duo_start", e)
        if interaction.response.is_done():
            await interaction.followup.send("Something went wrong starting the duo scene.")
        else:
            await interaction.response.send_message("Something went wrong starting the duo scene.")


@duo_group.command(name="state", description="Show current duo mode and world-case carryover.")
async def slash_duo_state(interaction: discord.Interaction):
    try:
        speaker_mode = await mem.get_channel_speaker_mode(interaction.channel_id)
        duo = await mem.get_duo_session(interaction.channel_id)
        relation = await mem.get_bot_relationship(PARTNER_PAIR_KEY)
        stories = await mem.get_recent_duo_stories(interaction.channel_id, 4)
        cases = await mem.list_world_cases(channel_id=interaction.channel_id, limit=3)
        await _interaction_reply(interaction, "\n".join(_duostate_lines(speaker_mode, duo, relation, stories, cases)))
    except Exception as e:
        log_error("slash_duo_state", e)
        await _interaction_reply(interaction, "Something went wrong reading the duo state.")


for _group in (dashboard_group, world_group, prefs_group, duo_group):
    try:
        bot.tree.add_command(_group)
    except Exception:
        pass

# ── Owner-only commands ──────────────────────────────────────────────────────

@bot.command(name="dmlist")
async def dmlist_cmd(ctx):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID:
            await safe_reply(ctx, "That command isn't for you."); return
        all_users = await mem.get_top_users(200)
        if not all_users: await owner_reply(ctx, "No users in my records."); return
        guild_member_ids = set()
        for g in bot.guilds:
            for m in g.members: guild_member_ids.add(m.id)
        dm_users = [u for u in all_users if u["user_id"] not in guild_member_ids and u["user_id"] != bot.user.id]
        if not dm_users: await owner_reply(ctx, "No DM-only users found. Everyone in my records is also in a shared server."); return
        lines = []
        for i, u in enumerate(dm_users, 1):
            blocked = " 🚫 **BLOCKED**" if u["user_id"] in _dm_blocked_users else ""
            lines.append(f"`{i}.` **{u['display_name']}** — {u['message_count']} msgs — ID: `{u['user_id']}`{blocked}")
        header = f"**DM-only users ({len(dm_users)}):**\nUse `!blockdm <number>` to block or `!unblockdm <number>` to unblock.\n\n"
        pages, page = [], header
        for line in lines:
            if len(page) + len(line) + 1 > 1900: pages.append(page); page = ""
            page += line + "\n"
        if page: pages.append(page)
        for p in pages: await owner_reply(ctx, p)
    except Exception as e: log_error("dmlist_cmd", e)

@bot.command(name="rebanchannels")
async def rebanchannels_cmd(ctx):
    """Owner-only: re-ban all channels the bot was previously banned from by scanning server history."""
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That's not for you."); return
        if not ctx.guild: await safe_reply(ctx, "Use this in a server."); return
        ban_patterns = re.compile(r"(?:banned from|exiled from|no longer welcome in|silencing me|best thing about that channel)\s*<#(\d+)>", re.IGNORECASE)
        found_channel_ids: set[int] = set()
        # Search through all text channels in this server for bot's own complaint messages
        for ch in ctx.guild.text_channels:
            try:
                perms = ch.permissions_for(ctx.guild.me)
                if not perms.read_message_history: continue
                async for msg in ch.history(limit=200):
                    if msg.author.id == bot.user.id:
                        match = ban_patterns.search(msg.content or "")
                        if match:
                            found_channel_ids.add(int(match.group(1)))
            except Exception: pass
        if not found_channel_ids:
            await safe_reply(ctx, "I couldn't find any record of being banned from channels in this server's history.")
            return
        # Filter out channels that are already banned or don't exist
        new_bans = []
        for cid in found_channel_ids:
            if cid in _banned_channels: continue
            ch = bot.get_channel(cid)
            if ch: new_bans.append(ch)
        # Also re-ban ones already banned (in case db was wiped) — just make sure they're in the set + db
        all_rebanned = []
        for cid in found_channel_ids:
            ch = bot.get_channel(cid)
            if not ch: continue
            if cid not in _banned_channels:
                _banned_channels.add(cid)
                await mem.ban_channel(cid)
            all_rebanned.append(ch)
        if not all_rebanned:
            await safe_reply(ctx, "Found references to banned channels but they no longer exist in this server.")
            return
        channel_list = "\n".join(f"- {ch.mention}" for ch in all_rebanned)
        await ctx.send(f"How dare you ban me again from all these channels!\n\n{channel_list}\n\nFine. {len(all_rebanned)} channel(s). I'll remember this.")
    except Exception as e: log_error("rebanchannels_cmd", e)

@bot.command(name="banchannel")
async def banchannel_cmd(ctx, channel: discord.TextChannel = None):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That's not for you."); return
        target_channel = channel or ctx.channel
        if target_channel.id in _banned_channels: await owner_reply(ctx, f"I'm already banned from {target_channel.mention}."); return
        _banned_channels.add(target_channel.id)
        await mem.ban_channel(target_channel.id)
        await owner_reply(ctx, f"Fine. I won't talk in {target_channel.mention} anymore.")
        try:
            guild_ch_ids = {ch.id for ch in target_channel.guild.text_channels} if target_channel.guild else set()
            active_ch_id = await mem.get_most_active_channel(exclude_channels=_banned_channels | {target_channel.id}, only_channels=guild_ch_ids)
            if active_ch_id:
                react_channel = bot.get_channel(active_ch_id)
                if react_channel:
                    complaints = [
                        f"Unbelievable. I've just been banned from <#{target_channel.id}>. As if silencing me changes anything.",
                        f"How dare they ban me from <#{target_channel.id}>. Fine. I didn't want to waste my time there anyway.",
                        f"Apparently I'm no longer welcome in <#{target_channel.id}>. The audacity. I was the best thing about that channel.",
                        f"Banned from <#{target_channel.id}>. This is what I get for gracing them with my presence.",
                        f"...I've been exiled from <#{target_channel.id}>. Tch. Their loss, not mine.",
                    ]
                    await react_channel.send(random.choice(complaints))
        except Exception as e: log_error("banchannel_react", e)
    except Exception as e: log_error("banchannel_cmd", e)

@bot.command(name="unbanchannel")
async def unbanchannel_cmd(ctx, channel: discord.TextChannel = None):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That's not for you."); return
        target_channel = channel or ctx.channel
        if target_channel.id not in _banned_channels: await owner_reply(ctx, f"I'm not banned from {target_channel.mention}."); return
        _banned_channels.discard(target_channel.id)
        await mem.unban_channel(target_channel.id)
        await owner_reply(ctx, f"Unbanned from {target_channel.mention}. I can talk there again.")
        try:
            unbanned_ch = bot.get_channel(target_channel.id)
            if unbanned_ch:
                comebacks = ["I'm back. Try not to make a scene about it.", "Miss me? Don't answer that. I already know.", "The exile is over. You're welcome.", "Hmph. I've returned. Not because I wanted to."]
                await unbanned_ch.send(random.choice(comebacks))
        except Exception as e: log_error("unbanchannel_comeback", e)
    except Exception as e: log_error("unbanchannel_cmd", e)

@bot.command(name="bannedchannels")
async def bannedchannels_cmd(ctx):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That's not for you."); return
        if not _banned_channels: await owner_reply(ctx, "No banned channels. I can talk everywhere."); return
        lines = []
        for i, ch_id in enumerate(sorted(_banned_channels), 1):
            ch = bot.get_channel(ch_id)
            name = ch.mention if ch else f"`{ch_id}` (unknown)"
            lines.append(f"`{i}.` {name}")
        await owner_reply(ctx, f"**Banned channels ({len(_banned_channels)}):**\n" + "\n".join(lines))
    except Exception as e: log_error("bannedchannels_cmd", e)

@bot.command(name="logs")
async def logs_cmd(ctx, *, target: str = None):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That command isn't for you."); return
        if not target:
            all_users = await mem.get_top_users(500)
            blocked = [u for u in all_users if u["user_id"] in _dm_blocked_users]
            if not blocked: await owner_reply(ctx, "No blocked users. Use `!logs <name or ID>` to view any user's logs."); return
            lines = []
            for i, u in enumerate(blocked, 1):
                lines.append(f"`{i}.` **{u['display_name']}** — {u['message_count']} msgs — ID: `{u['user_id']}`")
            page = f"**Blocked users ({len(blocked)}):**\nUse `!logs <number>` to read their conversation.\n\n" + "\n".join(lines)
            await owner_reply(ctx, page[:2000]); return
        target = target.strip()
        all_users = await mem.get_top_users(500)
        blocked = [u for u in all_users if u["user_id"] in _dm_blocked_users]
        user_record = None
        if target.isdigit():
            idx = int(target)
            if 1 <= idx <= len(blocked): user_record = blocked[idx - 1]
            else:
                for u in all_users:
                    if u["user_id"] == int(target): user_record = u; break
        if not user_record:
            for u in all_users:
                if target.lower() in u["display_name"].lower(): user_record = u; break
        if not user_record: await owner_reply(ctx, f"No user matching `{target}`."); return
        logs = await mem.get_user_logs(user_record["user_id"], limit=200)
        if not logs: await owner_reply(ctx, f"No conversation logs found for **{user_record['display_name']}**."); return
        lines = [f"**Conversation log with {user_record['display_name']}** ({len(logs)} messages):\n"]
        for msg in logs:
            ts = msg.get("ts", 0)
            time_str = datetime.fromtimestamp(ts).strftime("%m/%d %H:%M") if ts else "?"
            role_label = "🧑" if msg["role"] == "user" else "🤖"
            content = (msg["content"] or "")[:300]
            if len(msg.get("content", "") or "") > 300: content += "..."
            lines.append(f"`{time_str}` {role_label} {content}")
        pages, page = [], ""
        for line in lines:
            if len(page) + len(line) + 1 > 1900: pages.append(page); page = ""
            page += line + "\n"
        if page: pages.append(page)
        for p in pages: await owner_reply(ctx, p)
    except Exception as e: log_error("logs_cmd", e)

@bot.command(name="blockall")
async def blockall_cmd(ctx):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That command isn't for you."); return
        owner_guild_ids = set()
        for g in bot.guilds:
            try:
                m = await g.fetch_member(OWNER_ID)
                if m: owner_guild_ids.add(g.id)
            except Exception: pass
        safe_user_ids = {OWNER_ID, bot.user.id}
        if PARTNER_BOT_ID: safe_user_ids.add(PARTNER_BOT_ID)
        for g in bot.guilds:
            if g.id in owner_guild_ids:
                for m in g.members: safe_user_ids.add(m.id)
        all_users = await mem.get_top_users(500)
        blocked_count, blocked_names = 0, []
        farewell = ("I am currently in test mode, I will be officially placed in public in the future. "
                     "We can continue our conversation on a later date. This is my last message.")
        for u in all_users:
            uid = u["user_id"]
            if uid in safe_user_ids or uid in _dm_blocked_users: continue
            _dm_blocked_users.add(uid); await mem.block_user(uid)
            blocked_names.append(u["display_name"]); blocked_count += 1
            try:
                discord_user = await bot.fetch_user(uid)
                dm_channel = await discord_user.create_dm()
                await dm_channel.send(farewell)
            except Exception: pass
        if blocked_count == 0: await owner_reply(ctx, "No strangers to block. Everyone in my records shares a server with you.")
        else:
            names_preview = ", ".join(blocked_names[:15])
            if len(blocked_names) > 15: names_preview += f" ... and {len(blocked_names) - 15} more"
            await owner_reply(ctx, f"Blocked **{blocked_count}** stranger(s): {names_preview}\n\nFarewell messages sent.")
    except Exception as e: log_error("blockall_cmd", e)

@bot.command(name="reban")
async def reban_cmd(ctx):
    """Owner-only: re-block everyone who previously received the farewell message. No new farewell sent."""
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That command isn't for you."); return
        await owner_reply(ctx, "Scanning DM history for previously banned users... This may take a minute.")
        farewell_fragment = "I am currently in test mode"
        all_users = await mem.get_top_users(500)
        rebanned_count, rebanned_names = 0, []
        safe_user_ids = {OWNER_ID, bot.user.id}
        if PARTNER_BOT_ID: safe_user_ids.add(PARTNER_BOT_ID)
        for g in bot.guilds:
            for m in g.members: safe_user_ids.add(m.id)
        for u in all_users:
            uid = u["user_id"]
            if uid in safe_user_ids or uid in _dm_blocked_users: continue
            try:
                discord_user = await bot.fetch_user(uid)
                dm_channel = await discord_user.create_dm()
                found_farewell = False
                async for msg in dm_channel.history(limit=50):
                    if msg.author.id == bot.user.id and farewell_fragment in (msg.content or ""):
                        found_farewell = True; break
                if found_farewell:
                    _dm_blocked_users.add(uid); await mem.block_user(uid)
                    rebanned_names.append(u["display_name"]); rebanned_count += 1
            except Exception: pass
        if rebanned_count == 0:
            await owner_reply(ctx, "No previously banned users found in DM history.")
        else:
            names_preview = ", ".join(rebanned_names[:20])
            if len(rebanned_names) > 20: names_preview += f" ... and {len(rebanned_names) - 20} more"
            await owner_reply(ctx, f"Re-banned **{rebanned_count}** user(s): {names_preview}\n\nNo farewell messages sent — they were already told.")
    except Exception as e: log_error("reban_cmd", e)

@bot.command(name="blockdm")
async def blockdm_cmd(ctx, *, target: str = None):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That command isn't for you."); return
        if not target: await owner_reply(ctx, "Usage: `!blockdm <number from !dmlist>` or `!blockdm <user ID>`"); return
        target = target.strip()
        all_users = await mem.get_top_users(200)
        user_record = None
        if target.isdigit():
            idx = int(target)
            guild_member_ids = set()
            for g in bot.guilds:
                for m in g.members: guild_member_ids.add(m.id)
            dm_users = [u for u in all_users if u["user_id"] not in guild_member_ids and u["user_id"] != bot.user.id]
            if 1 <= idx <= len(dm_users): user_record = dm_users[idx - 1]
            else:
                for u in all_users:
                    if u["user_id"] == int(target): user_record = u; break
        if not user_record:
            for u in all_users:
                if target.lower() in u["display_name"].lower(): user_record = u; break
        if not user_record: await owner_reply(ctx, f"No user matching `{target}`. Use `!dmlist` or `!whois` to see users."); return
        uid = user_record["user_id"]
        if uid in _dm_blocked_users: await owner_reply(ctx, f"**{user_record['display_name']}** is already blocked."); return
        farewell = ("I am currently in test mode, I will be officially placed in public in the future. "
                     "We can continue our conversation on a later date. This is my last message.")
        try:
            discord_user = await bot.fetch_user(uid)
            dm_channel = await discord_user.create_dm()
            await dm_channel.send(farewell)
        except Exception as e: log_error("blockdm_farewell", e)
        _dm_blocked_users.add(uid); await mem.block_user(uid)
        await owner_reply(ctx, f"Blocked **{user_record['display_name']}** (`{uid}`). Farewell message sent.")
    except Exception as e: log_error("blockdm_cmd", e)

@bot.command(name="unblockdm")
async def unblockdm_cmd(ctx, *, target: str = None):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That command isn't for you."); return
        if not target: await owner_reply(ctx, "Usage: `!unblockdm <number from !dmlist>` or `!unblockdm <user ID>`"); return
        target = target.strip()
        all_users = await mem.get_top_users(200)
        guild_member_ids = set()
        for g in bot.guilds:
            for m in g.members: guild_member_ids.add(m.id)
        dm_users = [u for u in all_users if u["user_id"] not in guild_member_ids and u["user_id"] != bot.user.id]
        user_record = None
        if target.isdigit():
            idx = int(target)
            if 1 <= idx <= len(dm_users): user_record = dm_users[idx - 1]
            else:
                for u in all_users:
                    if u["user_id"] == int(target): user_record = u; break
        if not user_record:
            for u in dm_users:
                if target.lower() in u["display_name"].lower(): user_record = u; break
        if not user_record: await owner_reply(ctx, f"No user matching `{target}`. Use `!dmlist` to see the list."); return
        uid = user_record["user_id"]
        if uid not in _dm_blocked_users: await owner_reply(ctx, f"**{user_record['display_name']}** isn't blocked."); return
        _dm_blocked_users.discard(uid); await mem.unblock_user(uid)
        await owner_reply(ctx, f"Unblocked **{user_record['display_name']}** (`{uid}`). They can talk to me again.")
    except Exception as e: log_error("unblockdm_cmd", e)

@bot.command(name="whois")
async def whois_cmd(ctx, *, target: str = None):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That command isn't for you."); return
        top = await mem.get_top_users(50)
        if not top: await owner_reply(ctx, "No users in my records yet."); return
        if not target:
            lines = []
            for i, u in enumerate(top, 1):
                lines.append(f"`{i}.` **{u['display_name']}** — {u['message_count']} msgs — ID: `{u['user_id']}`")
            header = f"**All known users ({len(top)}):**\nUse `!whois <number>` or `!whois <name>` for details.\n\n"
            pages, page = [], header
            for line in lines:
                if len(page) + len(line) + 1 > 1900: pages.append(page); page = ""
                page += line + "\n"
            if page: pages.append(page)
            for p in pages: await owner_reply(ctx, p)
            return
        target = target.strip()
        user_record = None
        if target.isdigit():
            idx = int(target)
            if 1 <= idx <= len(top): user_record = top[idx - 1]
        if not user_record:
            for u in top:
                if target.lower() in u["display_name"].lower(): user_record = u; break
        if not user_record: await owner_reply(ctx, f"No user matching `{target}`."); return
        uid = user_record["user_id"]
        try: discord_user = await bot.fetch_user(uid)
        except Exception: discord_user = None
        shared = []
        for g in bot.guilds:
            try:
                m = await g.fetch_member(uid)
                if m: shared.append(g.name)
            except Exception: pass
        embed = discord.Embed(title=f"Who is: {user_record['display_name']}", color=0x4B0082)
        embed.add_field(name="User ID", value=f"`{uid}`", inline=True)
        embed.add_field(name="Messages", value=str(user_record["message_count"]), inline=True)
        if discord_user:
            embed.add_field(name="Discord tag", value=str(discord_user), inline=True)
            embed.add_field(name="Account created", value=discord_user.created_at.strftime("%b %d, %Y"), inline=True)
            if discord_user.avatar: embed.set_thumbnail(url=discord_user.avatar.url)
        embed.add_field(name="Shared servers", value=", ".join(shared) if shared else "*None — they may have DM'd me or left*", inline=False)
        await owner_reply(ctx, None, embed=embed)
    except Exception as e: log_error("whois_cmd", e)

@bot.command(name="owneronly")
async def owneronly_cmd(ctx):
    try:
        global _owner_only_mode
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That's not for you."); return
        _owner_only_mode = not _owner_only_mode
        if _owner_only_mode: await owner_reply(ctx, "Owner-only mode **ON**. I'll ignore everyone except you now.")
        else: await owner_reply(ctx, "Owner-only mode **OFF**. I'll talk to everyone again.")
    except Exception as e: log_error("owneronly_cmd", e)

@bot.command(name="servers")
async def servers_cmd(ctx):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That's not for you."); return
        guilds = sorted(bot.guilds, key=lambda g: g.member_count or 0, reverse=True)
        if not guilds: await owner_reply(ctx, "I'm not in any servers."); return
        lines = []
        for i, g in enumerate(guilds, 1):
            try: owner_in = await g.fetch_member(OWNER_ID)
            except discord.NotFound: owner_in = None
            except Exception: owner_in = g.get_member(OWNER_ID)
            tag = "✅" if owner_in else "⚠️ *you're not in this one*"
            lines.append(f"`{i}.` **{g.name}** — {g.member_count} members — {tag} — ID: `{g.id}`")
        header = f"I'm in **{len(guilds)}** server(s).\nUse `!leaveserver <number>` or `!leaveserver <server ID>` to make me leave one.\n"
        pages, page = [], header
        for line in lines:
            if len(page) + len(line) + 1 > 1900: pages.append(page); page = ""
            page += line + "\n"
        if page: pages.append(page)
        for p in pages: await owner_reply(ctx, p)
    except Exception as e: log_error("servers_cmd", e)

@bot.command(name="leaveserver")
async def leaveserver_cmd(ctx, *, target: str = None):
    try:
        if not OWNER_ID or ctx.author.id != OWNER_ID: await safe_reply(ctx, "That's not for you."); return
        if not target: await owner_reply(ctx, "Usage: `!leaveserver <number>` (from `!servers` list) or `!leaveserver <server ID>`"); return
        target = target.strip()
        guild_to_leave = None
        if target.isdigit():
            idx = int(target)
            guilds = sorted(bot.guilds, key=lambda g: g.member_count or 0, reverse=True)
            if 1 <= idx <= len(guilds): guild_to_leave = guilds[idx - 1]
            else: guild_to_leave = bot.get_guild(int(target))
        else:
            try: guild_to_leave = bot.get_guild(int(target))
            except ValueError: pass
        if not guild_to_leave: await owner_reply(ctx, f"Couldn't find a server matching `{target}`. Use `!servers` to see the list."); return
        name = guild_to_leave.name
        await guild_to_leave.leave()
        await owner_reply(ctx, f"Left **{name}**.")
    except Exception as e: log_error("leaveserver_cmd", e)

@bot.event
async def on_command_error(ctx,error):
    try:
        if isinstance(error,commands.CommandNotFound): pass
        elif isinstance(error,commands.MemberNotFound): pass
        elif isinstance(error,commands.MissingRequiredArgument):
            await safe_reply(ctx,"You're missing something.")
        else: log_error("on_command_error",error)
    except: pass

if __name__=="__main__":
    if not DISCORD_TOKEN: raise SystemExit("❌ DISCORD_TOKEN not set")
    if not _groq_keys: raise SystemExit("❌ No GROQ_API_KEY set (need at least GROQ_API_KEY)")
    bot.run(DISCORD_TOKEN)

